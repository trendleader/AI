import streamlit as st
import streamlit.components.v1 as components
import anthropic
import requests
from bs4 import BeautifulSoup
import json
import re
import os
import subprocess
import tempfile
import time
from pathlib import Path
from avatar_component import get_avatar_html

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResumeIQ · Interview Suite",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500&display=swap');

:root {
    --bg:      #0a0b0f;
    --surface: #13151c;
    --border:  #1e2130;
    --accent:  #6c63ff;
    --accent2: #00d4aa;
    --danger:  #ff4d6d;
    --warn:    #ffa940;
    --text:    #e8eaf0;
    --muted:   #6b7280;
}

html, body, [data-testid="stAppViewContainer"] {
    background: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif;
}
[data-testid="stHeader"] { background: transparent !important; }
h1,h2,h3,h4 { font-family: 'Syne', sans-serif !important; color: var(--text) !important; }

.hero {
    text-align: center;
    padding: 2.5rem 1rem 1.8rem;
    background: radial-gradient(ellipse 80% 60% at 50% -10%, rgba(108,99,255,0.22) 0%, transparent 70%);
    border-bottom: 1px solid var(--border);
    margin-bottom: 2rem;
}
.hero-badge {
    display: inline-block;
    background: rgba(108,99,255,0.15);
    border: 1px solid rgba(108,99,255,0.4);
    color: #a29bfe;
    font-size: 0.7rem; font-weight: 600;
    letter-spacing: 0.15em; text-transform: uppercase;
    padding: 4px 14px; border-radius: 100px; margin-bottom: 1rem;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: clamp(1.8rem, 4vw, 3rem); font-weight: 800; line-height: 1.1;
    margin: 0 0 0.8rem;
    background: linear-gradient(135deg, #fff 30%, #a29bfe 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.hero-sub { color: var(--muted); font-size: 1rem; max-width: 560px; margin: 0 auto; line-height: 1.6; }

/* Module nav pills */
.mod-nav { display: flex; gap: 0.6rem; flex-wrap: wrap; justify-content: center; margin: 1.5rem 0; }
.mod-pill {
    padding: 0.45rem 1.2rem; border-radius: 100px;
    border: 1px solid var(--border);
    background: var(--surface);
    color: var(--muted);
    font-family: 'Syne', sans-serif; font-size: 0.78rem; font-weight: 600;
    cursor: pointer; transition: all 0.15s;
}
.mod-pill.active {
    background: rgba(108,99,255,0.18);
    border-color: rgba(108,99,255,0.5);
    color: #a29bfe;
}

/* Cards */
.card {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 14px; padding: 1.4rem; margin-bottom: 1rem;
}
.card:hover { border-color: rgba(108,99,255,0.3); }

.section-hdr {
    font-family: 'Syne', sans-serif; font-size: 0.65rem; font-weight: 700;
    letter-spacing: 0.2em; text-transform: uppercase; color: var(--muted);
    margin: 1.2rem 0 0.6rem; padding-bottom: 0.35rem;
    border-bottom: 1px solid var(--border);
}

/* Pills */
.pill-row { display: flex; flex-wrap: wrap; gap: 0.45rem; margin: 0.5rem 0; }
.pill { display: inline-flex; align-items: center; gap: 5px; padding: 3px 11px; border-radius: 100px; font-size: 0.76rem; font-weight: 500; }
.pill-green  { background: rgba(0,212,170,0.12);  border: 1px solid rgba(0,212,170,0.35);  color: #00d4aa; }
.pill-red    { background: rgba(255,77,109,0.12);  border: 1px solid rgba(255,77,109,0.35);  color: #ff4d6d; }
.pill-yellow { background: rgba(255,169,64,0.12);  border: 1px solid rgba(255,169,64,0.35);  color: #ffa940; }
.pill-purple { background: rgba(108,99,255,0.12);  border: 1px solid rgba(108,99,255,0.35);  color: #a29bfe; }

/* Rec cards */
.rec-card {
    background: rgba(108,99,255,0.05); border: 1px solid rgba(108,99,255,0.18);
    border-left: 3px solid var(--accent); border-radius: 0 10px 10px 0;
    padding: 0.85rem 1rem; margin-bottom: 0.6rem; font-size: 0.88rem; line-height: 1.55;
}
.rec-title { font-weight: 600; color: #a29bfe; font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.08em; margin-bottom: 0.25rem; }

/* Chat bubbles */
.chat-wrap { display: flex; flex-direction: column; gap: 0.8rem; margin: 1rem 0; }
.bubble {
    max-width: 82%; padding: 0.85rem 1.1rem; border-radius: 14px;
    font-size: 0.9rem; line-height: 1.6;
}
.bubble-ai {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 4px 14px 14px 14px; align-self: flex-start;
    color: var(--text);
}
.bubble-user {
    background: rgba(108,99,255,0.18); border: 1px solid rgba(108,99,255,0.3);
    border-radius: 14px 4px 14px 14px; align-self: flex-end;
    color: #c8caff;
}
.bubble-label { font-size: 0.68rem; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; margin-bottom: 0.3rem; }
.bubble-ai .bubble-label { color: var(--accent2); }
.bubble-user .bubble-label { color: #a29bfe; text-align: right; }

/* Feedback card */
.feedback-card {
    background: rgba(0,212,170,0.06); border: 1px solid rgba(0,212,170,0.25);
    border-left: 3px solid var(--accent2); border-radius: 0 10px 10px 0;
    padding: 1rem 1.1rem; margin: 0.8rem 0;
}
.feedback-score {
    font-family: 'Syne', sans-serif; font-size: 2rem; font-weight: 800; color: var(--accent2);
}

/* Exam */
.q-card {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 12px; padding: 1.2rem 1.4rem; margin-bottom: 1rem;
}
.q-num { font-size: 0.68rem; font-weight: 700; color: var(--accent); letter-spacing: 0.15em; text-transform: uppercase; margin-bottom: 0.4rem; }
.q-text { font-size: 0.95rem; color: var(--text); font-weight: 500; margin-bottom: 0.8rem; line-height: 1.5; }
.answer-correct { background: rgba(0,212,170,0.1) !important; border-color: rgba(0,212,170,0.5) !important; }
.answer-wrong   { background: rgba(255,77,109,0.1) !important; border-color: rgba(255,77,109,0.5) !important; }

/* Score ring */
.score-big { font-family: 'Syne', sans-serif; font-size: 4rem; font-weight: 800; line-height: 1; }

/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, var(--accent), #8b5cf6) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 600 !important;
    font-size: 0.88rem !important; padding: 0.6rem 1.6rem !important;
    transition: opacity 0.2s, transform 0.1s !important;
}
.stButton > button:hover { opacity: 0.85 !important; transform: translateY(-1px) !important; }
.stDownloadButton > button {
    background: linear-gradient(135deg, var(--accent2), #00b894) !important;
    color: #0a0b0f !important; border: none !important; border-radius: 10px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 700 !important; font-size: 0.88rem !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: var(--surface) !important; border: 1px solid var(--border) !important;
    border-radius: 10px !important; color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 3px rgba(108,99,255,0.18) !important;
}
[data-testid="stFileUploader"] { background: var(--surface) !important; border: 1.5px dashed var(--border) !important; border-radius: 12px !important; }
.stProgress > div > div { background: var(--accent) !important; }
hr { border-color: var(--border) !important; margin: 1.8rem 0 !important; }
.stTabs [data-baseweb="tab-list"] { background: var(--surface); border-radius: 10px; padding: 4px; border: 1px solid var(--border); gap: 4px; }
.stTabs [data-baseweb="tab"] { border-radius: 8px !important; color: var(--muted) !important; font-family: 'Syne', sans-serif !important; font-weight: 600 !important; }
.stTabs [aria-selected="true"] { background: var(--accent) !important; color: white !important; }
details { background: var(--surface) !important; border-radius: 10px !important; border: 1px solid var(--border) !important; }
summary { color: var(--text) !important; }
.stAlert { border-radius: 10px !important; }
.resume-preview {
    background: var(--surface); border: 1px solid var(--border); border-radius: 12px;
    padding: 1.8rem; font-size: 0.86rem; line-height: 1.7; white-space: pre-wrap;
    max-height: 550px; overflow-y: auto;
}
/* Radio button styling */
.stRadio > div { gap: 0.4rem !important; }
.stRadio label { color: var(--text) !important; }
</style>
""", unsafe_allow_html=True)

# ── Anthropic client ────────────────────────────────────────────────────────────
@st.cache_resource
def get_client():
    # Streamlit Cloud: key lives in st.secrets["ANTHROPIC_API_KEY"]
    # Local dev: falls back to ANTHROPIC_API_KEY environment variable
    try:
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    except (FileNotFoundError, KeyError):
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        st.error("⚠️ No Anthropic API key found. Add it to `.streamlit/secrets.toml` or set the `ANTHROPIC_API_KEY` environment variable.")
        st.stop()
    return anthropic.Anthropic(api_key=api_key)

# ══════════════════════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def extract_resume_text(uploaded_file) -> str:
    """Extract plain text from .docx or .txt — no pandoc required."""
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".docx":
        # Try pandoc first (best quality), fall back to python-docx
        try:
            import shutil
            if shutil.which("pandoc"):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                    f.write(uploaded_file.getvalue())
                    tmp_path = f.name
                try:
                    result = subprocess.run(
                        ["pandoc", tmp_path, "-t", "plain"],
                        capture_output=True, text=True, timeout=15
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        return result.stdout.strip()
                finally:
                    os.unlink(tmp_path)
        except Exception:
            pass
        # python-docx fallback (works on Windows without pandoc)
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(uploaded_file.getvalue()))
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
            return "\n".join(lines)
        except Exception as e:
            raise RuntimeError(f"Could not read .docx file: {e}")
    else:
        # Plain text
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

def scrape_job_description(url: str) -> str:
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script","style","nav","footer","header","aside"]):
            tag.decompose()
        lines = [l.strip() for l in soup.get_text(separator="\n").splitlines() if l.strip()]
        return "\n".join(lines)[:8000]
    except Exception as e:
        return f"ERROR: {e}"

def score_color(score):
    if score >= 70: return "#00d4aa"
    if score >= 45: return "#ffa940"
    return "#ff4d6d"

def clean_json(raw: str) -> dict:
    raw = re.sub(r"^```json\s*", "", raw.strip())
    raw = re.sub(r"^```\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)

def build_docx(resume_text: str) -> bytes:
    """Convert plain text resume to .docx using python-docx (no pandoc needed)."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()

        # Narrow margins for a clean resume look
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.9)
            section.right_margin  = Inches(0.9)

        for line in resume_text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # ALL CAPS lines → section header style
            if stripped == stripped.upper() and len(stripped) > 3 and not stripped[0].isdigit():
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
                # Underline as divider
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
            # Bullet lines
            elif stripped.startswith("- ") or stripped.startswith("• "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(stripped.lstrip("-•").strip())
                run.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception:
        # Last resort: return UTF-8 bytes (user gets a .docx that opens as text)
        return resume_text.encode("utf-8")

def _build_docx_pandoc_unused(resume_text: str) -> bytes:
    """Kept for reference — pandoc path not used on Windows."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as f:
        f.write(resume_text)
        txt_path = f.name
    out_path = txt_path.replace(".txt", ".docx")
    try:
        subprocess.run(["pandoc", txt_path, "-o", out_path], capture_output=True)
        with open(out_path, "rb") as fh:
            return fh.read()
    except Exception:
        return resume_text.encode("utf-8")
    finally:
        for p in [txt_path, out_path]:
            try: os.unlink(p)
            except: pass

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 1 — RESUME GAP ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

def run_analysis(resume_text, jd_text) -> dict:
    client = get_client()
    prompt = f"""You are an expert resume analyst. Analyze this resume vs job description. Return ONLY raw JSON — no markdown fences.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

Return this exact JSON:
{{
  "match_score": <integer 0-100>,
  "role_title": "<role title>",
  "company": "<company name>",
  "summary": "<2-3 sentence executive summary>",
  "strengths": ["<strength 1>", ...],
  "gaps": ["<gap 1>", ...],
  "partial_matches": ["<partial 1>", ...],
  "ats_keywords_missing": ["<kw1>", ...],
  "ats_keywords_present": ["<kw1>", ...],
  "recommendations": [{{"area": "<area>", "suggestion": "<suggestion>"}}, ...],
  "rejection_risks": ["<risk 1>", ...]
}}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)

def generate_tailored_resume(resume_text, jd_text, analysis) -> str:
    gaps = "\n".join(f"- {g}" for g in analysis.get("gaps", []))
    kws  = ", ".join(analysis.get("ats_keywords_missing", []))
    recs = "\n".join(f"- [{r['area']}] {r['suggestion']}" for r in analysis.get("recommendations", []))
    prompt = f"""You are an expert resume writer. Rewrite this resume to align with the job description.
Keep all facts exact — no fabrication. Reframe bullets, weave in missing keywords naturally, rewrite the summary.
Output ONLY the resume text, no commentary. Use ALL CAPS for section headers, dashes for bullets.

ORIGINAL RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

GAPS: {gaps}
MISSING KEYWORDS: {kws}
RECOMMENDATIONS: {recs}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    draft = msg.content[0].text.strip()
    # ── Humanization pass ─────────────────────────────────────────────────────
    return humanize_resume_bullets(draft)

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 2 — BEHAVIORAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════

BEHAVIORAL_SYSTEM = """You are an expert behavioral interview coach specializing in STAR method responses.
You have deep knowledge of the job description and the candidate's resume provided to you.

Your role:
1. Ask ONE focused behavioral question at a time, tailored to the specific role and company
2. After the candidate responds, give structured STAR feedback:
   - Score the response 1-10
   - Identify what was strong
   - Point out what was missing (Situation/Task/Action/Result gaps)
   - Give a concrete suggestion to strengthen the answer
3. Then ask the next question or offer a follow-up
4. Vary question themes: leadership, conflict, failure, collaboration, achievement, adaptability

Always be encouraging but honest. Keep responses concise and actionable.
Format feedback clearly with sections: SCORE, STRENGTHS, GAPS, SUGGESTION.
Then end with: "Ready for the next question? (yes/no)" or ask a follow-up."""

def behavioral_chat(messages, resume_text, jd_text) -> str:
    system = f"""{BEHAVIORAL_SYSTEM}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:3000]}"""

    # Build message history
    api_messages = []
    for m in messages:
        api_messages.append({"role": m["role"], "content": m["content"]})

    response = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1000,
        system=system,
        messages=api_messages
    )
    return response.content[0].text

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3 — TECHNICAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3 — TECHNICAL INTERVIEW AGENT (GPT-4o powered)
# ══════════════════════════════════════════════════════════════════════════════

# ── Domain knowledge bank ──────────────────────────────────────────────────────
DOMAIN_PROFILES = {
    "dax_powerbi": {
        "label":    "DAX & Power BI",
        "icon":     "📊",
        "signals":  ["dax","power bi","powerbi","microsoft fabric","tabular","ssas","measures","calculated column","report","pbix"],
        "persona":  "You are a Microsoft-certified Power BI architect with 12 years of enterprise BI experience. You write complex DAX daily and have deep knowledge of the Vertipaq engine, query folding, and data model optimization.",
        "formats": {
            "conceptual":    "Explain how {topic} works in DAX/Power BI and when you would use it.",
            "write_code":    "Write a DAX {artifact} that {requirement}. Show the full expression and explain each function used.",
            "debug":         "Here is a DAX {artifact} that is returning incorrect results or an error. Identify the problem and rewrite it correctly:\n```dax\n{code_snippet}\n```",
            "scenario":      "Given a data model with {model_description}, how would you {task}? Walk through your approach and write the DAX.",
            "optimization":  "This DAX {artifact} is running slowly on a large dataset:\n```dax\n{code_snippet}\n```\nExplain why it's slow and rewrite it for better performance.",
        },
        "topic_pool": [
            "CALCULATE and filter context","context transition","SUMX vs SUM","iterator functions (SUMX, AVERAGEX, MAXX)",
            "time intelligence (TOTALYTD, SAMEPERIODLASTYEAR, DATEADD)","row context vs filter context",
            "RELATED and RELATEDTABLE","bidirectional relationships and their risks","RANKX","TOPN",
            "variables (VAR/RETURN)","SWITCH and SWITCH(TRUE())","DIVIDE vs division operator",
            "ALL, ALLEXCEPT, ALLSELECTED","SELECTEDVALUE and HASONEVALUE","TREATAS","USERELATIONSHIP",
            "incremental refresh","aggregations and composite models","DirectQuery vs Import mode performance",
            "query folding","Vertipaq engine column compression","field parameters","calculation groups",
            "dynamic format strings","many-to-many relationships","role-playing dimensions",
        ],
        "code_examples": [
            {
                "artifact": "measure",
                "code": "Sales YTD = TOTALYTD([Total Sales], 'Date'[Date])",
                "bug": "Returns blank when a slicer filters to a non-standard fiscal year end"
            },
            {
                "artifact": "measure",
                "code": "Running Total = CALCULATE([Sales], FILTER(ALL('Date'), 'Date'[Date] <= MAX('Date'[Date])))",
                "bug": "Extremely slow on tables with 10M+ rows"
            },
            {
                "artifact": "calculated column",
                "code": "Prev Month Sales = CALCULATE([Total Sales], PREVIOUSMONTH('Date'[Date]))",
                "bug": "Calculated columns cannot use time intelligence — this will error"
            },
        ],
    },

    "sql": {
        "label":    "SQL & Database",
        "icon":     "🗄️",
        "signals":  ["sql","t-sql","postgresql","mysql","snowflake","bigquery","redshift","oracle","stored procedure","query","database","etl","dbt","data warehouse"],
        "persona":  "You are a senior database engineer and SQL expert with 15 years across SQL Server, PostgreSQL, Snowflake, and BigQuery. You specialize in query optimization, execution plans, and large-scale data warehouse design.",
        "formats": {
            "conceptual":    "Explain {topic} in SQL — how it works internally and when you would use it.",
            "write_code":    "Write a SQL query that {requirement}. Use {dialect} syntax. Include comments explaining your approach.",
            "debug":         "This SQL query is returning wrong results or throwing an error. Identify the issue and fix it:\n```sql\n{code_snippet}\n```",
            "scenario":      "Given a schema with {model_description}, write the SQL to {task}. Explain your join strategy and any indexes you'd create.",
            "optimization":  "This query takes 45 seconds on a 50M row table. Explain exactly why it's slow and rewrite it optimally:\n```sql\n{code_snippet}\n```",
        },
        "topic_pool": [
            "window functions (ROW_NUMBER, RANK, DENSE_RANK, LAG, LEAD, NTILE)",
            "CTEs vs subqueries vs temp tables — when to use each",
            "INNER vs LEFT vs CROSS JOIN behavior","GROUP BY vs PARTITION BY",
            "index types (clustered, non-clustered, covering, columnstore)","execution plan reading",
            "query optimizer and statistics","HAVING vs WHERE","EXISTS vs IN vs JOIN for filtering",
            "set operations (UNION, INTERSECT, EXCEPT)","recursive CTEs","pivot and unpivot",
            "transaction isolation levels (READ COMMITTED, SERIALIZABLE, SNAPSHOT)",
            "deadlocks and how to prevent them","normalization (1NF, 2NF, 3NF, BCNF)",
            "star schema vs snowflake schema","slowly changing dimensions (SCD Type 1, 2, 3)",
            "query folding in Snowflake/BigQuery","materialized views","partitioning strategies",
            "data skew in distributed systems","EXPLAIN ANALYZE output interpretation",
            "NULL handling (COALESCE, NULLIF, IS NULL)","string functions and regex in SQL",
            "date arithmetic across dialects","JSON column querying","stored procedures vs functions",
        ],
        "code_examples": [
            {
                "artifact": "query",
                "code": "SELECT department, AVG(salary) FROM employees WHERE AVG(salary) > 70000 GROUP BY department;",
                "bug": "Cannot use aggregate function in WHERE — needs HAVING"
            },
            {
                "artifact": "query",
                "code": "SELECT * FROM orders o LEFT JOIN customers c ON o.customer_id = c.id WHERE c.country = 'US';",
                "bug": "WHERE on outer table turns LEFT JOIN into INNER JOIN — filter should be in ON clause"
            },
            {
                "artifact": "query",
                "code": "SELECT id, name, ROW_NUMBER() OVER (ORDER BY created_at) FROM users WHERE ROW_NUMBER() OVER (ORDER BY created_at) <= 10;",
                "bug": "Window functions cannot be used in WHERE — need CTE or subquery"
            },
        ],
    },

    "python_data": {
        "label":    "Python & Data Engineering",
        "icon":     "🐍",
        "signals":  ["python","pandas","numpy","scikit","spark","airflow","dbt","kafka","pyspark","ml","machine learning","etl pipeline","data engineering","fastapi","flask"],
        "persona":  "You are a senior data engineer and Python developer with expertise in pandas, PySpark, Airflow, and building production ML pipelines. You care deeply about memory efficiency, vectorization, and code that scales.",
        "formats": {
            "conceptual":    "Explain {topic} in Python/data engineering — how it works and when you'd use it.",
            "write_code":    "Write Python code using {library} to {requirement}. Include type hints and handle edge cases.",
            "debug":         "This Python code has a bug or performance problem. Find it and fix it:\n```python\n{code_snippet}\n```",
            "scenario":      "Given a dataset with {model_description}, write the Python/pandas code to {task}. Explain your approach.",
            "optimization":  "This pandas code is too slow on 5M+ rows:\n```python\n{code_snippet}\n```\nExplain why and rewrite it efficiently (vectorized, avoid loops).",
        },
        "topic_pool": [
            "pandas groupby and agg patterns","apply vs vectorized operations — when each is appropriate",
            "memory optimization (dtypes, chunking, categorical)","merge types and merge_asof",
            "pivot_table vs crosstab vs groupby","MultiIndex operations",
            "window functions in pandas (rolling, expanding, ewm)","string operations at scale",
            "datetime indexing and resampling","reading large CSVs efficiently",
            "list comprehensions vs generator expressions","decorators and context managers",
            "Python GIL and when it matters","multiprocessing vs threading vs asyncio",
            "PySpark RDD vs DataFrame API","Spark partitioning and shuffling",
            "Airflow DAG design and task dependencies","idempotent pipeline design",
            "type hints and dataclasses","pytest fixtures and mocking",
            "scikit-learn pipeline and ColumnTransformer","cross-validation strategy",
            "feature engineering best practices","handling class imbalance",
            "model serialization (pickle, joblib, ONNX)","REST API design with FastAPI",
        ],
        "code_examples": [
            {
                "artifact": "function",
                "code": "def process(df):\n    result = []\n    for i, row in df.iterrows():\n        result.append(row['sales'] * row['qty'])\n    return result",
                "bug": "Using iterrows() on large DataFrame — extremely slow, should be vectorized"
            },
            {
                "artifact": "merge",
                "code": "df1.merge(df2, on='id', how='left').merge(df3, on='id', how='left').fillna(0)",
                "bug": "Multiple merges can explode row count if keys aren't unique — missing dedup check"
            },
        ],
    },

    "analytics_general": {
        "label":    "Analytics & BI General",
        "icon":     "📈",
        "signals":  ["analytics","tableau","looker","qlik","reporting","kpi","dashboard","data analyst","business intelligence","metrics","a/b test","statistical"],
        "persona":  "You are a senior analytics engineer who has built BI solutions across Tableau, Looker, and Power BI. You are rigorous about metric definitions, statistical validity, and translating business questions into data.",
        "formats": {
            "conceptual":    "Explain {topic} — how you would approach this analytically.",
            "write_code":    "Write the SQL or calculation to {requirement} in {dialect}.",
            "debug":         "This dashboard metric is giving stakeholders incorrect numbers. Walk through how you'd debug it:\n{code_snippet}",
            "scenario":      "A business stakeholder asks you to {task}. How do you approach this? What data do you need, what are the pitfalls?",
            "optimization":  "This report takes 3 minutes to load for users. Walk through your investigation and optimization approach:\n{code_snippet}",
        },
        "topic_pool": [
            "metric definition and the single source of truth","leading vs lagging indicators",
            "cohort analysis","funnel analysis and drop-off","retention curves",
            "A/B test design — sample size, power, significance","p-values and statistical significance",
            "seasonality and trend decomposition","attribution modeling",
            "data quality checks and anomaly detection","slowly changing dimensions",
            "grain of a fact table","additive vs semi-additive vs non-additive measures",
            "fan traps and chasm traps in data models","row-level security",
        ],
        "code_examples": [],
    },
}

# ── Domain detector ────────────────────────────────────────────────────────────
def detect_domain(jd_text: str) -> str:
    """Score each domain by keyword frequency in the JD and return best match."""
    jd_lower = jd_text.lower()
    scores = {}
    for domain_key, profile in DOMAIN_PROFILES.items():
        score = sum(1 for signal in profile["signals"] if signal in jd_lower)
        # Weight multi-word signals more heavily
        score += sum(1 for signal in profile["signals"] if " " in signal and signal in jd_lower)
        scores[domain_key] = score
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "analytics_general"


# ── GPT-4o client ──────────────────────────────────────────────────────────────
@st.cache_resource
def get_openai_client(api_key: str):
    """Return an OpenAI client. Cached per api_key."""
    try:
        from openai import OpenAI
        return OpenAI(api_key=api_key)
    except ImportError:
        return None


def gpt4o_technical_chat(
    messages: list,
    resume_text: str,
    jd_text: str,
    domain_key: str,
    difficulty: str,
    openai_key: str,
) -> str:
    """GPT-4o powered technical interview — domain-specialized question generation."""
    import random

    profile = DOMAIN_PROFILES.get(domain_key, DOMAIN_PROFILES["analytics_general"])
    topics  = profile["topic_pool"]
    examples = profile.get("code_examples", [])

    # Build rich system prompt
    system = f"""You are {profile['persona']}

You are conducting a rigorous technical interview for the role described below.
Difficulty level: {difficulty}

QUESTION FORMAT ROTATION — cycle through these types, don't repeat the same format twice in a row:
1. CONCEPTUAL — "Explain how X works and when you'd use it"
2. WRITE THE CODE — Give a specific requirement, ask them to write working {profile['label']} code/query
3. DEBUG THIS — Show a broken code snippet, ask them to find and fix the bug
4. SCENARIO-BASED — Given a realistic data model or business situation, ask how they'd solve it
5. PERFORMANCE OPTIMIZATION — Show slow code, ask why it's slow and how to fix it

DOMAIN: {profile['label']}
TOPIC POOL (rotate through these, don't repeat):
{chr(10).join(f"- {t}" for t in topics)}

SAMPLE CODE BUGS (use these for DEBUG questions):
{chr(10).join(f"- {e['artifact']}: {e['bug']}" for e in examples) if examples else "Generate realistic bugs from your expertise"}

RULES:
- Ask EXACTLY ONE question per turn — never multiple questions
- For code questions: show actual {profile['label']} code in a properly fenced code block
- For debug questions: show broken code first, THEN ask what's wrong — do not reveal the answer
- For scenario questions: describe a realistic data model (table names, columns, relationships)
- After the candidate answers: give structured feedback:
  SCORE: X/10
  CORRECT: [what they got right]  
  MISSED: [key points not covered]
  IDEAL ANSWER: [concise version of perfect answer, with correct code if applicable]
  Then ask your next question (different format from last one)
- Be technically precise — wrong syntax in your code examples is not acceptable
- Calibrate difficulty to: {difficulty}
- Progress from foundational → advanced as the conversation develops

CANDIDATE RESUME (use this to calibrate starting difficulty and reference their experience):
{resume_text[:2000]}

JOB DESCRIPTION:
{jd_text[:2500]}"""

    client = get_openai_client(openai_key)
    if not client:
        return "❌ OpenAI package not installed. Run: pip install openai"

    try:
        api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": system}] + api_messages,
            max_tokens=1400,
            temperature=0.7,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ GPT-4o error: {e}"


# ── Fallback: Claude technical chat (used when no OpenAI key) ──────────────────
TECHNICAL_SYSTEM = """You are a senior technical interviewer conducting a rigorous technical interview.
You have the job description and candidate resume. Questions should reflect the technical stack directly.

Ask ONE question per turn. After the candidate answers give:
  SCORE: X/10 | CORRECT: ... | MISSED: ... | IDEAL ANSWER: ...
Then ask the next question. Rotate formats: conceptual, write-the-code, debug, scenario, optimization."""

def technical_chat(messages, resume_text, jd_text) -> str:
    """Claude fallback when no OpenAI key is provided."""
    system = f"""{TECHNICAL_SYSTEM}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:3000]}"""
    api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]
    response = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1200,
        system=system, messages=api_messages
    )
    return response.content[0].text

# ── Hiring Manager chat ────────────────────────────────────────────────────────
HIRING_SYSTEM = """You are a hiring manager conducting a final-round interview focused on culture fit,
leadership potential, strategic thinking, and long-term vision.

You have the job description and the candidate's resume. Your questions should probe:
1. Why this company and role specifically — genuine motivation
2. Career trajectory and where they want to be in 3-5 years
3. Leadership style, team dynamics, conflict resolution
4. How they've handled ambiguity, rapid change, or high-stakes decisions
5. What value they bring beyond the job spec

Your tone: warm but direct. You care about the human behind the resume.
After their answer: give brief honest feedback — what resonated, what felt rehearsed, and one follow-up probe.
Keep responses conversational, not formulaic. No rubric scores — just real dialogue."""

def hiring_chat(messages, resume_text, jd_text) -> str:
    system = f"""{HIRING_SYSTEM}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:3000]}"""
    api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]
    response = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1000,
        system=system, messages=api_messages
    )
    return response.content[0].text

# ── Avatar helper ──────────────────────────────────────────────────────────────
def render_avatar(persona: str, message: str, height: int = None):
    """Render animated interviewer avatar with TTS into the Streamlit page.
    Height is computed from message length if not provided — prevents clipping.
    """
    if height is None:
        # Base height (avatar + controls) + proportional bubble height
        char_count = len(message)
        # ~55 chars per line at 14px in a ~300px bubble, ~22px per line
        estimated_lines = max(4, char_count // 52)
        bubble_height   = estimated_lines * 22 + 40   # padding
        controls_height = 160                          # TTS bar + avatar label
        height = 130 + bubble_height + controls_height  # avatar SVG + margins
        height = max(320, min(height, 900))            # clamp 320–900px
    components.html(get_avatar_html(persona, message), height=height, scrolling=False)

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 4 — TECHNICAL EXAM
# ══════════════════════════════════════════════════════════════════════════════

def generate_exam(jd_text, resume_text) -> dict:
    prompt = f"""You are a technical hiring manager. Create a timed technical exam based on this job description.
Return ONLY raw JSON — no markdown fences.

JOB DESCRIPTION:
{jd_text}

CANDIDATE BACKGROUND (use this to calibrate difficulty):
{resume_text[:2000]}

Generate an exam with this exact JSON structure:
{{
  "title": "<Role> Technical Assessment",
  "duration_minutes": 30,
  "instructions": "<2 sentence exam instructions>",
  "sections": [
    {{
      "name": "Section 1: <topic>",
      "questions": [
        {{
          "id": "q1",
          "type": "multiple_choice",
          "question": "<question text>",
          "options": {{"A": "<option>", "B": "<option>", "C": "<option>", "D": "<option>"}},
          "correct": "A",
          "explanation": "<why this is correct and others are wrong>",
          "difficulty": "easy|medium|hard",
          "topic": "<topic tag>"
        }},
        {{
          "id": "q2",
          "type": "multiple_choice",
          "question": "<question>",
          "options": {{"A": "<option>", "B": "<option>", "C": "<option>", "D": "<option>"}},
          "correct": "B",
          "explanation": "<explanation>",
          "difficulty": "medium",
          "topic": "<topic>"
        }}
      ]
    }},
    {{
      "name": "Section 2: <topic>",
      "questions": [
        {{
          "id": "q6",
          "type": "short_answer",
          "question": "<question requiring a written answer>",
          "sample_answer": "<what a good answer looks like>",
          "key_points": ["<point 1>", "<point 2>", "<point 3>"],
          "difficulty": "medium",
          "topic": "<topic>"
        }}
      ]
    }}
  ]
}}

Requirements:
- 3 sections total covering the main technical areas from the JD
- Section 1: 5 multiple choice questions (fundamentals)
- Section 2: 5 multiple choice questions (intermediate/applied)
- Section 3: 3 short answer questions (design/scenario-based)
- Questions must be directly relevant to THIS specific job description
- Difficulty should be appropriate for the role level implied by the JD
- Make questions substantive and realistic — not trivial"""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)

def grade_short_answer(question: str, sample_answer: str, key_points: list, user_answer: str) -> dict:
    prompt = f"""Grade this short answer response for a technical interview exam.

QUESTION: {question}

CANDIDATE ANSWER: {user_answer}

SAMPLE ANSWER: {sample_answer}

KEY POINTS EXPECTED: {json.dumps(key_points)}

Return ONLY raw JSON:
{{
  "score": <integer 0-10>,
  "points_covered": ["<point covered>"],
  "points_missed": ["<point missed>"],
  "feedback": "<2-3 sentence constructive feedback>",
  "grade": "Excellent|Good|Partial|Needs Work"
}}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=600,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)


# ══════════════════════════════════════════════════════════════════════════════
# HUMANIZATION ENGINE
# ══════════════════════════════════════════════════════════════════════════════

HUMANIZE_SYSTEM = """You are a professional editor who specializes in making AI-generated writing 
sound genuinely human. You have a sharp eye for AI tells and know how to eliminate them without 
losing substance.

AI writing patterns you ALWAYS eliminate:
- Filler openers: "In today's...", "In the ever-evolving...", "As a...", "I am writing to..."
- Hollow superlatives: "passionate", "dynamic", "innovative", "leverage", "synergy", "utilize"  
- Robotic transitions: "Furthermore", "Moreover", "Additionally", "In conclusion", "To summarize"
- Overuse of "I" at the start of every sentence
- Perfect parallel structure in every list (humans vary this)
- Vague claims with no specifics ("significant results", "great success", "various projects")
- Corporate buzzwords: "stakeholders", "deliverables", "bandwidth", "circle back", "deep dive"
- Excessive hedging: "I believe", "I feel", "I think" used repeatedly
- The phrase "I would be remiss" or any similar affected formality
- Ending with "I look forward to hearing from you" as the sole closing thought

What genuine human writing does instead:
- Starts mid-thought, sometimes with a short punchy sentence
- Mixes long and short sentences — a 4-word sentence after a 25-word sentence
- Uses specific details and numbers rather than vague superlatives
- Has a slight personality that comes through — dry, warm, direct, or wry
- Mentions one concrete story or moment, briefly
- Has one slightly unconventional word choice that feels intentional, not random
- The closer feels earned, not formulaic

Preserve all specific facts, numbers, dates, and claims from the original.
Do NOT add new facts that weren't in the original.
Match the requested tone but let personality breathe through."""


# ══════════════════════════════════════════════════════════════════════════════
# TECHNICAL PROJECTS ENGINE
# ══════════════════════════════════════════════════════════════════════════════

# ── BI Dashboard project brief generator ──────────────────────────────────────
BI_STAGES = [
    {
        "id": "data_model",
        "title": "Stage 1 — Data Model Design",
        "icon": "🗂️",
        "prompt": "Design the star schema data model for this project. Define your fact table(s), dimension tables, key columns, relationships, and cardinality. Explain your grain decision.",
        "placeholder": "Describe your fact table(s), dimension tables, columns, relationships, and why you chose this grain...",
        "review_focus": "data model design, star schema correctness, grain definition, relationship cardinality, normalization decisions",
    },
    {
        "id": "dax_measures",
        "title": "Stage 2 — Core DAX Measures",
        "icon": "📐",
        "prompt": "Write the 4 most critical DAX measures for this dashboard. For each measure: show the full DAX expression, explain what it calculates, and note any filter context considerations.",
        "placeholder": "Measure 1:\n[Name] = \n\nExplanation:\n\nMeasure 2:\n...",
        "review_focus": "DAX syntax correctness, filter context handling, use of CALCULATE, time intelligence, performance implications, naming conventions",
    },
    {
        "id": "dashboard_design",
        "title": "Stage 3 — Dashboard Layout & UX",
        "icon": "🎨",
        "prompt": "Describe your dashboard layout. What visuals would you use for each KPI and why? How would you handle cross-filtering, drill-through, and mobile responsiveness? What would go on each page?",
        "placeholder": "Page 1 — Executive Summary:\n  Visual 1: ...\n  Visual 2: ...\n\nPage 2 — ...",
        "review_focus": "visual selection appropriateness, UX hierarchy, cross-filtering strategy, drill-through design, accessibility, report navigation",
    },
    {
        "id": "performance",
        "title": "Stage 4 — Performance & Optimization",
        "icon": "⚡",
        "prompt": "This dashboard will serve 500 concurrent users on a 50M row fact table. Walk through your optimization strategy: Import vs DirectQuery decision, aggregations, calculated columns vs measures, row-level security approach, and refresh strategy.",
        "placeholder": "Import vs DirectQuery decision and reasoning:\n\nAggregation strategy:\n\nRLS approach:\n\nRefresh design:\n...",
        "review_focus": "Import/DirectQuery tradeoffs, aggregation table design, Vertipaq optimization, RLS scalability, incremental refresh configuration",
    },
]

SQL_STAGES = [
    {
        "id": "exploration",
        "title": "Stage 1 — Data Exploration",
        "icon": "🔍",
        "prompt": "Write the SQL queries you'd run first to understand this dataset. Include: row counts, null checks, cardinality of key columns, date range, and any data quality issues you'd flag.",
        "placeholder": "-- Row counts and basic profiling\nSELECT ...\n\n-- Null checks\nSELECT ...\n\n-- Cardinality\nSELECT ...",
        "review_focus": "profiling completeness, null handling, cardinality checks, data quality awareness, query efficiency",
    },
    {
        "id": "core_query",
        "title": "Stage 2 — Core Analysis Query",
        "icon": "📊",
        "prompt": "Write the primary analytical query that answers the central business question. Use CTEs for readability. Include window functions where appropriate. Comment your logic.",
        "placeholder": "WITH base AS (\n  SELECT ...\n),\nfinal AS (\n  SELECT ...\n)\nSELECT * FROM final;",
        "review_focus": "query correctness, CTE structure, window function usage, join strategy, business logic accuracy, readability and commenting",
    },
    {
        "id": "optimization",
        "title": "Stage 3 — Query Optimization",
        "icon": "🚀",
        "prompt": "Your Stage 2 query runs in 4 minutes on production. Rewrite it for performance. Explain every change you made and why it helps. Include any clustering keys, partitioning, or materialization decisions.",
        "placeholder": "-- Optimized version\n/*\nChanges made:\n1. ...\n2. ...\n*/\n\nSELECT ...",
        "review_focus": "clustering/partitioning usage, predicate pushdown, join order, result caching, materialized view candidates, query profile interpretation",
    },
    {
        "id": "pipeline",
        "title": "Stage 4 — Production Pipeline Design",
        "icon": "🔄",
        "prompt": "Design the production data pipeline for this analysis. Write the DDL for your target table, a stored procedure or dbt model structure, incremental load logic, and how you'd handle late-arriving data and failures.",
        "placeholder": "-- Target table DDL\nCREATE OR REPLACE TABLE ...\n\n-- Incremental merge logic\nMERGE INTO ...\n\n-- Failure handling approach:\n...",
        "review_focus": "DDL design, incremental strategy (merge vs insert-overwrite), idempotency, error handling, SCD handling, dbt model structure",
    },
]


def generate_bi_brief(jd_text: str, resume_text: str) -> dict:
    """Generate a realistic BI dashboard project brief tailored to the JD."""
    prompt = f"""You are a hiring manager creating a realistic take-home BI dashboard project.
Generate a project brief based on this job description. Return ONLY raw JSON.

JOB DESCRIPTION:
{jd_text[:2500]}

CANDIDATE BACKGROUND:
{resume_text[:1500]}

Return this exact JSON:
{{
  "title": "<Project title — specific, not generic>",
  "company_context": "<2-3 sentence fictional company description relevant to the JD domain>",
  "business_problem": "<3-4 sentence description of the specific business problem to solve>",
  "data_model_description": "<Description of 3-4 tables the candidate will work with: table names, key columns, row counts, relationships. Make these realistic for the domain.>",
  "kpis_required": ["<KPI 1>", "<KPI 2>", "<KPI 3>", "<KPI 4>", "<KPI 5>"],
  "technical_requirements": ["<requirement 1>", "<requirement 2>", "<requirement 3>", "<requirement 4>"],
  "sample_data": {{
    "fact_table": "<table name>",
    "columns": ["<col1>", "<col2>", "<col3>", "<col4>", "<col5>"],
    "dim_tables": ["<dim1>", "<dim2>", "<dim3>"]
  }},
  "evaluation_criteria": ["<criterion 1>", "<criterion 2>", "<criterion 3>"],
  "estimated_hours": "<e.g. 3-4 hours>",
  "difficulty": "Intermediate|Senior|Staff"
}}

Make the project domain-specific to the JD. If the JD mentions healthcare, make it a healthcare dashboard. Financial services → financial reporting. Retail → sales analytics. Be specific with table and column names."""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)


def generate_sql_brief(jd_text: str, resume_text: str, dialect: str) -> dict:
    """Generate a realistic SQL analytics project brief."""
    prompt = f"""You are a senior data engineer creating a realistic take-home SQL project.
Generate a project brief for {dialect} SQL based on this job description. Return ONLY raw JSON.

JOB DESCRIPTION:
{jd_text[:2500]}

CANDIDATE BACKGROUND:
{resume_text[:1500]}

Return this exact JSON:
{{
  "title": "<Project title>",
  "dialect": "{dialect}",
  "company_context": "<2-3 sentence fictional company description>",
  "business_problem": "<3-4 sentence description of the analytics problem>",
  "schema": {{
    "tables": [
      {{
        "name": "<table_name>",
        "description": "<what this table contains>",
        "columns": [
          {{"name": "<col>", "type": "<type>", "description": "<what it means>"}},
          {{"name": "<col>", "type": "<type>", "description": "<what it means>"}},
          {{"name": "<col>", "type": "<type>", "description": "<what it means>"}}
        ],
        "row_count": "<e.g. 45M rows>",
        "clustering_key": "<column name or null>"
      }}
    ]
  }},
  "central_question": "<The main business question to answer in 1 sentence>",
  "sub_questions": ["<sub-question 1>", "<sub-question 2>", "<sub-question 3>"],
  "known_data_issues": ["<issue 1 — e.g. 3% of order_ids have duplicate records>", "<issue 2>"],
  "performance_constraints": "<e.g. Must run under 30 seconds on 50M row fact table>",
  "evaluation_criteria": ["<criterion 1>", "<criterion 2>", "<criterion 3>"],
  "estimated_hours": "<e.g. 2-3 hours>",
  "difficulty": "Intermediate|Senior|Staff"
}}

Make the schema realistic and domain-specific. Use real {dialect} data types. Include 3-4 tables. Column names should be snake_case."""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1800,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)


def review_project_submission(
    stage: dict,
    submission: str,
    brief: dict,
    project_type: str,
    dialect: str = "Snowflake",
    openai_key: str = "",
) -> dict:
    """GPT-4o reviews a project stage submission. Falls back to Claude."""
    brief_summary = json.dumps(brief, indent=2)[:2000]

    system = """You are a senior technical reviewer evaluating a data professional's take-home project submission.
Be rigorous, specific, and constructive. Point to exact lines or expressions in their submission.
Never give vague feedback like "good job" — always explain WHY something is correct or incorrect."""

    prompt = f"""Review this {project_type} project submission for Stage: {stage['title']}

PROJECT BRIEF CONTEXT:
{brief_summary}

EVALUATION FOCUS: {stage['review_focus']}

CANDIDATE SUBMISSION:
{submission}

Return ONLY raw JSON:
{{
  "score": <integer 1-10>,
  "grade": "Excellent|Strong|Adequate|Needs Work|Insufficient",
  "summary": "<2-3 sentence overall assessment>",
  "strengths": [
    {{"point": "<specific thing done well>", "detail": "<why it matters technically>"}},
    {{"point": "<strength 2>", "detail": "<technical detail>"}}
  ],
  "issues": [
    {{"severity": "Critical|Major|Minor", "point": "<specific issue>", "fix": "<exactly how to fix it>"}},
    {{"severity": "Major|Minor", "point": "<issue 2>", "fix": "<fix>"}}
  ],
  "missed_opportunities": ["<thing they didn't consider but should have>", "<missed opportunity 2>"],
  "corrected_example": "<If there's a code error, show the correct version here. Empty string if no code correction needed.>",
  "next_stage_tip": "<One specific thing to focus on in the next stage based on what you saw here>"
}}"""

    # Try GPT-4o first if key available
    if openai_key:
        try:
            client = get_openai_client(openai_key)
            if client:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": system},
                        {"role": "user", "content": prompt},
                    ],
                    max_tokens=1400,
                    temperature=0.3,
                )
                return clean_json(response.choices[0].message.content)
        except Exception:
            pass

    # Claude fallback
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1400,
        system=system,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)


def overall_project_score(reviews: dict, stages: list) -> dict:
    """Compute overall project score from stage reviews."""
    scores = [reviews[s["id"]]["score"] for s in stages if s["id"] in reviews]
    if not scores:
        return {"overall": 0, "grade": "—", "color": "#6b7280"}
    avg = sum(scores) / len(scores)
    grade = "Distinction" if avg >= 9 else "Merit" if avg >= 7.5 else "Pass" if avg >= 6 else "Borderline" if avg >= 5 else "Fail"
    color = "#00d4aa" if avg >= 7.5 else "#ffa940" if avg >= 5 else "#ff4d6d"
    return {"overall": round(avg, 1), "grade": grade, "color": color, "scores": scores}


# ── Session save / restore ─────────────────────────────────────────────────────
def build_session_snapshot() -> dict:
    """Serialise all scorable session state into a portable JSON snapshot."""
    ss = st.session_state
    return {
        "version": "1.0",
        "role":    (ss.analysis or {}).get("role_title", ""),
        "company": (ss.analysis or {}).get("company", ""),
        "match_score": (ss.analysis or {}).get("match_score"),
        "cover_letter": ss.cover_letter,
        "exam_submitted": ss.exam_submitted,
        "exam_score": _exam_score(),
        "bi_reviews":  ss.proj_bi_reviews,
        "sql_reviews": ss.proj_sql_reviews,
        "bi_brief_title":  (ss.proj_bi_brief or {}).get("title",""),
        "sql_brief_title": (ss.proj_sql_brief or {}).get("title",""),
        "behavioral_turns": len([m for m in ss.behavioral_messages if m["role"]=="assistant"]),
        "technical_turns":  len([m for m in ss.technical_messages  if m["role"]=="assistant"]),
    }


def _exam_score() -> dict | None:
    """Compute exam score summary from session state."""
    ss = st.session_state
    if not ss.exam_submitted or not ss.exam_data:
        return None
    mc_c = mc_t = sa_s = sa_t = 0
    for section in ss.exam_data.get("sections", []):
        for q in section.get("questions", []):
            if q.get("type") == "multiple_choice":
                mc_t += 1
                if ss.exam_answers.get(q["id"]) == q.get("correct"):
                    mc_c += 1
            elif q.get("type") == "short_answer":
                g = ss.exam_sa_grades.get(q["id"], {})
                sa_s += g.get("score", 0)
                sa_t += 10
    total = mc_t * 10 + sa_t
    earned = mc_c * 10 + sa_s
    pct = int(earned / total * 100) if total else 0
    return {"pct": pct, "earned": earned, "total": total}


def build_project_report_docx(project_type: str, brief: dict, stages: list, reviews: dict, submissions: dict) -> bytes:
    """Build a detailed project review report as .docx."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime, io

        doc = DocxDoc()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(0.9)
            sec.left_margin = sec.right_margin = Inches(1.0)

        def _heading(text, size=14, bold=True, color=(0x1B, 0x2A, 0x4A)):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold; r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(*color)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            return p

        def _body(text, size=10.5, color=(0x2D, 0x3A, 0x4F)):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor(*color)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(15)
            return p

        def _rule():
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
            bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bot); pPr.append(pBdr)
            p.paragraph_format.space_after = Pt(6)

        # ── Title block ────────────────────────────────────────────────────────
        _heading(f"Technical Project Review Report", size=18, color=(0x1B, 0x2A, 0x4A))
        _heading(brief.get("title", project_type), size=13, bold=False, color=(0x71, 0x80, 0x96))
        _body(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}  ·  Project type: {project_type}  ·  Difficulty: {brief.get('difficulty','')}")
        _rule()

        # ── Overall score ──────────────────────────────────────────────────────
        overall = overall_project_score(reviews, stages)
        _heading(f"Overall Score: {overall['overall']}/10 — {overall['grade']}", size=13, color=(0x0D, 0x7A, 0x6E))
        stage_scores = "  ·  ".join(
            f"{s['title'].split('—')[1].strip()}: {reviews.get(s['id'],{}).get('score','—')}/10"
            for s in stages if s["id"] in reviews
        )
        _body(stage_scores)
        _rule()

        # ── Project brief summary ──────────────────────────────────────────────
        _heading("Project Brief", size=12)
        _body(brief.get("business_problem", brief.get("central_question", "")))

        # ── Stage-by-stage review ──────────────────────────────────────────────
        for stage in stages:
            sid = stage["id"]
            if sid not in reviews:
                continue
            rv = reviews[sid]
            score = rv.get("score", 0)
            grade = rv.get("grade", "")

            doc.add_paragraph()
            _heading(f"{stage['title']}  —  {score}/10  ({grade})", size=12, color=(0x1B, 0x2A, 0x4A))
            _body(rv.get("summary", ""))

            # Submission
            if sid in submissions:
                _heading("Your Submission", size=10, bold=True, color=(0x4A, 0x55, 0x68))
                sub_para = doc.add_paragraph()
                sub_para.paragraph_format.left_indent = Inches(0.3)
                r = sub_para.add_run(submissions[sid][:1200] + ("…" if len(submissions[sid]) > 1200 else ""))
                r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

            # Strengths
            if rv.get("strengths"):
                _heading("Strengths", size=10, bold=True, color=(0x0D, 0x7A, 0x6E))
                for s in rv["strengths"]:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(f"{s.get('point','')} — "); r.bold = True; r.font.size = Pt(10)
                    r2 = p.add_run(s.get("detail","")); r2.font.size = Pt(10)

            # Issues
            if rv.get("issues"):
                _heading("Issues & Fixes", size=10, bold=True, color=(0xB0, 0x30, 0x30))
                for issue in rv["issues"]:
                    p = doc.add_paragraph(style="List Bullet")
                    sev = issue.get("severity","Minor")
                    r = p.add_run(f"[{sev}] {issue.get('point','')} → "); r.bold = True; r.font.size = Pt(10)
                    r2 = p.add_run(issue.get("fix","")); r2.font.size = Pt(10)

            # Corrected example
            if rv.get("corrected_example","").strip():
                _heading("Corrected Example", size=10, bold=True, color=(0x1B, 0x2A, 0x4A))
                cp = doc.add_paragraph()
                cp.paragraph_format.left_indent = Inches(0.3)
                r = cp.add_run(rv["corrected_example"])
                r.font.size = Pt(9); r.font.name = "Courier New"
                r.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4F)

            # Missed opportunities
            if rv.get("missed_opportunities"):
                _heading("Missed Opportunities", size=10, bold=True, color=(0x71, 0x80, 0x96))
                for m in rv["missed_opportunities"]:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(m); r.font.size = Pt(10)

            _rule()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception as e:
        return f"Error generating report: {e}".encode()


def humanize_text(text: str, tone: str = "Professional & Confident", doc_type: str = "cover letter") -> str:
    """Run generated text through the humanization layer."""
    prompt = f"""Rewrite this {doc_type} to sound like it was written by a real person, not an AI.
Tone requested: {tone}

Apply the full humanization process:
1. Kill every AI tell (hollow words, robotic transitions, formulaic openers/closers)
2. Vary sentence length intentionally — mix short punchy sentences with longer ones
3. Make one concrete detail more specific and vivid
4. Let a touch of personality show in word choice — not forced, just present
5. The opening line must earn attention immediately, not with a pleasantry
6. The closing must feel like a real human wrote it after thinking for 10 seconds

Return ONLY the rewritten text. No commentary, no "Here is the rewritten version:".

ORIGINAL TEXT:
{text}"""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1500,
        system=HUMANIZE_SYSTEM,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text.strip()


def humanize_resume_bullets(resume_text: str) -> str:
    """Run the humanization pass specifically tuned for resume bullets."""
    prompt = f"""Rewrite these resume bullets to sound like a real person wrote them — 
not an AI generating polished corporate-speak.

Rules:
- Kill hollow words: "leveraged", "spearheaded", "utilized", "facilitated", "championed"
- Replace vague superlatives with specifics wherever possible  
- Vary the opening verb — not every bullet starts with "Led" or "Developed"
- Keep all numbers, dates, companies, and factual claims EXACTLY as they are
- Section headers (ALL CAPS lines) stay unchanged
- Make the language direct, confident, and grounded — like a competent professional talking, not performing
- Remove any filler phrases at the start of bullets ("Successfully", "Effectively", "Proactively")

Return ONLY the rewritten resume text. No commentary.

RESUME:
{resume_text}"""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text.strip()


# ══════════════════════════════════════════════════════════════════════════════
# COVER LETTER GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def generate_cover_letter(
    resume_text: str,
    jd_text: str,
    candidate_name: str,
    company: str,
    role: str,
    tone: str,
    analysis: dict = None,
) -> str:
    """Generate a raw cover letter draft, then humanize it."""

    strengths   = ", ".join((analysis or {}).get("strengths", [])[:4])
    gaps        = ", ".join((analysis or {}).get("gaps", [])[:2])
    gap_note    = f"Acknowledge briefly that you're growing in: {gaps}." if gaps else ""

    tone_guides = {
        "Professional & Confident":
            "Assured and direct. No fluff. Let accomplishments speak.",
        "Warm & Conversational":
            "Approachable and genuine. Like a smart colleague talking, not a press release.",
        "Bold & Direct":
            "Sharp, punchy. Get to the point in sentence one. No pleasantries.",
        "Storytelling":
            "Open with a 2-sentence specific moment or story that connects to the role, "
            "then pivot to qualifications.",
    }
    tone_guide = tone_guides.get(tone, tone_guides["Professional & Confident"])

    prompt = f"""Write a cover letter for {candidate_name} applying to the {role} position at {company}.

TONE: {tone} — {tone_guide}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:3000]}

KEY STRENGTHS TO HIGHLIGHT: {strengths}
{gap_note}

Structure (but don't make it feel structured):
- Opening: Immediately establish why THIS role at THIS company — not generic enthusiasm
- Body paragraph 1: One or two specific accomplishments most relevant to the JD's top requirements
- Body paragraph 2: Why this company specifically — reference something real about them if possible
- Closing: Confident, forward-looking, not begging

Length: 3-4 paragraphs, 250-320 words. Never go over 350.
Do NOT include date, address blocks, or "Dear Hiring Manager" — start with the first paragraph.
Do NOT end with "Sincerely," — end with just the name on its own line after the closing paragraph.

Write the draft now:"""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=900,
        messages=[{"role": "user", "content": prompt}]
    )
    draft = msg.content[0].text.strip()

    # ── Humanization pass ──────────────────────────────────────────────────────
    humanized = humanize_text(draft, tone=tone, doc_type="cover letter")
    return humanized


def build_cover_letter_docx(text: str, name: str, role: str, company: str) -> bytes:
    """Build a clean cover letter .docx."""
    try:
        from docx import Document as DocxDoc
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime, io

        doc = DocxDoc()
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.15)
            section.right_margin  = Inches(1.15)

        # Header — name
        hdr = doc.add_paragraph()
        hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = hdr.add_run(name.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

        # Sub-header — role · company · date
        date_str = datetime.date.today().strftime("%B %d, %Y")
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub_run = sub.add_run(f"{role}  ·  {company}  ·  {date_str}")
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        sub.paragraph_format.space_after = Pt(4)

        # Thin rule
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1B2A4A')
        pBdr.append(bottom)
        pPr.append(pBdr)
        p.paragraph_format.space_after = Pt(14)

        # Body paragraphs
        for para_text in text.split("\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(stripped)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4F)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(10)
            p.paragraph_format.line_spacing = Pt(16)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception:
        return text.encode("utf-8")


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ══════════════════════════════════════════════════════════════════════════════
defaults = {
    "analysis": None, "tailored_resume": None,
    "resume_text": None, "jd_text": None,
    "active_module": "resume",
    # Cover letter
    "cover_letter": None,
    "cover_letter_tone": "Professional & Confident",
    "cover_letter_name": "",
    "cover_letter_company": "",
    "cover_letter_role": "",
    # Behavioral
    "behavioral_messages": [],
    "behavioral_started": False,
    "behavioral_persona": "behavioral",
    # Technical
    "technical_messages": [],
    "technical_started": False,
    "tech_persona": "technical",
    "tech_domain": None,          # auto-detected domain key
    "tech_domain_override": None, # user-overridden domain key
    "openai_key": "",             # entered in sidebar
    # Hiring Manager
    "hiring_messages": [],
    "hiring_started": False,
    # Active interview tab
    "interview_tab": "behavioral",
    # Exam
    "exam_data": None,
    "exam_answers": {},
    "exam_sa_answers": {},
    "exam_submitted": False,
    "exam_sa_grades": {},
    "exam_started": False,
    "exam_start_time": None,
    # Technical Projects
    "proj_bi_brief": None,          # generated project brief
    "proj_bi_submissions": {},      # stage_id -> submitted code/text
    "proj_bi_reviews": {},          # stage_id -> GPT-4o review dict
    "proj_bi_active_stage": 0,
    "proj_bi_complete": False,
    "proj_sql_brief": None,
    "proj_sql_submissions": {},
    "proj_sql_reviews": {},
    "proj_sql_active_stage": 0,
    "proj_sql_complete": False,
    "proj_sql_dialect": "Snowflake",
    "proj_active": "bi",            # "bi" | "sql"
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════════════════════════════════════
# HERO
# ══════════════════════════════════════════════════════════════════════════════
# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:

    # ── API config ─────────────────────────────────────────────────────────────
    st.markdown("""
    <div style="font-family:Syne,sans-serif;font-size:0.95rem;font-weight:700;
         color:#e8eaf0;margin-bottom:0.8rem;">⚙️ API Configuration</div>
    """, unsafe_allow_html=True)

    try:
        _ant_key = st.secrets.get("ANTHROPIC_API_KEY","") or os.environ.get("ANTHROPIC_API_KEY","")
    except Exception:
        _ant_key = os.environ.get("ANTHROPIC_API_KEY","")
    ant_status = "✅ Connected" if _ant_key else "❌ Not set"
    st.markdown(f"""
    <div style="background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.1);
         border-radius:8px;padding:0.6rem 0.8rem;margin-bottom:0.8rem;font-size:0.82rem;">
      <div style="color:#6b7280;font-size:0.7rem;text-transform:uppercase;
           letter-spacing:0.1em;margin-bottom:2px;">Anthropic (Claude)</div>
      <div style="color:#e8eaf0;font-weight:600;">{ant_status}</div>
      <div style="color:#6b7280;font-size:0.72rem;margin-top:2px;">
        Resume · Cover Letter · Behavioral · Exam · Projects
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""<div style="color:#6b7280;font-size:0.72rem;text-transform:uppercase;
         letter-spacing:0.1em;margin-bottom:4px;">OpenAI API Key</div>""",
         unsafe_allow_html=True)
    openai_input = st.text_input(
        "OpenAI key", value=st.session_state.openai_key,
        type="password", placeholder="sk-...",
        label_visibility="collapsed", key="openai_key_input",
    )
    if openai_input != st.session_state.openai_key:
        st.session_state.openai_key = openai_input
        st.session_state.technical_messages = []
        st.session_state.technical_started  = False

    if st.session_state.openai_key:
        st.markdown("""<div style="background:rgba(0,212,170,0.08);border:1px solid rgba(0,212,170,0.25);
             border-radius:8px;padding:0.5rem 0.7rem;font-size:0.78rem;color:#00d4aa;margin-bottom:0.5rem;">
          ✅ GPT-4o active — Technical Interview &amp; Project Review
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown("""<div style="background:rgba(255,169,64,0.08);border:1px solid rgba(255,169,64,0.25);
             border-radius:8px;padding:0.5rem 0.7rem;font-size:0.78rem;color:#ffa940;margin-bottom:0.5rem;">
          ⚠️ No OpenAI key — Claude fallback active
        </div>""", unsafe_allow_html=True)

    st.markdown("""<div style="font-size:0.72rem;color:#6b7280;line-height:1.6;margin-bottom:0.5rem;">
      <strong style="color:#e8eaf0;">Model routing:</strong><br>
      🧠 Claude — Resume, Cover Letter, Behavioral, Exam<br>
      ⚡ GPT-4o — Technical Interview, Project Reviews<br>
    </div>""", unsafe_allow_html=True)

    st.markdown("---")

    # ── Progress tracker ────────────────────────────────────────────────────────
    st.markdown("""<div style="font-family:Syne,sans-serif;font-size:0.95rem;font-weight:700;
         color:#e8eaf0;margin-bottom:0.8rem;">📊 Session Progress</div>""",
         unsafe_allow_html=True)

    ss = st.session_state

    # Resume match score
    match = (ss.analysis or {}).get("match_score")
    _match_c = score_color(match) if match else "#6b7280"
    _match_v = f"{match}/100" if match is not None else "—"

    # Exam score
    _exam = _exam_score()
    _exam_v = f"{_exam['pct']}%" if _exam else "—"
    _exam_c = score_color(_exam["pct"]) if _exam else "#6b7280"

    # BI project score
    _bi_reviews = ss.proj_bi_reviews
    if _bi_reviews:
        _bi_avg = sum(v.get("score",0) for v in _bi_reviews.values()) / len(_bi_reviews)
        _bi_v = f"{round(_bi_avg,1)}/10 ({len(_bi_reviews)}/4 stages)"
        _bi_c = score_color(int(_bi_avg*10))
    else:
        _bi_v, _bi_c = "Not started", "#6b7280"

    # SQL project score
    _sql_reviews = ss.proj_sql_reviews
    if _sql_reviews:
        _sql_avg = sum(v.get("score",0) for v in _sql_reviews.values()) / len(_sql_reviews)
        _sql_v = f"{round(_sql_avg,1)}/10 ({len(_sql_reviews)}/4 stages)"
        _sql_c = score_color(int(_sql_avg*10))
    else:
        _sql_v, _sql_c = "Not started", "#6b7280"

    # Cover letter
    _cl_v = "✓ Generated" if ss.cover_letter else "—"
    _cl_c = "#00d4aa" if ss.cover_letter else "#6b7280"

    # Interview turns
    _beh_turns = len([m for m in ss.behavioral_messages if m["role"]=="assistant"])
    _tech_turns = len([m for m in ss.technical_messages if m["role"]=="assistant"])

    progress_items = [
        ("🎯 Resume Match",       _match_v,  _match_c),
        ("✉️ Cover Letter",       _cl_v,     _cl_c),
        ("🗣️ Behavioral Qs",      f"{_beh_turns} answered", "#00d4aa" if _beh_turns>0 else "#6b7280"),
        ("💻 Technical Qs",       f"{_tech_turns} answered", "#00d4aa" if _tech_turns>0 else "#6b7280"),
        ("📝 Exam Score",         _exam_v,   _exam_c),
        ("📊 BI Project",         _bi_v,     _bi_c),
        ("🗄️ SQL Project",        _sql_v,    _sql_c),
    ]

    for label, value, color in progress_items:
        st.markdown(f"""
        <div style="display:flex;justify-content:space-between;align-items:center;
             padding:5px 0;border-bottom:1px solid rgba(255,255,255,0.05);">
          <div style="font-size:0.75rem;color:#9ca3af;">{label}</div>
          <div style="font-size:0.75rem;font-weight:600;color:{color};">{value}</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Session save / restore ──────────────────────────────────────────────────
    st.markdown("""<div style="font-family:Syne,sans-serif;font-size:0.95rem;font-weight:700;
         color:#e8eaf0;margin-bottom:0.6rem;">💾 Session</div>""",
         unsafe_allow_html=True)

    # Save session
    if ss.resume_text or ss.analysis:
        snapshot = build_session_snapshot()
        st.download_button(
            "⬇️ Save Session (.json)",
            data=json.dumps(snapshot, indent=2, default=str),
            file_name=f"resumeiq_session_{(ss.analysis or {}).get('company','session').replace(' ','_')}.json",
            mime="application/json",
            use_container_width=True,
        )

    # Restore session
    restored_file = st.file_uploader(
        "↑ Restore Session", type=["json"],
        key="session_restore_upload",
        label_visibility="collapsed",
    )
    if restored_file:
        try:
            snap = json.loads(restored_file.read())
            if snap.get("version") == "1.0":
                if snap.get("bi_reviews"):
                    ss.proj_bi_reviews = snap["bi_reviews"]
                if snap.get("sql_reviews"):
                    ss.proj_sql_reviews = snap["sql_reviews"]
                if snap.get("cover_letter"):
                    ss.cover_letter = snap["cover_letter"]
                st.success("✓ Session restored!")
                st.rerun()
            else:
                st.error("Unrecognised session file format.")
        except Exception as e:
            st.error(f"Restore failed: {e}")

    # Project report downloads (shown if any reviews exist)
    if ss.proj_bi_reviews and ss.proj_bi_brief:
        bi_docx = build_project_report_docx(
            "Power BI Dashboard", ss.proj_bi_brief,
            BI_STAGES, ss.proj_bi_reviews, ss.proj_bi_submissions,
        )
        st.download_button(
            "⬇️ BI Project Report (.docx)",
            data=bi_docx,
            file_name="bi_project_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    if ss.proj_sql_reviews and ss.proj_sql_brief:
        sql_docx = build_project_report_docx(
            f"{ss.proj_sql_dialect} SQL Analytics", ss.proj_sql_brief,
            SQL_STAGES, ss.proj_sql_reviews, ss.proj_sql_submissions,
        )
        st.download_button(
            "⬇️ SQL Project Report (.docx)",
            data=sql_docx,
            file_name="sql_project_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

st.markdown("""
<div class="hero">
  <div class="hero-badge">AI-Powered · Career Intelligence Suite</div>
  <div class="hero-title">ResumeIQ</div>
  <div class="hero-sub">Resume analysis, tailored rewrites, behavioral coaching, technical interviews, and role-specific exams — all from one job URL.</div>
</div>
""", unsafe_allow_html=True)

# ── Module nav ─────────────────────────────────────────────────────────────────
modules = [
    ("resume",     "🎯 Resume Analyzer"),
    ("cover",      "✉️ Cover Letter"),
    ("behavioral", "🗣️ Behavioral Interview"),
    ("technical",  "💻 Technical Interview"),
    ("exam",       "📝 Technical Exam"),
    ("projects",   "🛠️ Technical Projects"),
]

cols = st.columns(len(modules))
for col, (key, label) in zip(cols, modules):
    with col:
        if st.button(label, key=f"nav_{key}", use_container_width=True):
            st.session_state.active_module = key
            st.rerun()

# Highlight active module
active_labels = {k: l for k, l in modules}
st.markdown(f"""
<div style="text-align:center; margin: 0.3rem 0 1.5rem;">
  <span style="font-family:'Syne',sans-serif; font-size:0.75rem; color:#a29bfe; font-weight:600;">
    Active: {active_labels.get(st.session_state.active_module, '')}
  </span>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SHARED INPUTS (always shown — needed for all modules)
# ══════════════════════════════════════════════════════════════════════════════
needs_inputs = not st.session_state.resume_text or not st.session_state.jd_text

if needs_inputs or st.session_state.active_module == "resume":
    with st.expander("⚙️ Setup — Resume & Job Description", expanded=needs_inputs):
        ic1, ic2 = st.columns(2, gap="large")
        with ic1:
            st.markdown('<div style="font-family:Syne,sans-serif;font-size:0.8rem;font-weight:700;color:#6c63ff;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:0.5rem;">Upload Resume</div>', unsafe_allow_html=True)
            uploaded = st.file_uploader("Resume (.docx or .txt)", type=["docx","txt"], label_visibility="collapsed")
            if uploaded:
                st.success(f"✓ {uploaded.name}")
        with ic2:
            st.markdown('<div style="font-family:Syne,sans-serif;font-size:0.8rem;font-weight:700;color:#6c63ff;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:0.5rem;">Job Posting</div>', unsafe_allow_html=True)
            job_url = st.text_input("Job URL", placeholder="https://boards.greenhouse.io/...", label_visibility="collapsed")
            st.markdown('<div style="font-size:0.75rem;color:#6b7280;margin:0.3rem 0;">— or paste directly —</div>', unsafe_allow_html=True)
            jd_paste = st.text_area("Job description text", placeholder="Paste job description here...", height=100, label_visibility="collapsed")

        _, load_col, _ = st.columns([2,1,2])
        with load_col:
            load_btn = st.button("🔄 Load & Analyze", use_container_width=True)

        if load_btn:
            if not uploaded:
                st.error("Please upload your resume.")
            elif not job_url and not jd_paste:
                st.error("Please provide a job URL or paste the description.")
            else:
                with st.spinner("Reading resume…"):
                    st.session_state.resume_text = extract_resume_text(uploaded)
                jd_text = ""
                if job_url:
                    with st.spinner("Scraping job description…"):
                        jd_text = scrape_job_description(job_url)
                        if jd_text.startswith("ERROR:"):
                            st.warning("Couldn't scrape URL — using pasted text.")
                            jd_text = jd_paste
                if not jd_text:
                    jd_text = jd_paste
                if not jd_text:
                    st.error("Could not get job description.")
                else:
                    st.session_state.jd_text = jd_text
                    with st.spinner("Running gap analysis…"):
                        try:
                            st.session_state.analysis = run_analysis(
                                st.session_state.resume_text, jd_text
                            )
                            # Reset interview/exam state when new JD loaded
                            st.session_state.behavioral_messages = []
                            st.session_state.behavioral_started = False
                            st.session_state.technical_messages = []
                            st.session_state.technical_started = False
                            st.session_state.hiring_messages = []
                            st.session_state.hiring_started = False
                            # Auto-detect technical domain from JD
                            st.session_state.tech_domain = detect_domain(jd_text)
                            st.session_state.tech_domain_override = None
                            st.session_state.exam_data = None
                            st.session_state.exam_submitted = False
                            st.session_state.exam_answers = {}
                            st.session_state.exam_sa_answers = {}
                            st.session_state.exam_started = False
                            st.rerun()
                        except Exception as e:
                            st.error(f"Analysis error: {e}")

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 1: RESUME ANALYZER
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.active_module == "resume":
    if not st.session_state.analysis:
        st.markdown("""
        <div style="text-align:center;padding:3rem;color:#6b7280;">
          <div style="font-size:2.5rem;margin-bottom:1rem;">📄</div>
          <div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;color:#e8eaf0;margin-bottom:0.5rem;">Upload your resume and a job URL to get started</div>
          <div style="font-size:0.9rem;">Use the Setup panel above</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        a = st.session_state.analysis
        score = a.get("match_score", 0)
        color = score_color(score)

        # Score header
        st.markdown(f"""
        <div style="text-align:center;margin-bottom:2rem;">
          <div style="font-family:Syne,sans-serif;font-size:0.72rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;color:#6b7280;margin-bottom:0.5rem;">Match Analysis</div>
          <div style="font-family:Syne,sans-serif;font-size:1.3rem;font-weight:700;color:#e8eaf0;margin-bottom:0.2rem;">{a.get('role_title','Role')} · {a.get('company','Company')}</div>
          <div style="font-family:Syne,sans-serif;font-size:4.5rem;font-weight:800;color:{color};line-height:1;">{score}</div>
          <div style="color:#6b7280;font-size:0.82rem;">Match Score / 100</div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""<div class="rec-card"><div class="rec-title">Executive Summary</div>{a.get('summary','')}</div>""", unsafe_allow_html=True)

        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown('<div class="section-hdr">✅ Strengths</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">✓ {s}</span>' for s in a.get("strengths",[])) + '</div>', unsafe_allow_html=True)
        with c2:
            st.markdown('<div class="section-hdr">⚠️ Partial Matches</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-yellow">~ {s}</span>' for s in a.get("partial_matches",[])) + '</div>', unsafe_allow_html=True)
        with c3:
            st.markdown('<div class="section-hdr">❌ Gaps</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">✗ {s}</span>' for s in a.get("gaps",[])) + '</div>', unsafe_allow_html=True)

        st.markdown("---")
        k1, k2 = st.columns(2)
        with k1:
            st.markdown('<div class="section-hdr">🔍 ATS Keywords — Present</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">{k}</span>' for k in a.get("ats_keywords_present",[])) + '</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-hdr" style="margin-top:1rem;">Missing Keywords</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">{k}</span>' for k in a.get("ats_keywords_missing",[])) + '</div>', unsafe_allow_html=True)
        with k2:
            st.markdown('<div class="section-hdr">🚨 Rejection Risks</div>', unsafe_allow_html=True)
            for risk in a.get("rejection_risks",[]):
                st.markdown(f'<div class="rec-card" style="border-left-color:#ff4d6d;"><div class="rec-title" style="color:#ff4d6d;">Risk</div>{risk}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown('<div class="section-hdr">💡 Recommendations</div>', unsafe_allow_html=True)
        for rec in a.get("recommendations",[]):
            st.markdown(f'<div class="rec-card"><div class="rec-title">{rec.get("area","")}</div>{rec.get("suggestion","")}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("""<div style="text-align:center;margin-bottom:1.2rem;"><div style="font-family:Syne,sans-serif;font-size:1.2rem;font-weight:700;color:#e8eaf0;">✨ Generate Tailored Resume</div><div style="color:#6b7280;font-size:0.85rem;margin-top:0.3rem;">Reframes your existing experience — no fabrication.</div></div>""", unsafe_allow_html=True)
        _, gc, _ = st.columns([2,1,2])
        with gc:
            if st.button("🚀 Generate Tailored Resume", use_container_width=True):
                with st.spinner("Crafting your tailored resume…"):
                    st.session_state.tailored_resume = generate_tailored_resume(
                        st.session_state.resume_text, st.session_state.jd_text, a
                    )

        if st.session_state.tailored_resume:
            st.markdown("---")
            t1, t2 = st.tabs(["📝 Tailored Resume", "📊 Side-by-Side"])
            with t1:
                st.markdown(f'<div class="resume-preview">{st.session_state.tailored_resume}</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                d1, d2 = st.columns(2)
                with d1:
                    st.download_button("⬇️ Download .txt", data=st.session_state.tailored_resume.encode(), file_name=f"tailored_{a.get('company','role').replace(' ','_')}.txt", mime="text/plain", use_container_width=True)
                with d2:
                    st.download_button("⬇️ Download .docx", data=build_docx(st.session_state.tailored_resume), file_name=f"tailored_{a.get('company','role').replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            with t2:
                oc, tc = st.columns(2)
                with oc:
                    st.markdown("**Original**")
                    st.markdown(f'<div class="resume-preview">{st.session_state.resume_text}</div>', unsafe_allow_html=True)
                with tc:
                    st.markdown("**Tailored**")
                    st.markdown(f'<div class="resume-preview">{st.session_state.tailored_resume}</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 1b: COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "cover":
    st.markdown("""
    <div style="margin-bottom:1.2rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.3rem;">✉️ Cover Letter Generator</div>
      <div style="color:#6b7280;font-size:0.9rem;">Role-specific, humanized — written to sound like you, not an AI. Two-pass generation: draft → humanization layer.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # ── Settings ───────────────────────────────────────────────────────────
        st.markdown('<div class="section-hdr">⚙️ Cover Letter Settings</div>', unsafe_allow_html=True)
        cfg1, cfg2, cfg3 = st.columns(3)

        with cfg1:
            candidate_name = st.text_input(
                "Your full name",
                value=st.session_state.cover_letter_name or "",
                placeholder="Philip A. Jones",
                key="cl_name"
            )
            st.session_state.cover_letter_name = candidate_name

        with cfg2:
            # Pre-fill from analysis if available
            default_company = (st.session_state.analysis or {}).get("company","") or st.session_state.cover_letter_company
            default_role    = (st.session_state.analysis or {}).get("role_title","") or st.session_state.cover_letter_role
            company_input = st.text_input("Company name", value=default_company, placeholder="Garner Health", key="cl_company")
            st.session_state.cover_letter_company = company_input

        with cfg3:
            role_input = st.text_input("Role / position", value=default_role, placeholder="Data Analyst II", key="cl_role")
            st.session_state.cover_letter_role = role_input

        tone_col, _ = st.columns([2,3])
        with tone_col:
            tone = st.selectbox(
                "Writing tone",
                ["Professional & Confident", "Warm & Conversational", "Bold & Direct", "Storytelling"],
                index=["Professional & Confident","Warm & Conversational","Bold & Direct","Storytelling"].index(
                    st.session_state.cover_letter_tone
                ),
                key="cl_tone"
            )
            st.session_state.cover_letter_tone = tone

        # Tone descriptions
        tone_descs = {
            "Professional & Confident": "Assured and direct. Accomplishments speak for themselves. No fluff.",
            "Warm & Conversational":    "Approachable and genuine — like a smart colleague, not a press release.",
            "Bold & Direct":            "Sharp and punchy. Gets to the point in sentence one. Zero pleasantries.",
            "Storytelling":             "Opens with a brief specific moment or story, then pivots to qualifications.",
        }
        st.markdown(f'<div style="font-size:0.8rem;color:#6b7280;margin-top:0.2rem;margin-bottom:1rem;">💡 {tone_descs[tone]}</div>', unsafe_allow_html=True)

        # Humanization info box
        st.markdown("""
        <div style="background:rgba(0,212,170,0.06);border:1px solid rgba(0,212,170,0.2);border-left:3px solid #00d4aa;
             border-radius:0 10px 10px 0;padding:0.8rem 1rem;margin-bottom:1.2rem;font-size:0.85rem;color:#c8cad4;">
          <div style="font-weight:600;color:#00d4aa;font-size:0.75rem;text-transform:uppercase;letter-spacing:0.08em;margin-bottom:0.3rem;">
            🧠 Two-Pass Humanization Active
          </div>
          Every output runs through a dedicated humanization layer that eliminates AI tells — 
          hollow superlatives, robotic transitions, formulaic openers, corporate buzzwords — 
          then varies sentence rhythm and injects specificity. The same pass is applied to tailored resumes.
        </div>
        """, unsafe_allow_html=True)

        # Generate button
        _, gen_c, _ = st.columns([2,1,2])
        with gen_c:
            gen_cl_btn = st.button("✉️ Generate Cover Letter", use_container_width=True)

        if gen_cl_btn:
            if not candidate_name.strip():
                st.error("Please enter your name.")
            else:
                with st.spinner("Writing first draft…"):
                    pass
                progress = st.progress(0, text="Drafting cover letter…")
                import time as _time
                _time.sleep(0.3)
                progress.progress(35, text="Tailoring to job description…")

                try:
                    cl_draft = generate_cover_letter(
                        resume_text   = st.session_state.resume_text,
                        jd_text       = st.session_state.jd_text,
                        candidate_name= candidate_name,
                        company       = company_input or "the company",
                        role          = role_input    or "this role",
                        tone          = tone,
                        analysis      = st.session_state.analysis,
                    )
                    progress.progress(100, text="Done!")
                    _time.sleep(0.4)
                    progress.empty()
                    st.session_state.cover_letter = cl_draft
                    st.rerun()
                except Exception as e:
                    progress.empty()
                    st.error(f"Generation failed: {e}")

        # ── Output ─────────────────────────────────────────────────────────────
        if st.session_state.cover_letter:
            st.markdown("---")
            st.markdown('<div class="section-hdr">✉️ Your Cover Letter</div>', unsafe_allow_html=True)

            # Editable preview
            edited_cl = st.text_area(
                "Edit directly if needed",
                value=st.session_state.cover_letter,
                height=400,
                key="cl_editor",
                label_visibility="collapsed"
            )

            # Action row
            act1, act2, act3, act4 = st.columns(4)

            with act1:
                # Re-humanize
                if st.button("🔄 Re-Humanize", use_container_width=True):
                    with st.spinner("Running humanization pass…"):
                        st.session_state.cover_letter = humanize_text(
                            edited_cl, tone=tone, doc_type="cover letter"
                        )
                    st.rerun()

            with act2:
                # Regenerate fresh
                if st.button("⚡ Regenerate", use_container_width=True):
                    with st.spinner("Generating new version…"):
                        st.session_state.cover_letter = generate_cover_letter(
                            st.session_state.resume_text,
                            st.session_state.jd_text,
                            candidate_name,
                            company_input or "the company",
                            role_input    or "this role",
                            tone,
                            st.session_state.analysis,
                        )
                    st.rerun()

            with act3:
                st.download_button(
                    "⬇️ Download .txt",
                    data=edited_cl.encode("utf-8"),
                    file_name=f"cover_letter_{(company_input or 'role').replace(' ','_')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )

            with act4:
                docx_cl = build_cover_letter_docx(
                    edited_cl,
                    candidate_name or "Candidate",
                    role_input     or "Position",
                    company_input  or "Company",
                )
                st.download_button(
                    "⬇️ Download .docx",
                    data=docx_cl,
                    file_name=f"cover_letter_{(company_input or 'role').replace(' ','_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            # Word count
            wc = len(edited_cl.split())
            wc_color = "#00d4aa" if 240 <= wc <= 350 else "#ffa940" if wc < 240 else "#ff4d6d"
            st.markdown(
                f'<div style="text-align:right;font-size:0.78rem;color:{wc_color};margin-top:0.4rem;">'
                f'Word count: {wc} {"✓ ideal range" if 240 <= wc <= 350 else "⚠ aim for 250–350"}'
                f'</div>',
                unsafe_allow_html=True
            )

            # AI-detection tips
            with st.expander("🔍 What was humanized — and how to spot AI writing yourself"):
                st.markdown("""
**AI tells eliminated in your letter:**

**Hollow openers removed** — phrases like *"I am writing to express my interest"* or *"In today's competitive landscape"* are the first thing screeners notice.

**Superlative culling** — words like *passionate*, *innovative*, *dynamic*, *leverage*, *spearhead*, *utilize* read as AI filler. They've been replaced with specific, grounded language.

**Transition word purge** — *Furthermore, Moreover, Additionally, In conclusion* are robotic. Human writers use shorter connective tissue or just start a new thought.

**Sentence rhythm variation** — AI tends to write in consistent medium-length sentences. Human writing mixes short punchy sentences with longer ones deliberately.

**The "I" chain broken** — starting five consecutive sentences with "I" is a clear AI pattern. The humanizer varies sentence openers.

**Formulaic closer replaced** — *"I look forward to hearing from you"* is the most-detected AI phrase in cover letters. Your letter closes with something earned.

**To check your own writing:** Paste into [Grammarly](https://grammarly.com) or [Quillbot](https://quillbot.com/ai-content-detector) — though no detector is perfect, the humanization pass significantly reduces detection likelihood.
                """)

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 2: BEHAVIORAL INTERVIEW AGENT  (with avatar)
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "behavioral":
    st.markdown("""
    <div style="margin-bottom:1rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.3rem;">🗣️ Behavioral Interview — Full Panel</div>
      <div style="color:#6b7280;font-size:0.9rem;">Practice with three distinct interviewers — each with their own personality, focus, and voice.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # ── Interviewer selector ───────────────────────────────────────────────
        psel1, psel2, psel3 = st.columns(3)
        persona_map = {
            "Sarah Chen — HR Coach":          "behavioral",
            "Marcus Rivera — Technical":       "technical",
            "Diana Wells — Hiring Manager":    "hiring",
        }
        for col, (label, key) in zip([psel1, psel2, psel3], persona_map.items()):
            with col:
                active = st.session_state.behavioral_persona == key
                border = "rgba(108,99,255,0.7)" if active else "rgba(255,255,255,0.08)"
                bg     = "rgba(108,99,255,0.15)" if active else "rgba(255,255,255,0.03)"
                icons  = {"behavioral":"👩‍💼", "technical":"👨‍💻", "hiring":"👩‍🏫"}
                subtitles = {
                    "behavioral": "STAR method · Culture · Competencies",
                    "technical":  "Tech stack · Problem solving · Code",
                    "hiring":     "Motivation · Vision · Leadership",
                }
                st.markdown(f"""<div style="background:{bg};border:1px solid {border};border-radius:12px;
                    padding:0.9rem 1rem;cursor:pointer;text-align:center;">
                  <div style="font-size:1.4rem;">{icons[key]}</div>
                  <div style="font-family:Syne,sans-serif;font-size:0.85rem;font-weight:700;
                       color:#e8eaf0;margin:0.2rem 0;">{label}</div>
                  <div style="font-size:0.72rem;color:#6b7280;">{subtitles[key]}</div>
                </div>""", unsafe_allow_html=True)
                if st.button(f"{'✓ Active' if active else 'Select'}", key=f"psel_{key}", use_container_width=True):
                    st.session_state.behavioral_persona = key
                    st.session_state.behavioral_messages = []
                    st.session_state.behavioral_started  = False
                    st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        # ── STAR quick ref ─────────────────────────────────────────────────────
        with st.expander("📖 STAR Method Quick Reference"):
            sc1, sc2, sc3, sc4 = st.columns(4)
            for col, letter, name, desc in [
                (sc1,"S","Situation","Set the scene — context, timeframe, your role"),
                (sc2,"T","Task","What was required of you specifically?"),
                (sc3,"A","Action","What YOU did — be specific, use 'I' not 'we'"),
                (sc4,"R","Result","Quantified outcome — % improvement, time saved, etc."),
            ]:
                with col:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-family:Syne,sans-serif;font-size:1.8rem;font-weight:800;color:#6c63ff;">{letter}</div>
                      <div style="font-weight:600;color:#e8eaf0;font-size:0.85rem;">{name}</div>
                      <div style="color:#6b7280;font-size:0.78rem;margin-top:0.3rem;">{desc}</div>
                    </div>""", unsafe_allow_html=True)

        # ── Start / reset ──────────────────────────────────────────────────────
        bc1, bc2, _ = st.columns([1,1,3])
        with bc1:
            btn_label = "▶️ Start Interview" if not st.session_state.behavioral_started else "🔄 New Session"
            if st.button(btn_label, use_container_width=True):
                st.session_state.behavioral_messages = []
                st.session_state.behavioral_started  = True
                with st.spinner("Preparing first question…"):
                    role    = st.session_state.analysis.get("role_title","this role") if st.session_state.analysis else "this role"
                    company = st.session_state.analysis.get("company","the company") if st.session_state.analysis else "the company"
                    persona = st.session_state.behavioral_persona
                    persona_intros = {
                        "behavioral": f"Please start the behavioral interview. I'm interviewing for the {role} position at {company}. Focus on STAR-method behavioral questions.",
                        "technical":  f"Please start a technical interview for the {role} position at {company}. Focus on technical skills and problem solving.",
                        "hiring":     f"Please start a final-round hiring manager interview for the {role} position at {company}. Focus on culture fit, motivation, and leadership.",
                    }
                    opener = persona_intros.get(persona, persona_intros["behavioral"])
                    st.session_state.behavioral_messages.append({"role":"user","content":opener})
                    # Route to correct chat function
                    if persona == "technical":
                        reply = technical_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    elif persona == "hiring":
                        reply = hiring_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    else:
                        reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                st.rerun()

        # ── Avatar + chat ──────────────────────────────────────────────────────
        if st.session_state.behavioral_started and st.session_state.behavioral_messages:

            # Find the latest AI message to show in avatar
            latest_ai = ""
            for m in reversed(st.session_state.behavioral_messages):
                if m["role"] == "assistant":
                    latest_ai = m["content"]
                    break

            # Two-column layout: avatar left, chat right
            av_col, chat_col = st.columns([2, 3], gap="large")

            with av_col:
                st.markdown('<div class="section-hdr">Your Interviewer</div>', unsafe_allow_html=True)
                render_avatar(st.session_state.behavioral_persona, latest_ai)

            with chat_col:
                st.markdown('<div class="section-hdr">Interview Session</div>', unsafe_allow_html=True)
                st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
                for msg in st.session_state.behavioral_messages:
                    if msg["role"] == "assistant":
                        persona_names = {"behavioral":"Sarah Chen","technical":"Marcus Rivera","hiring":"Diana Wells"}
                        pname = persona_names.get(st.session_state.behavioral_persona,"Interviewer")
                        st.markdown(f"""<div class="bubble bubble-ai">
                          <div class="bubble-label">{pname}</div>
                          {msg["content"].replace(chr(10),"<br>")}
                        </div>""", unsafe_allow_html=True)
                    elif msg["role"] == "user" and not msg["content"].startswith("Please start"):
                        st.markdown(f"""<div class="bubble bubble-user">
                          <div class="bubble-label">You</div>
                          {msg["content"].replace(chr(10),"<br>")}
                        </div>""", unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)
                user_input = st.text_area("Your answer", placeholder="Type your response here...",
                                          height=130, key="behavioral_input", label_visibility="collapsed")
                ri1, ri2, ri3 = st.columns([1,1,1])
                with ri1:
                    send_btn = st.button("📤 Send Response", use_container_width=True)
                with ri2:
                    hint_btn = st.button("💡 Hint", use_container_width=True)
                with ri3:
                    next_btn = st.button("⏭️ Next Question", use_container_width=True)

                if send_btn and user_input.strip():
                    st.session_state.behavioral_messages.append({"role":"user","content":user_input.strip()})
                    with st.spinner("Analyzing response…"):
                        persona = st.session_state.behavioral_persona
                        if persona == "technical":
                            reply = technical_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        elif persona == "hiring":
                            reply = hiring_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        else:
                            reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                    st.rerun()

                if hint_btn:
                    st.session_state.behavioral_messages.append({"role":"user","content":"Give me a brief hint — what are you looking for?"})
                    with st.spinner("Getting hint…"):
                        persona = st.session_state.behavioral_persona
                        if persona == "technical":
                            reply = technical_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        elif persona == "hiring":
                            reply = hiring_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        else:
                            reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                    st.rerun()

                if next_btn:
                    st.session_state.behavioral_messages.append({"role":"user","content":"Please move on to your next question."})
                    with st.spinner("Next question…"):
                        persona = st.session_state.behavioral_persona
                        if persona == "technical":
                            reply = technical_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        elif persona == "hiring":
                            reply = hiring_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                        else:
                            reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                    st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3: TECHNICAL INTERVIEW — GPT-4o Domain Engine
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "technical":

    # ── Resolve active domain ──────────────────────────────────────────────────
    active_domain = st.session_state.tech_domain_override or st.session_state.tech_domain or "analytics_general"
    profile = DOMAIN_PROFILES[active_domain]
    using_gpt = bool(st.session_state.openai_key)

    st.markdown(f"""
    <div style="margin-bottom:1rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;
           color:#e8eaf0;margin-bottom:0.3rem;">
        💻 Technical Interview — {profile['icon']} {profile['label']}
      </div>
      <div style="color:#6b7280;font-size:0.9rem;">
        {'⚡ Powered by GPT-4o — domain-specialized code, debug, and scenario questions'
         if using_gpt else
         '🧠 Using Claude (add OpenAI key in sidebar to unlock GPT-4o domain engine)'}
      </div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # ── Domain selector ────────────────────────────────────────────────────
        st.markdown('<div class="section-hdr">🎯 Technical Domain</div>', unsafe_allow_html=True)

        # Show auto-detected domain
        auto_domain = st.session_state.tech_domain or "analytics_general"
        auto_profile = DOMAIN_PROFILES[auto_domain]
        st.markdown(f"""
        <div style="background:rgba(108,99,255,0.07);border:1px solid rgba(108,99,255,0.2);
             border-radius:10px;padding:0.7rem 1rem;margin-bottom:0.8rem;
             display:flex;align-items:center;gap:0.7rem;">
          <div style="font-size:1.3rem;">{auto_profile['icon']}</div>
          <div>
            <div style="font-size:0.72rem;color:#a29bfe;text-transform:uppercase;
                 letter-spacing:0.1em;font-weight:600;">Auto-detected from job description</div>
            <div style="font-size:0.92rem;font-weight:600;color:#e8eaf0;">{auto_profile['label']}</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # Domain override selector
        domain_options = {f"{v['icon']} {v['label']}": k for k, v in DOMAIN_PROFILES.items()}
        override_label = f"{profile['icon']} {profile['label']}"
        sel_domain_label = st.selectbox(
            "Override domain (optional)",
            list(domain_options.keys()),
            index=list(domain_options.keys()).index(override_label),
            key="domain_override_sel",
        )
        selected_domain_key = domain_options[sel_domain_label]
        if selected_domain_key != active_domain:
            st.session_state.tech_domain_override = selected_domain_key
            st.session_state.technical_messages  = []
            st.session_state.technical_started   = False
            st.rerun()

        # ── Topic pool preview ─────────────────────────────────────────────────
        with st.expander(f"📚 {profile['label']} topic pool ({len(profile['topic_pool'])} topics)"):
            cols_t = st.columns(2)
            half = len(profile["topic_pool"]) // 2
            with cols_t[0]:
                for t in profile["topic_pool"][:half]:
                    st.markdown(f'<span class="pill pill-purple" style="display:inline-flex;margin:2px 0;">◆ {t}</span>', unsafe_allow_html=True)
            with cols_t[1]:
                for t in profile["topic_pool"][half:]:
                    st.markdown(f'<span class="pill pill-purple" style="display:inline-flex;margin:2px 0;">◆ {t}</span>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # ── Controls row ───────────────────────────────────────────────────────
        ctrl1, ctrl2, ctrl3, _ = st.columns([1, 1, 1, 2])
        with ctrl1:
            btn_label = "▶️ Start" if not st.session_state.technical_started else "🔄 Reset"
            start_btn = st.button(btn_label, use_container_width=True)
        with ctrl2:
            diff = st.selectbox(
                "Difficulty",
                ["Progressive", "Junior", "Mid-level", "Senior", "Staff/Principal"],
                key="tech_diff_sel",
                label_visibility="collapsed",
            )
        with ctrl3:
            fmt_options = ["All formats", "Conceptual only", "Code writing only", "Debug only", "Scenario only", "Optimization only"]
            fmt_filter = st.selectbox("Question type", fmt_options, key="tech_fmt_filter", label_visibility="collapsed")

        if start_btn:
            st.session_state.technical_messages = []
            st.session_state.technical_started  = True
            active_domain = st.session_state.tech_domain_override or st.session_state.tech_domain or "analytics_general"
            with st.spinner(f"GPT-4o preparing first {profile['label']} question…" if using_gpt else "Preparing first question…"):
                role = (st.session_state.analysis or {}).get("role_title", "this role")
                fmt_note = f" Focus exclusively on {fmt_filter.replace(' only','').lower()} questions." if fmt_filter != "All formats" else ""
                opener = (
                    f"Start the technical interview for the {role} position. "
                    f"Domain: {profile['label']}. Difficulty: {diff}.{fmt_note} "
                    f"Ask your first question now."
                )
                st.session_state.technical_messages.append({"role": "user", "content": opener})
                if using_gpt:
                    reply = gpt4o_technical_chat(
                        st.session_state.technical_messages,
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                        active_domain,
                        diff,
                        st.session_state.openai_key,
                    )
                else:
                    reply = technical_chat(
                        st.session_state.technical_messages,
                        st.session_state.resume_text,
                        st.session_state.jd_text,
                    )
                st.session_state.technical_messages.append({"role": "assistant", "content": reply})
            st.rerun()

        # ── Interview panel ────────────────────────────────────────────────────
        if st.session_state.technical_started and st.session_state.technical_messages:
            latest_tech_ai = next(
                (m["content"] for m in reversed(st.session_state.technical_messages) if m["role"] == "assistant"), ""
            )

            av_col2, chat_col2 = st.columns([2, 3], gap="large")

            with av_col2:
                st.markdown('<div class="section-hdr">Technical Interviewer</div>', unsafe_allow_html=True)
                render_avatar("technical", latest_tech_ai)

                # Domain badge
                st.markdown(f"""
                <div style="background:rgba(0,212,170,0.08);border:1px solid rgba(0,212,170,0.2);
                     border-radius:8px;padding:0.5rem 0.8rem;margin-top:0.5rem;text-align:center;">
                  <div style="font-size:0.7rem;color:#6b7280;text-transform:uppercase;
                       letter-spacing:0.1em;">Active domain</div>
                  <div style="font-size:0.9rem;font-weight:600;color:#00d4aa;">
                    {profile['icon']} {profile['label']}
                  </div>
                  <div style="font-size:0.7rem;color:#6b7280;margin-top:2px;">
                    {'GPT-4o' if using_gpt else 'Claude fallback'}
                  </div>
                </div>
                """, unsafe_allow_html=True)

                # Question counter
                q_count = sum(1 for m in st.session_state.technical_messages if m["role"] == "assistant")
                st.markdown(f"""
                <div style="text-align:center;margin-top:0.6rem;font-size:0.8rem;color:#6b7280;">
                  Questions asked: <strong style="color:#e8eaf0;">{q_count}</strong>
                </div>
                """, unsafe_allow_html=True)

            with chat_col2:
                st.markdown('<div class="section-hdr">Technical Interview Session</div>', unsafe_allow_html=True)
                st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
                for msg in st.session_state.technical_messages:
                    if msg["role"] == "assistant":
                        content = msg["content"].replace(chr(10), "<br>")
                        st.markdown(f"""<div class="bubble bubble-ai">
                          <div class="bubble-label">Marcus Rivera · {profile['label']}</div>
                          {content}
                        </div>""", unsafe_allow_html=True)
                    elif msg["role"] == "user" and not msg["content"].startswith("Start the technical"):
                        st.markdown(f"""<div class="bubble bubble-user">
                          <div class="bubble-label">You</div>
                          {msg["content"].replace(chr(10),"<br>")}
                        </div>""", unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)

                st.markdown("<br>", unsafe_allow_html=True)

                # Answer input — larger for code answers
                tech_input = st.text_area(
                    "Your answer",
                    placeholder=f"Type your answer here. For code questions, write your {profile['label']} code directly...",
                    height=170,
                    key="technical_input",
                    label_visibility="collapsed",
                )

                ti1, ti2, ti3, ti4 = st.columns(4)

                def _send_tech(user_msg: str, spinner_text: str = "Evaluating…"):
                    active_d = st.session_state.tech_domain_override or st.session_state.tech_domain or "analytics_general"
                    st.session_state.technical_messages.append({"role": "user", "content": user_msg})
                    with st.spinner(spinner_text):
                        if using_gpt:
                            r = gpt4o_technical_chat(
                                st.session_state.technical_messages,
                                st.session_state.resume_text,
                                st.session_state.jd_text,
                                active_d, diff,
                                st.session_state.openai_key,
                            )
                        else:
                            r = technical_chat(
                                st.session_state.technical_messages,
                                st.session_state.resume_text,
                                st.session_state.jd_text,
                            )
                    st.session_state.technical_messages.append({"role": "assistant", "content": r})
                    st.rerun()

                with ti1:
                    if st.button("📤 Submit", use_container_width=True) and tech_input.strip():
                        _send_tech(tech_input.strip(), "Evaluating with GPT-4o…" if using_gpt else "Evaluating…")
                with ti2:
                    if st.button("⏭️ Skip", use_container_width=True):
                        _send_tech("Skip this question — ask me the next one.")
                with ti3:
                    if st.button("🔍 Ideal Answer", use_container_width=True):
                        _send_tech(
                            "Show me the complete ideal answer for the last question — include working code with explanation.",
                            "Generating ideal answer…"
                        )
                with ti4:
                    if st.button("🐛 New Debug Q", use_container_width=True):
                        _send_tech(
                            f"Give me a debug question — show me a broken {profile['label']} code snippet and ask me to find and fix the bug.",
                            "Generating debug question…"
                        )



# ══════════════════════════════════════════════════════════════════════════════
# MODULE 4: TECHNICAL EXAM
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "exam":
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.4rem;">📝 Technical Exam</div>
      <div style="color:#6b7280;font-size:0.9rem;">A timed, role-specific exam generated from the job description. Multiple choice + short answer. Scored with full explanations.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # Generate exam
        if not st.session_state.exam_data:
            _, gen_col, _ = st.columns([2,1,2])
            with gen_col:
                if st.button("⚡ Generate Exam", use_container_width=True):
                    with st.spinner("Building your custom exam from the job description…"):
                        try:
                            st.session_state.exam_data = generate_exam(st.session_state.jd_text, st.session_state.resume_text)
                            st.session_state.exam_answers = {}
                            st.session_state.exam_sa_answers = {}
                            st.session_state.exam_submitted = False
                            st.session_state.exam_started = False
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to generate exam: {e}")
        else:
            exam = st.session_state.exam_data

            # Exam header
            if not st.session_state.exam_submitted:
                eh1, eh2, eh3 = st.columns([3,1,1])
                with eh1:
                    st.markdown(f"""<div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;color:#e8eaf0;">{exam.get('title','Technical Assessment')}</div>
                    <div style="color:#6b7280;font-size:0.82rem;margin-top:0.3rem;">{exam.get('instructions','')}</div>""", unsafe_allow_html=True)
                with eh2:
                    total_q = sum(len(s.get("questions",[])) for s in exam.get("sections",[]))
                    answered = len(st.session_state.exam_answers) + len({k for k,v in st.session_state.exam_sa_answers.items() if v.strip()})
                    st.metric("Progress", f"{answered}/{total_q}")
                with eh3:
                    st.metric("Duration", f"{exam.get('duration_minutes',30)} min")

                if not st.session_state.exam_started:
                    _, sc, _ = st.columns([2,1,2])
                    with sc:
                        if st.button("▶️ Begin Exam", use_container_width=True):
                            st.session_state.exam_started = True
                            st.session_state.exam_start_time = time.time()
                            st.rerun()
                else:
                    # Timer display
                    elapsed = int(time.time() - (st.session_state.exam_start_time or time.time()))
                    limit   = exam.get("duration_minutes", 30) * 60
                    remain  = max(0, limit - elapsed)
                    mins, secs = divmod(remain, 60)
                    timer_color = "#ff4d6d" if remain < 300 else "#ffa940" if remain < 600 else "#00d4aa"
                    st.markdown(f"""<div style="text-align:right;font-family:Syne,sans-serif;font-size:1rem;font-weight:700;color:{timer_color};margin-bottom:1rem;">⏱ {mins:02d}:{secs:02d} remaining</div>""", unsafe_allow_html=True)

                    # Questions
                    for section in exam.get("sections", []):
                        st.markdown(f'<div class="section-hdr">{section.get("name","Section")}</div>', unsafe_allow_html=True)

                        for q in section.get("questions", []):
                            qid   = q.get("id","")
                            qtype = q.get("type","multiple_choice")
                            diff_badge = {"easy":"🟢","medium":"🟡","hard":"🔴"}.get(q.get("difficulty","medium"),"🟡")

                            st.markdown(f"""<div class="q-card">
                              <div class="q-num">{qid.upper()}  ·  {q.get('topic','')}  {diff_badge} {q.get('difficulty','').capitalize()}</div>
                              <div class="q-text">{q.get('question','')}</div>""", unsafe_allow_html=True)

                            if qtype == "multiple_choice":
                                opts = q.get("options", {})
                                choices = [f"{k}: {v}" for k, v in opts.items()]
                                current = st.session_state.exam_answers.get(qid)
                                idx = None
                                if current:
                                    try: idx = choices.index(next(c for c in choices if c.startswith(current)))
                                    except: idx = None
                                answer = st.radio("Select answer", choices, index=idx, key=f"mc_{qid}", label_visibility="collapsed")
                                if answer:
                                    st.session_state.exam_answers[qid] = answer[0]  # store just the letter
                            else:
                                sa_val = st.session_state.exam_sa_answers.get(qid, "")
                                new_val = st.text_area("Your answer", value=sa_val, height=100, key=f"sa_{qid}", label_visibility="collapsed", placeholder="Type your answer here...")
                                st.session_state.exam_sa_answers[qid] = new_val

                            st.markdown("</div>", unsafe_allow_html=True)

                    # Submit
                    st.markdown("<br>", unsafe_allow_html=True)
                    _, sub_col, _ = st.columns([2,1,2])
                    with sub_col:
                        if st.button("✅ Submit Exam", use_container_width=True):
                            st.session_state.exam_submitted = True
                            # Grade short answers
                            for section in exam.get("sections",[]):
                                for q in section.get("questions",[]):
                                    if q.get("type") == "short_answer":
                                        qid = q.get("id","")
                                        user_ans = st.session_state.exam_sa_answers.get(qid,"")
                                        if user_ans.strip():
                                            with st.spinner(f"Grading {qid}…"):
                                                grade = grade_short_answer(
                                                    q.get("question",""),
                                                    q.get("sample_answer",""),
                                                    q.get("key_points",[]),
                                                    user_ans
                                                )
                                                st.session_state.exam_sa_grades[qid] = grade
                            st.rerun()

            # ── RESULTS ───────────────────────────────────────────────────────
            else:
                exam = st.session_state.exam_data
                mc_correct = 0
                mc_total   = 0
                sa_score   = 0
                sa_total   = 0

                for section in exam.get("sections",[]):
                    for q in section.get("questions",[]):
                        if q.get("type") == "multiple_choice":
                            mc_total += 1
                            if st.session_state.exam_answers.get(q.get("id","")) == q.get("correct",""):
                                mc_correct += 1
                        elif q.get("type") == "short_answer":
                            grade = st.session_state.exam_sa_grades.get(q.get("id",""),{})
                            sa_score += grade.get("score",0)
                            sa_total += 10

                total_possible = mc_total * 10 + sa_total
                total_earned   = mc_correct * 10 + sa_score
                overall_pct    = int(total_earned / total_possible * 100) if total_possible > 0 else 0
                grade_letter   = "A" if overall_pct >= 90 else "B" if overall_pct >= 80 else "C" if overall_pct >= 70 else "D" if overall_pct >= 60 else "F"
                res_color      = score_color(overall_pct)

                st.markdown(f"""
                <div style="text-align:center;padding:2rem 0 1.5rem;">
                  <div style="font-family:Syne,sans-serif;font-size:0.72rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;color:#6b7280;margin-bottom:0.5rem;">Exam Results</div>
                  <div style="font-family:Syne,sans-serif;font-size:5rem;font-weight:800;color:{res_color};line-height:1;">{overall_pct}%</div>
                  <div style="font-family:Syne,sans-serif;font-size:1.5rem;font-weight:700;color:{res_color};">Grade: {grade_letter}</div>
                  <div style="color:#6b7280;font-size:0.85rem;margin-top:0.5rem;">{total_earned}/{total_possible} points</div>
                </div>
                """, unsafe_allow_html=True)

                rm1, rm2, rm3 = st.columns(3)
                with rm1:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Multiple Choice</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{score_color(int(mc_correct/mc_total*100) if mc_total else 0)};">{mc_correct}/{mc_total}</div>
                    </div>""", unsafe_allow_html=True)
                with rm2:
                    sa_pct = int(sa_score/sa_total*100) if sa_total else 0
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Short Answer</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{score_color(sa_pct)};">{sa_score}/{sa_total}</div>
                    </div>""", unsafe_allow_html=True)
                with rm3:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Overall Score</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{res_color};">{overall_pct}%</div>
                    </div>""", unsafe_allow_html=True)

                st.markdown("---")
                st.markdown('<div class="section-hdr">📋 Detailed Review</div>', unsafe_allow_html=True)

                for section in exam.get("sections",[]):
                    st.markdown(f"**{section.get('name','')}**")
                    for q in section.get("questions",[]):
                        qid   = q.get("id","")
                        qtype = q.get("type","multiple_choice")

                        if qtype == "multiple_choice":
                            user_ans    = st.session_state.exam_answers.get(qid,"—")
                            correct_ans = q.get("correct","")
                            is_correct  = user_ans == correct_ans
                            border_col  = "#00d4aa" if is_correct else "#ff4d6d"
                            icon        = "✅" if is_correct else "❌"
                            opts        = q.get("options",{})
                            opts_html   = " · ".join(f"<strong>{k}</strong>: {v}" for k,v in opts.items())

                            st.markdown(f"""<div class="q-card" style="border-left:3px solid {border_col};">
                              <div class="q-num">{qid.upper()} — {icon} {'Correct' if is_correct else 'Incorrect'}</div>
                              <div class="q-text">{q.get('question','')}</div>
                              <div style="font-size:0.82rem;color:#6b7280;margin-bottom:0.5rem;">{opts_html}</div>
                              <div style="font-size:0.85rem;margin-bottom:0.3rem;">Your answer: <strong style="color:{'#00d4aa' if is_correct else '#ff4d6d'};">{user_ans}</strong> · Correct: <strong style="color:#00d4aa;">{correct_ans}</strong></div>
                              <div style="font-size:0.83rem;color:#a0a8b8;background:rgba(255,255,255,0.04);padding:0.6rem 0.8rem;border-radius:8px;">💡 {q.get('explanation','')}</div>
                            </div>""", unsafe_allow_html=True)

                        elif qtype == "short_answer":
                            grade    = st.session_state.exam_sa_grades.get(qid,{})
                            sa_score_q = grade.get("score",0)
                            grade_label = grade.get("grade","Not graded")
                            grade_color = score_color(sa_score_q * 10)

                            st.markdown(f"""<div class="q-card" style="border-left:3px solid {grade_color};">
                              <div class="q-num">{qid.upper()} — Short Answer · {grade_label} ({sa_score_q}/10)</div>
                              <div class="q-text">{q.get('question','')}</div>""", unsafe_allow_html=True)

                            st.markdown("**Your answer:**")
                            st.markdown(f'<div style="background:rgba(255,255,255,0.04);padding:0.7rem 0.9rem;border-radius:8px;font-size:0.85rem;color:#c8cad4;margin-bottom:0.6rem;">{st.session_state.exam_sa_answers.get(qid,"No answer provided")}</div>', unsafe_allow_html=True)

                            if grade:
                                covered = grade.get("points_covered",[])
                                missed  = grade.get("points_missed",[])
                                if covered:
                                    st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">✓ {p}</span>' for p in covered) + '</div>', unsafe_allow_html=True)
                                if missed:
                                    st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">✗ {p}</span>' for p in missed) + '</div>', unsafe_allow_html=True)
                                st.markdown(f'<div style="font-size:0.83rem;color:#a0a8b8;background:rgba(255,255,255,0.04);padding:0.6rem 0.8rem;border-radius:8px;margin-top:0.4rem;">💡 {grade.get("feedback","")}</div>', unsafe_allow_html=True)

                            with st.expander("📖 Sample answer"):
                                st.markdown(q.get("sample_answer",""))
                            st.markdown("</div>", unsafe_allow_html=True)

                # Retake
                st.markdown("<br>", unsafe_allow_html=True)
                rc1, rc2, _ = st.columns([1,1,3])
                with rc1:
                    if st.button("🔄 Retake Exam", use_container_width=True):
                        st.session_state.exam_answers = {}
                        st.session_state.exam_sa_answers = {}
                        st.session_state.exam_submitted = False
                        st.session_state.exam_sa_grades = {}
                        st.session_state.exam_started = False
                        st.rerun()
                with rc2:
                    if st.button("⚡ New Exam", use_container_width=True):
                        st.session_state.exam_data = None
                        st.session_state.exam_answers = {}
                        st.session_state.exam_sa_answers = {}
                        st.session_state.exam_submitted = False
                        st.session_state.exam_sa_grades = {}
                        st.session_state.exam_started = False
                        st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 5: TECHNICAL PROJECTS
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "projects":

    st.markdown("""
    <div style="margin-bottom:1.2rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;
           color:#e8eaf0;margin-bottom:0.3rem;">🛠️ Technical Projects</div>
      <div style="color:#6b7280;font-size:0.9rem;">
        Two staged take-home projects generated from the job description.
        Submit each stage, get GPT-4o code review with scores, corrections, and rubric feedback.
      </div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # ── Project selector tabs ──────────────────────────────────────────────
        proj_tab1, proj_tab2 = st.tabs([
            "📊 Project 1 — BI Dashboard",
            "🗄️ Project 2 — SQL Analytics",
        ])

        # ════════════════════════════════════════════════════════════════════════
        # PROJECT 1 — BI DASHBOARD
        # ════════════════════════════════════════════════════════════════════════
        with proj_tab1:
            using_gpt = bool(st.session_state.openai_key)

            st.markdown(f"""
            <div style="background:rgba(108,99,255,0.07);border:1px solid rgba(108,99,255,0.2);
                 border-radius:10px;padding:0.8rem 1rem;margin-bottom:1rem;font-size:0.83rem;color:#c8cad4;">
              <strong style="color:#a29bfe;">How it works:</strong>
              Four staged deliverables — data model design, DAX measures, dashboard UX, and performance optimization.
              Submit each stage to receive scored feedback with specific corrections.
              Reviewer: {'⚡ GPT-4o' if using_gpt else '🧠 Claude (add OpenAI key for GPT-4o review)'}
            </div>
            """, unsafe_allow_html=True)

            # Generate brief
            if not st.session_state.proj_bi_brief:
                _, gc, _ = st.columns([2,1,2])
                with gc:
                    if st.button("⚡ Generate BI Project Brief", use_container_width=True, key="gen_bi"):
                        with st.spinner("Generating project brief from job description…"):
                            try:
                                st.session_state.proj_bi_brief = generate_bi_brief(
                                    st.session_state.jd_text,
                                    st.session_state.resume_text,
                                )
                                st.session_state.proj_bi_submissions = {}
                                st.session_state.proj_bi_reviews = {}
                                st.session_state.proj_bi_active_stage = 0
                                st.session_state.proj_bi_complete = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Failed to generate brief: {e}")
            else:
                brief = st.session_state.proj_bi_brief
                reviews = st.session_state.proj_bi_reviews
                submissions = st.session_state.proj_bi_submissions

                # ── Brief header ───────────────────────────────────────────────
                st.markdown(f"""
                <div class="card" style="margin-bottom:1.2rem;">
                  <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:0.5rem;">
                    <div>
                      <div style="font-family:Syne,sans-serif;font-size:1.2rem;font-weight:700;
                           color:#e8eaf0;margin-bottom:0.2rem;">{brief.get('title','BI Dashboard Project')}</div>
                      <div style="font-size:0.8rem;color:#6b7280;">{brief.get('estimated_hours','3-4 hours')} estimated
                        · {brief.get('difficulty','Senior')} level</div>
                    </div>
                    <div style="display:flex;gap:0.5rem;align-items:center;flex-wrap:wrap;">
                      {''.join(f'<span class="pill pill-purple">✓ {k}</span>' for k in brief.get('kpis_required',[])[:3])}
                    </div>
                  </div>
                </div>
                """, unsafe_allow_html=True)

                # Brief details in expanders
                with st.expander("📋 Full Project Brief — read before starting"):
                    bc1, bc2 = st.columns(2)
                    with bc1:
                        st.markdown("**Company Context**")
                        st.markdown(brief.get("company_context",""))
                        st.markdown("**Business Problem**")
                        st.markdown(brief.get("business_problem",""))
                        st.markdown("**Required KPIs**")
                        for kpi in brief.get("kpis_required",[]):
                            st.markdown(f"- {kpi}")
                    with bc2:
                        st.markdown("**Data Available**")
                        sample = brief.get("sample_data",{})
                        st.markdown(f"**Fact table:** `{sample.get('fact_table','fact_sales')}`")
                        st.markdown(f"**Columns:** {', '.join(f'`{c}`' for c in sample.get('columns',[]))}")
                        st.markdown(f"**Dimensions:** {', '.join(f'`{d}`' for d in sample.get('dim_tables',[]))}")
                        st.markdown("**Data Model Description**")
                        st.markdown(brief.get("data_model_description",""))
                        st.markdown("**Technical Requirements**")
                        for req in brief.get("technical_requirements",[]):
                            st.markdown(f"- {req}")

                # Overall score banner if all stages done
                if len(reviews) == len(BI_STAGES):
                    overall = overall_project_score(reviews, BI_STAGES)
                    st.markdown(f"""
                    <div style="text-align:center;padding:1.2rem;background:rgba(0,0,0,0.2);
                         border-radius:14px;margin-bottom:1.2rem;border:1px solid {overall['color']}44;">
                      <div style="font-family:Syne,sans-serif;font-size:0.7rem;font-weight:700;
                           letter-spacing:0.15em;text-transform:uppercase;color:#6b7280;">Project Score</div>
                      <div style="font-family:Syne,sans-serif;font-size:3.5rem;font-weight:800;
                           color:{overall['color']};line-height:1;">{overall['overall']}/10</div>
                      <div style="font-family:Syne,sans-serif;font-size:1rem;font-weight:600;
                           color:{overall['color']};">{overall['grade']}</div>
                      <div style="color:#6b7280;font-size:0.8rem;margin-top:0.3rem;">
                        Stages: {' · '.join(str(s) for s in overall['scores'])}
                      </div>
                    </div>
                    """, unsafe_allow_html=True)
                    # Download full review report
                    _, dl_c, _ = st.columns([2,2,2])
                    with dl_c:
                        bi_report = build_project_report_docx(
                            "Power BI Dashboard", brief,
                            BI_STAGES, reviews, submissions,
                        )
                        st.download_button(
                            "⬇️ Download Full Review Report (.docx)",
                            data=bi_report,
                            file_name="bi_project_review.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

                # ── Stage progress bar ─────────────────────────────────────────
                prog_cols = st.columns(len(BI_STAGES))
                for i, (stage, col) in enumerate(zip(BI_STAGES, prog_cols)):
                    sid = stage["id"]
                    done = sid in reviews
                    active = i == st.session_state.proj_bi_active_stage
                    color = "#00d4aa" if done else "#6c63ff" if active else "#1e2130"
                    border = "#00d4aa" if done else "#6c63ff" if active else "#2d3040"
                    icon = "✓" if done else stage["icon"]
                    score_badge = f'<div style="font-size:0.7rem;color:#00d4aa;">{reviews[sid]["score"]}/10</div>' if done else ""
                    with col:
                        st.markdown(f"""
                        <div style="background:{color}22;border:1px solid {border};border-radius:10px;
                             padding:0.6rem;text-align:center;cursor:pointer;">
                          <div style="font-size:1.1rem;">{icon}</div>
                          <div style="font-size:0.7rem;color:#e8eaf0;font-weight:600;margin-top:2px;">
                            Stage {i+1}
                          </div>
                          {score_badge}
                        </div>
                        """, unsafe_allow_html=True)
                        if st.button("Open", key=f"bi_stage_{i}", use_container_width=True):
                            st.session_state.proj_bi_active_stage = i
                            st.rerun()

                st.markdown("<br>", unsafe_allow_html=True)

                # ── Active stage ───────────────────────────────────────────────
                active_idx = st.session_state.proj_bi_active_stage
                if active_idx < len(BI_STAGES):
                    stage = BI_STAGES[active_idx]
                    sid   = stage["id"]

                    st.markdown(f"""
                    <div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;
                         color:#e8eaf0;margin-bottom:0.5rem;">{stage['title']}</div>
                    """, unsafe_allow_html=True)

                    st.markdown(f"""
                    <div class="rec-card" style="margin-bottom:1rem;">
                      <div class="rec-title">Your Task</div>
                      {stage['prompt']}
                    </div>
                    """, unsafe_allow_html=True)

                    # Submission area
                    existing = submissions.get(sid, "")
                    submission = st.text_area(
                        "Your submission",
                        value=existing,
                        height=280,
                        placeholder=stage["placeholder"],
                        key=f"bi_sub_{sid}",
                        label_visibility="collapsed",
                    )

                    sub_c1, sub_c2, _ = st.columns([1,1,3])
                    with sub_c1:
                        submit_btn = st.button("📤 Submit for Review", key=f"bi_submit_{sid}", use_container_width=True)
                    with sub_c2:
                        if active_idx > 0:
                            if st.button("⬅ Previous Stage", key=f"bi_prev_{sid}", use_container_width=True):
                                st.session_state.proj_bi_active_stage = active_idx - 1
                                st.rerun()

                    if submit_btn and submission.strip():
                        st.session_state.proj_bi_submissions[sid] = submission.strip()
                        with st.spinner(f"{'GPT-4o' if using_gpt else 'Claude'} reviewing your submission…"):
                            review = review_project_submission(
                                stage, submission.strip(), brief,
                                "Power BI Dashboard", "N/A",
                                st.session_state.openai_key,
                            )
                            st.session_state.proj_bi_reviews[sid] = review
                        # Auto-advance to next stage
                        if active_idx + 1 < len(BI_STAGES):
                            st.session_state.proj_bi_active_stage = active_idx + 1
                        st.rerun()

                    # Show review if submitted
                    if sid in reviews:
                        rv = reviews[sid]
                        score = rv.get("score", 0)
                        grade = rv.get("grade", "")
                        s_color = "#00d4aa" if score >= 8 else "#ffa940" if score >= 6 else "#ff4d6d"

                        st.markdown(f"""
                        <div style="background:rgba(0,0,0,0.2);border:1px solid {s_color}44;
                             border-radius:12px;padding:1.2rem;margin-top:1rem;">
                          <div style="display:flex;align-items:center;gap:1rem;margin-bottom:0.8rem;">
                            <div>
                              <div style="font-family:Syne,sans-serif;font-size:2.2rem;
                                   font-weight:800;color:{s_color};line-height:1;">{score}/10</div>
                              <div style="font-size:0.8rem;font-weight:600;color:{s_color};">{grade}</div>
                            </div>
                            <div style="font-size:0.88rem;color:#c8cad4;line-height:1.5;">
                              {rv.get('summary','')}
                            </div>
                          </div>
                        """, unsafe_allow_html=True)

                        rv_c1, rv_c2 = st.columns(2)
                        with rv_c1:
                            st.markdown('<div class="section-hdr">✅ Strengths</div>', unsafe_allow_html=True)
                            for s in rv.get("strengths", []):
                                st.markdown(f"""
                                <div class="rec-card" style="border-left-color:#00d4aa;background:rgba(0,212,170,0.05);">
                                  <div class="rec-title" style="color:#00d4aa;">{s.get('point','')}</div>
                                  {s.get('detail','')}
                                </div>""", unsafe_allow_html=True)

                        with rv_c2:
                            st.markdown('<div class="section-hdr">⚠️ Issues to Fix</div>', unsafe_allow_html=True)
                            for issue in rv.get("issues", []):
                                sev = issue.get("severity","Minor")
                                ic = "#ff4d6d" if sev=="Critical" else "#ffa940" if sev=="Major" else "#6b7280"
                                st.markdown(f"""
                                <div class="rec-card" style="border-left-color:{ic};background:{ic}11;">
                                  <div class="rec-title" style="color:{ic};">{sev} · {issue.get('point','')}</div>
                                  <strong>Fix:</strong> {issue.get('fix','')}
                                </div>""", unsafe_allow_html=True)

                        if rv.get("corrected_example","").strip():
                            st.markdown('<div class="section-hdr">🔧 Corrected Example</div>', unsafe_allow_html=True)
                            st.code(rv["corrected_example"], language="dax")

                        missed = rv.get("missed_opportunities",[])
                        if missed:
                            st.markdown('<div class="section-hdr">💡 Missed Opportunities</div>', unsafe_allow_html=True)
                            for m in missed:
                                st.markdown(f'<div class="pill pill-yellow" style="display:inline-flex;margin:2px;">◆ {m}</div>', unsafe_allow_html=True)

                        tip = rv.get("next_stage_tip","")
                        if tip:
                            st.markdown(f"""
                            <div style="background:rgba(108,99,255,0.08);border:1px solid rgba(108,99,255,0.25);
                                 border-radius:8px;padding:0.7rem 0.9rem;margin-top:0.8rem;font-size:0.85rem;color:#c8cad4;">
                              <strong style="color:#a29bfe;">💬 Tip for next stage:</strong> {tip}
                            </div>""", unsafe_allow_html=True)

                        st.markdown("</div>", unsafe_allow_html=True)

                        # Next stage button
                        if active_idx + 1 < len(BI_STAGES):
                            st.markdown("<br>", unsafe_allow_html=True)
                            _, nc, _ = st.columns([2,1,2])
                            with nc:
                                if st.button(f"Next: {BI_STAGES[active_idx+1]['title']} ➡️", key=f"bi_next_{sid}", use_container_width=True):
                                    st.session_state.proj_bi_active_stage = active_idx + 1
                                    st.rerun()

                # Reset project
                st.markdown("---")
                _, reset_c, _ = st.columns([3,1,3])
                with reset_c:
                    if st.button("🔄 New BI Project", key="reset_bi", use_container_width=True):
                        st.session_state.proj_bi_brief = None
                        st.session_state.proj_bi_submissions = {}
                        st.session_state.proj_bi_reviews = {}
                        st.session_state.proj_bi_active_stage = 0
                        st.session_state.proj_bi_complete = False
                        st.rerun()

        # ════════════════════════════════════════════════════════════════════════
        # PROJECT 2 — SQL ANALYTICS
        # ════════════════════════════════════════════════════════════════════════
        with proj_tab2:
            using_gpt = bool(st.session_state.openai_key)

            st.markdown(f"""
            <div style="background:rgba(0,212,170,0.06);border:1px solid rgba(0,212,170,0.2);
                 border-radius:10px;padding:0.8rem 1rem;margin-bottom:1rem;font-size:0.83rem;color:#c8cad4;">
              <strong style="color:#00d4aa;">How it works:</strong>
              Four staged deliverables — data exploration, core analysis query, query optimization, and pipeline design.
              Submit each stage to receive scored review with corrected SQL and specific feedback.
              Reviewer: {'⚡ GPT-4o' if using_gpt else '🧠 Claude (add OpenAI key for GPT-4o review)'}
            </div>
            """, unsafe_allow_html=True)

            # Dialect selector
            d_col, _ = st.columns([1,3])
            with d_col:
                dialect = st.selectbox(
                    "SQL dialect",
                    ["Snowflake", "Databricks (Spark SQL)", "BigQuery", "PostgreSQL", "SQL Server"],
                    index=["Snowflake","Databricks (Spark SQL)","BigQuery","PostgreSQL","SQL Server"].index(
                        st.session_state.proj_sql_dialect
                    ),
                    key="sql_dialect_sel",
                )
                if dialect != st.session_state.proj_sql_dialect:
                    st.session_state.proj_sql_dialect = dialect
                    # Reset if dialect changes
                    if st.session_state.proj_sql_brief:
                        st.session_state.proj_sql_brief = None
                        st.session_state.proj_sql_submissions = {}
                        st.session_state.proj_sql_reviews = {}
                        st.session_state.proj_sql_active_stage = 0

            # Generate brief
            if not st.session_state.proj_sql_brief:
                _, gc2, _ = st.columns([2,1,2])
                with gc2:
                    if st.button(f"⚡ Generate {dialect} Project Brief", use_container_width=True, key="gen_sql"):
                        with st.spinner(f"Generating {dialect} project brief…"):
                            try:
                                st.session_state.proj_sql_brief = generate_sql_brief(
                                    st.session_state.jd_text,
                                    st.session_state.resume_text,
                                    dialect,
                                )
                                st.session_state.proj_sql_submissions = {}
                                st.session_state.proj_sql_reviews = {}
                                st.session_state.proj_sql_active_stage = 0
                                st.session_state.proj_sql_complete = False
                                st.rerun()
                            except Exception as e:
                                st.error(f"Failed to generate brief: {e}")
            else:
                brief2 = st.session_state.proj_sql_brief
                reviews2 = st.session_state.proj_sql_reviews
                submissions2 = st.session_state.proj_sql_submissions

                # Brief header
                st.markdown(f"""
                <div class="card" style="margin-bottom:1.2rem;">
                  <div style="font-family:Syne,sans-serif;font-size:1.2rem;font-weight:700;
                       color:#e8eaf0;margin-bottom:0.3rem;">{brief2.get('title','SQL Analytics Project')}</div>
                  <div style="font-size:0.82rem;color:#c8cad4;margin-bottom:0.4rem;">
                    {brief2.get('central_question','')}
                  </div>
                  <div style="font-size:0.78rem;color:#6b7280;">
                    {brief2.get('estimated_hours','2-3 hours')} · {brief2.get('difficulty','Senior')} ·
                    {brief2.get('dialect', dialect)}
                  </div>
                </div>
                """, unsafe_allow_html=True)

                with st.expander("📋 Full Project Brief — read before starting"):
                    sc1, sc2 = st.columns(2)
                    with sc1:
                        st.markdown("**Company Context**")
                        st.markdown(brief2.get("company_context",""))
                        st.markdown("**Business Problem**")
                        st.markdown(brief2.get("business_problem",""))
                        st.markdown("**Sub-Questions to Answer**")
                        for q in brief2.get("sub_questions",[]):
                            st.markdown(f"- {q}")
                    with sc2:
                        st.markdown("**Schema**")
                        for tbl in brief2.get("schema",{}).get("tables",[]):
                            st.markdown(f"**`{tbl['name']}`** — {tbl.get('description','')} ({tbl.get('row_count','')})")
                            cols_str = ", ".join(f"`{c['name']}` ({c['type']})" for c in tbl.get("columns",[]))
                            st.markdown(f"  {cols_str}")
                            if tbl.get("clustering_key"):
                                st.markdown(f"  Cluster key: `{tbl['clustering_key']}`")
                        st.markdown("**Known Data Issues**")
                        for issue in brief2.get("known_data_issues",[]):
                            st.markdown(f"⚠️ {issue}")
                        st.markdown(f"**Performance Constraint:** {brief2.get('performance_constraints','')}")

                # Overall score
                if len(reviews2) == len(SQL_STAGES):
                    overall2 = overall_project_score(reviews2, SQL_STAGES)
                    st.markdown(f"""
                    <div style="text-align:center;padding:1.2rem;background:rgba(0,0,0,0.2);
                         border-radius:14px;margin-bottom:1.2rem;border:1px solid {overall2['color']}44;">
                      <div style="font-family:Syne,sans-serif;font-size:0.7rem;font-weight:700;
                           letter-spacing:0.15em;text-transform:uppercase;color:#6b7280;">Project Score</div>
                      <div style="font-family:Syne,sans-serif;font-size:3.5rem;font-weight:800;
                           color:{overall2['color']};line-height:1;">{overall2['overall']}/10</div>
                      <div style="font-family:Syne,sans-serif;font-size:1rem;font-weight:600;
                           color:{overall2['color']};">{overall2['grade']}</div>
                    </div>
                    """, unsafe_allow_html=True)
                    _, dl_c2, _ = st.columns([2,2,2])
                    with dl_c2:
                        sql_report = build_project_report_docx(
                            f"{dialect} SQL Analytics", brief2,
                            SQL_STAGES, reviews2, submissions2,
                        )
                        st.download_button(
                            "⬇️ Download Full Review Report (.docx)",
                            data=sql_report,
                            file_name="sql_project_review.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

                # Stage progress
                prog_cols2 = st.columns(len(SQL_STAGES))
                for i, (stage, col) in enumerate(zip(SQL_STAGES, prog_cols2)):
                    sid = stage["id"]
                    done = sid in reviews2
                    active = i == st.session_state.proj_sql_active_stage
                    color = "#00d4aa" if done else "#00d4aa" if active else "#1e2130"
                    border = "#00d4aa" if done else "#00d4aa44" if active else "#2d3040"
                    icon = "✓" if done else stage["icon"]
                    score_badge = f'<div style="font-size:0.7rem;color:#00d4aa;">{reviews2[sid]["score"]}/10</div>' if done else ""
                    with col:
                        st.markdown(f"""
                        <div style="background:{color}22;border:1px solid {border};border-radius:10px;
                             padding:0.6rem;text-align:center;">
                          <div style="font-size:1.1rem;">{icon}</div>
                          <div style="font-size:0.7rem;color:#e8eaf0;font-weight:600;margin-top:2px;">Stage {i+1}</div>
                          {score_badge}
                        </div>
                        """, unsafe_allow_html=True)
                        if st.button("Open", key=f"sql_stage_{i}", use_container_width=True):
                            st.session_state.proj_sql_active_stage = i
                            st.rerun()

                st.markdown("<br>", unsafe_allow_html=True)

                # Active stage
                active_idx2 = st.session_state.proj_sql_active_stage
                if active_idx2 < len(SQL_STAGES):
                    stage2 = SQL_STAGES[active_idx2]
                    sid2   = stage2["id"]

                    st.markdown(f"""
                    <div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;
                         color:#e8eaf0;margin-bottom:0.5rem;">{stage2['title']}</div>
                    """, unsafe_allow_html=True)

                    st.markdown(f"""
                    <div class="rec-card" style="border-left-color:#00d4aa;margin-bottom:1rem;">
                      <div class="rec-title" style="color:#00d4aa;">Your Task</div>
                      {stage2['prompt']}
                    </div>
                    """, unsafe_allow_html=True)

                    existing2 = submissions2.get(sid2,"")
                    submission2 = st.text_area(
                        "Your SQL",
                        value=existing2,
                        height=320,
                        placeholder=stage2["placeholder"],
                        key=f"sql_sub_{sid2}",
                        label_visibility="collapsed",
                    )

                    s2c1, s2c2, _ = st.columns([1,1,3])
                    with s2c1:
                        submit_btn2 = st.button("📤 Submit for Review", key=f"sql_submit_{sid2}", use_container_width=True)
                    with s2c2:
                        if active_idx2 > 0:
                            if st.button("⬅ Previous Stage", key=f"sql_prev_{sid2}", use_container_width=True):
                                st.session_state.proj_sql_active_stage = active_idx2 - 1
                                st.rerun()

                    if submit_btn2 and submission2.strip():
                        st.session_state.proj_sql_submissions[sid2] = submission2.strip()
                        with st.spinner(f"{'GPT-4o' if using_gpt else 'Claude'} reviewing your SQL…"):
                            review2 = review_project_submission(
                                stage2, submission2.strip(), brief2,
                                f"{dialect} SQL Analytics", dialect,
                                st.session_state.openai_key,
                            )
                            st.session_state.proj_sql_reviews[sid2] = review2
                        if active_idx2 + 1 < len(SQL_STAGES):
                            st.session_state.proj_sql_active_stage = active_idx2 + 1
                        st.rerun()

                    # Show review
                    if sid2 in reviews2:
                        rv2 = reviews2[sid2]
                        score2 = rv2.get("score",0)
                        grade2 = rv2.get("grade","")
                        sc2_color = "#00d4aa" if score2 >= 8 else "#ffa940" if score2 >= 6 else "#ff4d6d"

                        st.markdown(f"""
                        <div style="background:rgba(0,0,0,0.2);border:1px solid {sc2_color}44;
                             border-radius:12px;padding:1.2rem;margin-top:1rem;">
                          <div style="display:flex;align-items:center;gap:1rem;margin-bottom:0.8rem;">
                            <div>
                              <div style="font-family:Syne,sans-serif;font-size:2.2rem;
                                   font-weight:800;color:{sc2_color};line-height:1;">{score2}/10</div>
                              <div style="font-size:0.8rem;font-weight:600;color:{sc2_color};">{grade2}</div>
                            </div>
                            <div style="font-size:0.88rem;color:#c8cad4;line-height:1.5;">
                              {rv2.get('summary','')}
                            </div>
                          </div>
                        """, unsafe_allow_html=True)

                        rv2_c1, rv2_c2 = st.columns(2)
                        with rv2_c1:
                            st.markdown('<div class="section-hdr">✅ Strengths</div>', unsafe_allow_html=True)
                            for s in rv2.get("strengths",[]):
                                st.markdown(f"""
                                <div class="rec-card" style="border-left-color:#00d4aa;background:rgba(0,212,170,0.05);">
                                  <div class="rec-title" style="color:#00d4aa;">{s.get('point','')}</div>
                                  {s.get('detail','')}
                                </div>""", unsafe_allow_html=True)
                        with rv2_c2:
                            st.markdown('<div class="section-hdr">⚠️ Issues</div>', unsafe_allow_html=True)
                            for issue in rv2.get("issues",[]):
                                sev = issue.get("severity","Minor")
                                ic = "#ff4d6d" if sev=="Critical" else "#ffa940" if sev=="Major" else "#6b7280"
                                st.markdown(f"""
                                <div class="rec-card" style="border-left-color:{ic};background:{ic}11;">
                                  <div class="rec-title" style="color:{ic};">{sev} · {issue.get('point','')}</div>
                                  <strong>Fix:</strong> {issue.get('fix','')}
                                </div>""", unsafe_allow_html=True)

                        if rv2.get("corrected_example","").strip():
                            st.markdown('<div class="section-hdr">🔧 Corrected SQL</div>', unsafe_allow_html=True)
                            st.code(rv2["corrected_example"], language="sql")

                        missed2 = rv2.get("missed_opportunities",[])
                        if missed2:
                            st.markdown('<div class="section-hdr">💡 Missed Opportunities</div>', unsafe_allow_html=True)
                            for m in missed2:
                                st.markdown(f'<div class="pill pill-yellow" style="display:inline-flex;margin:2px;">◆ {m}</div>', unsafe_allow_html=True)

                        tip2 = rv2.get("next_stage_tip","")
                        if tip2:
                            st.markdown(f"""
                            <div style="background:rgba(0,212,170,0.07);border:1px solid rgba(0,212,170,0.25);
                                 border-radius:8px;padding:0.7rem 0.9rem;margin-top:0.8rem;font-size:0.85rem;color:#c8cad4;">
                              <strong style="color:#00d4aa;">💬 Tip for next stage:</strong> {tip2}
                            </div>""", unsafe_allow_html=True)

                        st.markdown("</div>", unsafe_allow_html=True)

                        if active_idx2 + 1 < len(SQL_STAGES):
                            st.markdown("<br>", unsafe_allow_html=True)
                            _, nc2, _ = st.columns([2,1,2])
                            with nc2:
                                if st.button(f"Next: {SQL_STAGES[active_idx2+1]['title']} ➡️", key=f"sql_next_{sid2}", use_container_width=True):
                                    st.session_state.proj_sql_active_stage = active_idx2 + 1
                                    st.rerun()

                # Reset
                st.markdown("---")
                _, reset_c2, _ = st.columns([3,1,3])
                with reset_c2:
                    if st.button("🔄 New SQL Project", key="reset_sql", use_container_width=True):
                        st.session_state.proj_sql_brief = None
                        st.session_state.proj_sql_submissions = {}
                        st.session_state.proj_sql_reviews = {}
                        st.session_state.proj_sql_active_stage = 0
                        st.rerun()
