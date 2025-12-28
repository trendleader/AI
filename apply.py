"""
Resume Tailoring Agent
A Streamlit application that analyzes job postings, matches keywords,
and generates ATS-optimized, professionally formatted resumes.
"""

import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import json
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PyPDF2
import io
from collections import Counter
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Page configuration
st.set_page_config(
    page_title="Resume Tailoring Agent",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1E3A5F;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.1rem;
        color: #666;
        text-align: center;
        margin-bottom: 2rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1.5rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .metric-value {
        font-size: 2.5rem;
        font-weight: 700;
    }
    .metric-label {
        font-size: 0.9rem;
        opacity: 0.9;
    }
    .keyword-badge {
        display: inline-block;
        padding: 0.3rem 0.8rem;
        margin: 0.2rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 500;
    }
    .keyword-matched {
        background-color: #d4edda;
        color: #155724;
        border: 1px solid #c3e6cb;
    }
    .keyword-missing {
        background-color: #f8d7da;
        color: #721c24;
        border: 1px solid #f5c6cb;
    }
    .keyword-added {
        background-color: #cce5ff;
        color: #004085;
        border: 1px solid #b8daff;
    }
    .section-header {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1E3A5F;
        border-bottom: 2px solid #667eea;
        padding-bottom: 0.5rem;
        margin: 1.5rem 0 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
    }
</style>
""", unsafe_allow_html=True)


# Initialize session state
if 'resume_text' not in st.session_state:
    st.session_state.resume_text = None
if 'job_data' not in st.session_state:
    st.session_state.job_data = None
if 'tailored_resume' not in st.session_state:
    st.session_state.tailored_resume = None
if 'keyword_analysis' not in st.session_state:
    st.session_state.keyword_analysis = None
if 'cover_letter' not in st.session_state:
    st.session_state.cover_letter = None
if 'interview_prep' not in st.session_state:
    st.session_state.interview_prep = None


@st.cache_resource
def load_nlp_model():
    """Load spaCy model for NLP processing with fallback."""
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        try:
            import subprocess
            subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"], 
                         capture_output=True, timeout=60)
            return spacy.load("en_core_web_sm")
        except:
            # Return None to trigger fallback keyword extraction
            return None


def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file."""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text.strip()


def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX file."""
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text.strip()


def scrape_job_posting(url):
    """Scrape job posting from URL with multiple fallback strategies."""
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
        'Accept-Language': 'en-US,en;q=0.5',
        'Accept-Encoding': 'gzip, deflate, br',
        'Connection': 'keep-alive',
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            element.decompose()
        
        # Try to find job-specific content containers
        job_selectors = [
            {'class_': re.compile(r'job[-_]?(description|details|content|posting|body)', re.I)},
            {'class_': re.compile(r'description[-_]?(content|body|text)', re.I)},
            {'id': re.compile(r'job[-_]?(description|details|content)', re.I)},
            {'class_': 'content'},
            {'role': 'main'},
            {'itemprop': 'description'},
        ]
        
        job_content = None
        for selector in job_selectors:
            job_content = soup.find(['div', 'section', 'article', 'main'], selector)
            if job_content:
                break
        
        if not job_content:
            job_content = soup.find('body')
        
        # Extract title
        title = None
        title_selectors = [
            soup.find('h1'),
            soup.find(class_=re.compile(r'job[-_]?title', re.I)),
            soup.find('title'),
        ]
        for t in title_selectors:
            if t:
                title = t.get_text(strip=True)
                break
        
        # Get clean text
        text = job_content.get_text(separator='\n', strip=True) if job_content else ""
        
        # Clean up the text
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        text = '\n'.join(lines)
        
        # Limit text length
        if len(text) > 15000:
            text = text[:15000]
        
        return {
            'title': title or 'Job Posting',
            'content': text,
            'url': url
        }
        
    except Exception as e:
        return {'error': str(e)}


def extract_keywords(text, nlp, top_n=50):
    """Extract important keywords and phrases from text using NLP."""
    
    # Common tech and business keywords to prioritize
    tech_keywords = {
        'python', 'sql', 'java', 'javascript', 'react', 'node', 'aws', 'azure', 
        'gcp', 'docker', 'kubernetes', 'api', 'rest', 'graphql', 'machine learning',
        'data science', 'analytics', 'tableau', 'power bi', 'excel', 'agile',
        'scrum', 'project management', 'leadership', 'communication', 'collaboration',
        'strategic', 'analysis', 'reporting', 'dashboard', 'etl', 'data warehouse',
        'cloud', 'devops', 'ci/cd', 'git', 'database', 'mongodb', 'postgresql',
        'tensorflow', 'pytorch', 'nlp', 'deep learning', 'statistics', 'r',
        'spark', 'hadoop', 'airflow', 'kafka', 'redis', 'elasticsearch',
        'salesforce', 'sap', 'oracle', 'jira', 'confluence', 'slack',
        'microsoft', 'google', 'linux', 'windows', 'macos', 'ios', 'android',
        'html', 'css', 'typescript', 'vue', 'angular', 'django', 'flask',
        'numpy', 'pandas', 'scikit-learn', 'matplotlib', 'seaborn',
        'stakeholder', 'requirements', 'documentation', 'testing', 'qa',
        'budget', 'revenue', 'profit', 'growth', 'optimization', 'efficiency',
        'team', 'cross-functional', 'mentoring', 'training', 'presentation'
    }
    
    # Fallback extraction without spaCy
    if nlp is None:
        text_lower = text.lower()
        
        # Simple tokenization
        words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9+#.-]*[a-zA-Z0-9]\b|\b[a-zA-Z]\b', text_lower)
        
        # Filter stopwords
        stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                    'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
                    'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                    'should', 'may', 'might', 'must', 'shall', 'can', 'need', 'dare', 'ought',
                    'used', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it',
                    'we', 'they', 'what', 'which', 'who', 'whom', 'whose', 'where', 'when',
                    'why', 'how', 'all', 'each', 'every', 'both', 'few', 'more', 'most',
                    'other', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same',
                    'so', 'than', 'too', 'very', 'just', 'also', 'now', 'here', 'there',
                    'then', 'once', 'if', 'unless', 'until', 'while', 'about', 'above',
                    'after', 'again', 'against', 'before', 'below', 'between', 'into',
                    'through', 'during', 'over', 'under', 'further', 'then', 'once'}
        
        filtered_words = [w for w in words if w not in stopwords and len(w) > 2]
        
        # Extract bigrams
        bigrams = []
        for i in range(len(filtered_words) - 1):
            bigram = f"{filtered_words[i]} {filtered_words[i+1]}"
            bigrams.append(bigram)
        
        # Count frequencies
        phrase_counts = Counter(filtered_words + bigrams)
        
        # Boost tech keywords
        boosted_counts = {}
        for phrase, count in phrase_counts.items():
            if any(kw in phrase for kw in tech_keywords):
                boosted_counts[phrase] = count * 3
            else:
                boosted_counts[phrase] = count
        
        top_keywords = sorted(boosted_counts.items(), key=lambda x: x[1], reverse=True)[:top_n]
        return [kw[0] for kw in top_keywords]
    
    # spaCy-based extraction
    doc = nlp(text.lower())
    
    # Extract noun phrases and named entities
    phrases = []
    
    # Get noun chunks
    for chunk in doc.noun_chunks:
        phrase = chunk.text.strip()
        if len(phrase) > 2 and not phrase.isdigit():
            phrases.append(phrase)
    
    # Get named entities
    for ent in doc.ents:
        if ent.label_ in ['ORG', 'PRODUCT', 'SKILL', 'WORK_OF_ART']:
            phrases.append(ent.text.lower())
    
    # Also get individual important tokens
    important_pos = {'NOUN', 'PROPN', 'ADJ'}
    for token in doc:
        if (token.pos_ in important_pos and 
            not token.is_stop and 
            len(token.text) > 2 and 
            token.is_alpha):
            phrases.append(token.text)
    
    # Count frequencies
    phrase_counts = Counter(phrases)
    
    # Boost scores for tech keywords
    boosted_counts = {}
    for phrase, count in phrase_counts.items():
        if any(kw in phrase for kw in tech_keywords):
            boosted_counts[phrase] = count * 3
        else:
            boosted_counts[phrase] = count
    
    # Get top keywords
    top_keywords = sorted(boosted_counts.items(), key=lambda x: x[1], reverse=True)[:top_n]
    
    return [kw[0] for kw in top_keywords]


def analyze_keyword_match(resume_keywords, job_keywords):
    """Analyze keyword match between resume and job posting."""
    resume_set = set(kw.lower() for kw in resume_keywords)
    job_set = set(kw.lower() for kw in job_keywords)
    
    matched = resume_set.intersection(job_set)
    missing = job_set - resume_set
    
    # Calculate match score
    match_score = len(matched) / len(job_set) * 100 if job_set else 0
    
    return {
        'matched': list(matched),
        'missing': list(missing),
        'match_score': match_score,
        'resume_keywords': list(resume_set),
        'job_keywords': list(job_set)
    }


def calculate_similarity_score(resume_text, job_text):
    """Calculate TF-IDF cosine similarity between resume and job posting."""
    vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
    try:
        tfidf_matrix = vectorizer.fit_transform([resume_text, job_text])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return round(similarity * 100, 1)
    except:
        return 0


def tailor_resume_with_ai(resume_text, job_data, keyword_analysis, api_key):
    """Use OpenAI to tailor resume content for the job posting."""
    client = OpenAI(api_key=api_key)
    
    missing_keywords = keyword_analysis['missing'][:20]
    
    prompt = f"""You are an expert resume writer and ATS optimization specialist. Your task is to tailor the provided resume to match the job posting while maintaining authenticity and truthfulness.

JOB POSTING:
Title: {job_data.get('title', 'Position')}
Description:
{job_data.get('content', '')[:4000]}

KEYWORDS TO INCORPORATE (currently missing from resume):
{', '.join(missing_keywords)}

ORIGINAL RESUME:
{resume_text}

INSTRUCTIONS:
1. Rewrite the professional summary to directly address the job requirements
2. Reorder and rewrite bullet points to emphasize relevant experience
3. Naturally incorporate the missing keywords where truthful and appropriate
4. Use strong action verbs and quantifiable achievements
5. Optimize for ATS systems while maintaining readability
6. Keep the same overall structure but enhance relevance
7. Do NOT fabricate experience or skills the candidate doesn't have
8. Ensure each bullet point demonstrates impact with metrics where possible

Return the tailored resume in the following JSON structure:
{{
    "contact": {{
        "name": "Full Name",
        "email": "email@example.com",
        "phone": "phone number",
        "linkedin": "LinkedIn URL (optional)",
        "location": "City, State"
    }},
    "summary": "2-3 sentence professional summary tailored to this specific role",
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "location": "City, State",
            "dates": "Start - End",
            "bullets": [
                "Achievement-focused bullet point with metrics",
                "Another bullet point"
            ]
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "school": "School Name",
            "location": "City, State",
            "date": "Graduation Year",
            "details": "Honors, GPA if notable, relevant coursework (optional)"
        }}
    ],
    "skills": {{
        "technical": ["Skill 1", "Skill 2"],
        "tools": ["Tool 1", "Tool 2"],
        "certifications": ["Cert 1", "Cert 2"]
    }},
    "keywords_added": ["list of missing keywords successfully incorporated"]
}}

Ensure all content is truthful based on the original resume. Only incorporate keywords where they genuinely apply to the candidate's experience."""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert ATS resume optimizer. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Clean up JSON if wrapped in markdown
        if response_text.startswith('```'):
            response_text = re.sub(r'^```(?:json)?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)
        
        return json.loads(response_text)
        
    except json.JSONDecodeError as e:
        st.error(f"Error parsing AI response: {e}")
        return None
    except Exception as e:
        st.error(f"Error calling OpenAI API: {e}")
        return None


def generate_cover_letter(resume_data, job_data, api_key, tone="professional"):
    """Generate a tailored cover letter using AI."""
    client = OpenAI(api_key=api_key)
    
    tone_instructions = {
        "professional": "Write in a formal, professional tone suitable for corporate environments.",
        "conversational": "Write in a warm, conversational yet professional tone that shows personality.",
        "confident": "Write with confident, assertive language that showcases accomplishments boldly.",
        "enthusiastic": "Write with genuine enthusiasm and passion for the role and company."
    }
    
    prompt = f"""You are an expert cover letter writer. Create a compelling, personalized cover letter for this job application.

JOB POSTING:
Title: {job_data.get('title', 'Position')}
Company: Extract from job posting if available
Description:
{job_data.get('content', '')[:3000]}

CANDIDATE RESUME DATA:
{json.dumps(resume_data, indent=2)}

TONE: {tone_instructions.get(tone, tone_instructions['professional'])}

INSTRUCTIONS:
1. Start with a strong opening hook that shows genuine interest in the specific role
2. Highlight 2-3 most relevant achievements that directly address job requirements
3. Demonstrate knowledge of the company/industry if details are available
4. Show how the candidate's experience solves the employer's problems
5. Include specific metrics and accomplishments from the resume
6. End with a confident call to action
7. Keep it to 3-4 paragraphs, approximately 300-400 words
8. Do NOT use generic phrases like "I am writing to apply" or "I believe I would be a great fit"
9. Make it specific to THIS role and THIS candidate's experience

Return the cover letter in JSON format:
{{
    "opening_paragraph": "Strong opening that hooks the reader...",
    "body_paragraph_1": "First body paragraph highlighting key achievement...",
    "body_paragraph_2": "Second body paragraph with another relevant experience...",
    "closing_paragraph": "Confident closing with call to action...",
    "full_letter": "The complete cover letter as a single formatted text block"
}}"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert cover letter writer. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8,
            max_tokens=2000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        if response_text.startswith('```'):
            response_text = re.sub(r'^```(?:json)?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)
        
        return json.loads(response_text)
        
    except Exception as e:
        st.error(f"Error generating cover letter: {e}")
        return None


def generate_interview_prep(resume_data, job_data, api_key):
    """Generate behavioral questions, technical questions, and coding challenges."""
    client = OpenAI(api_key=api_key)
    
    prompt = f"""You are an expert interview coach and technical recruiter. Create a comprehensive interview preparation guide based on this job posting and candidate background.

JOB POSTING:
Title: {job_data.get('title', 'Position')}
Description:
{job_data.get('content', '')[:4000]}

CANDIDATE BACKGROUND:
{json.dumps(resume_data, indent=2) if isinstance(resume_data, dict) else resume_data[:2000]}

Create interview preparation materials in the following JSON structure:

{{
    "job_analysis": {{
        "key_responsibilities": ["Main responsibility 1", "Main responsibility 2"],
        "required_skills": ["Skill 1", "Skill 2"],
        "culture_indicators": ["What the job posting reveals about company culture"],
        "red_flags_to_address": ["Any gaps or concerns the candidate should prepare to address"]
    }},
    "behavioral_questions": [
        {{
            "question": "Tell me about a time when...",
            "why_asked": "What the interviewer is looking for",
            "star_framework_tips": "Specific tips for answering using STAR method",
            "sample_answer_outline": "Key points to hit based on candidate's experience"
        }}
    ],
    "technical_questions": [
        {{
            "question": "Technical question text",
            "difficulty": "Easy/Medium/Hard",
            "topic": "The technical area being tested",
            "key_points": ["Point 1 to cover", "Point 2 to cover"],
            "sample_answer": "Concise model answer"
        }}
    ],
    "coding_challenges": [
        {{
            "title": "Challenge title",
            "difficulty": "Easy/Medium/Hard",
            "description": "Full problem description",
            "example_input": "Sample input",
            "example_output": "Expected output",
            "hints": ["Hint 1", "Hint 2"],
            "solution_approach": "Optimal approach explanation",
            "python_solution": "Complete Python solution code",
            "time_complexity": "O(n) etc",
            "space_complexity": "O(1) etc"
        }}
    ],
    "system_design_questions": [
        {{
            "question": "Design a system that...",
            "key_components": ["Component 1", "Component 2"],
            "discussion_points": ["Point to discuss"],
            "follow_up_questions": ["Possible follow-up"]
        }}
    ],
    "questions_to_ask": [
        {{
            "question": "Question to ask interviewer",
            "why_effective": "Why this question shows engagement/insight"
        }}
    ],
    "salary_negotiation": {{
        "market_range": "Estimated salary range based on role",
        "negotiation_tips": ["Tip 1", "Tip 2"],
        "value_propositions": ["How to justify higher compensation based on experience"]
    }}
}}

IMPORTANT:
- Generate 8-10 behavioral questions covering leadership, conflict, failure, success, teamwork
- Generate 8-10 technical questions specific to the technologies mentioned in the job posting
- Generate 4-5 coding challenges ranging from Easy to Hard, relevant to the role
- Include 2-3 system design questions if the role is senior/technical
- Provide 5-7 smart questions for the candidate to ask
- Tailor everything to the SPECIFIC job posting and candidate background"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert interview coach. Return only valid JSON with comprehensive interview prep materials."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=6000
        )
        
        response_text = response.choices[0].message.content.strip()
        
        if response_text.startswith('```'):
            response_text = re.sub(r'^```(?:json)?\n?', '', response_text)
            response_text = re.sub(r'\n?```$', '', response_text)
        
        return json.loads(response_text)
        
    except Exception as e:
        st.error(f"Error generating interview prep: {e}")
        return None


def create_cover_letter_docx(cover_letter_data, contact_info, job_title=""):
    """Create a professionally formatted cover letter Word document."""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    primary_color = RGBColor(30, 58, 95)
    text_color = RGBColor(51, 51, 51)
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Header with contact info
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(contact_info.get('name', 'Your Name'))
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name_run.font.color.rgb = primary_color
    name_run.font.name = 'Calibri'
    
    # Contact line
    contact_parts = []
    if contact_info.get('email'):
        contact_parts.append(contact_info['email'])
    if contact_info.get('phone'):
        contact_parts.append(contact_info['phone'])
    if contact_info.get('location'):
        contact_parts.append(contact_info['location'])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = text_color
        contact_run.font.name = 'Calibri'
    
    # Date
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.space_before = Pt(24)
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(11)
    date_run.font.name = 'Calibri'
    
    # Hiring Manager line
    hiring_para = doc.add_paragraph()
    hiring_para.space_before = Pt(12)
    hiring_run = hiring_para.add_run("Dear Hiring Manager,")
    hiring_run.font.size = Pt(11)
    hiring_run.font.name = 'Calibri'
    
    # Letter body
    letter_text = cover_letter_data.get('full_letter', '')
    paragraphs = letter_text.split('\n\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph()
            para.space_before = Pt(12)
            run = para.add_run(para_text.strip())
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.font.color.rgb = text_color
    
    # Closing
    closing_para = doc.add_paragraph()
    closing_para.space_before = Pt(24)
    closing_run = closing_para.add_run("Sincerely,")
    closing_run.font.size = Pt(11)
    closing_run.font.name = 'Calibri'
    
    # Signature
    sig_para = doc.add_paragraph()
    sig_para.space_before = Pt(36)
    sig_run = sig_para.add_run(contact_info.get('name', 'Your Name'))
    sig_run.font.size = Pt(11)
    sig_run.font.name = 'Calibri'
    sig_run.font.bold = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_interview_prep_docx(interview_data, job_title=""):
    """Create a comprehensive interview preparation document."""
    doc = Document()
    
    primary_color = RGBColor(30, 58, 95)
    accent_color = RGBColor(102, 126, 234)
    text_color = RGBColor(51, 51, 51)
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    def add_title(text):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = primary_color
        run.font.name = 'Calibri'
        para.space_after = Pt(6)
    
    def add_section_header(text):
        para = doc.add_paragraph()
        para.space_before = Pt(18)
        para.space_after = Pt(8)
        run = para.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = accent_color
        run.font.name = 'Calibri'
        
        # Add bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '667EEA')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def add_subsection(text):
        para = doc.add_paragraph()
        para.space_before = Pt(12)
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = text_color
        run.font.name = 'Calibri'
    
    def add_body_text(text, indent=False):
        para = doc.add_paragraph()
        if indent:
            para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = text_color
        para.space_after = Pt(4)
    
    def add_bullet(text):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        run = para.add_run("• " + text)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        para.space_after = Pt(2)
    
    # Title
    add_title(f"Interview Preparation Guide")
    if job_title:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(job_title)
        sub_run.font.size = Pt(14)
        sub_run.font.italic = True
        sub_run.font.name = 'Calibri'
    
    # Job Analysis
    if interview_data.get('job_analysis'):
        add_section_header("📋 Job Analysis")
        analysis = interview_data['job_analysis']
        
        if analysis.get('key_responsibilities'):
            add_subsection("Key Responsibilities:")
            for resp in analysis['key_responsibilities']:
                add_bullet(resp)
        
        if analysis.get('required_skills'):
            add_subsection("Required Skills:")
            for skill in analysis['required_skills']:
                add_bullet(skill)
        
        if analysis.get('red_flags_to_address'):
            add_subsection("Areas to Address:")
            for flag in analysis['red_flags_to_address']:
                add_bullet(flag)
    
    # Behavioral Questions
    if interview_data.get('behavioral_questions'):
        add_section_header("🗣️ Behavioral Questions")
        for i, q in enumerate(interview_data['behavioral_questions'], 1):
            add_subsection(f"Q{i}: {q.get('question', '')}")
            if q.get('why_asked'):
                add_body_text(f"Why Asked: {q['why_asked']}", indent=True)
            if q.get('star_framework_tips'):
                add_body_text(f"STAR Tips: {q['star_framework_tips']}", indent=True)
            if q.get('sample_answer_outline'):
                add_body_text(f"Answer Outline: {q['sample_answer_outline']}", indent=True)
    
    # Technical Questions
    if interview_data.get('technical_questions'):
        add_section_header("💻 Technical Questions")
        for i, q in enumerate(interview_data['technical_questions'], 1):
            difficulty = q.get('difficulty', '')
            topic = q.get('topic', '')
            add_subsection(f"Q{i}: {q.get('question', '')} [{difficulty}]")
            if topic:
                add_body_text(f"Topic: {topic}", indent=True)
            if q.get('key_points'):
                add_body_text("Key Points:", indent=True)
                for point in q['key_points']:
                    add_bullet(point)
            if q.get('sample_answer'):
                add_body_text(f"Answer: {q['sample_answer']}", indent=True)
    
    # Coding Challenges
    if interview_data.get('coding_challenges'):
        add_section_header("🧩 Coding Challenges")
        for i, challenge in enumerate(interview_data['coding_challenges'], 1):
            add_subsection(f"Challenge {i}: {challenge.get('title', '')} [{challenge.get('difficulty', '')}]")
            add_body_text(f"Problem: {challenge.get('description', '')}", indent=True)
            
            if challenge.get('example_input'):
                add_body_text(f"Example Input: {challenge['example_input']}", indent=True)
            if challenge.get('example_output'):
                add_body_text(f"Example Output: {challenge['example_output']}", indent=True)
            
            if challenge.get('hints'):
                add_body_text("Hints:", indent=True)
                for hint in challenge['hints']:
                    add_bullet(hint)
            
            if challenge.get('solution_approach'):
                add_body_text(f"Approach: {challenge['solution_approach']}", indent=True)
            
            if challenge.get('python_solution'):
                add_body_text("Solution:", indent=True)
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_run = code_para.add_run(challenge['python_solution'])
                code_run.font.size = Pt(9)
                code_run.font.name = 'Consolas'
            
            if challenge.get('time_complexity') or challenge.get('space_complexity'):
                complexity = f"Time: {challenge.get('time_complexity', 'N/A')} | Space: {challenge.get('space_complexity', 'N/A')}"
                add_body_text(f"Complexity: {complexity}", indent=True)
    
    # System Design
    if interview_data.get('system_design_questions'):
        add_section_header("🏗️ System Design Questions")
        for i, q in enumerate(interview_data['system_design_questions'], 1):
            add_subsection(f"Q{i}: {q.get('question', '')}")
            if q.get('key_components'):
                add_body_text("Key Components:", indent=True)
                for comp in q['key_components']:
                    add_bullet(comp)
            if q.get('discussion_points'):
                add_body_text("Discussion Points:", indent=True)
                for point in q['discussion_points']:
                    add_bullet(point)
    
    # Questions to Ask
    if interview_data.get('questions_to_ask'):
        add_section_header("❓ Questions to Ask the Interviewer")
        for q in interview_data['questions_to_ask']:
            add_subsection(q.get('question', ''))
            if q.get('why_effective'):
                add_body_text(f"Why Effective: {q['why_effective']}", indent=True)
    
    # Salary Negotiation
    if interview_data.get('salary_negotiation'):
        add_section_header("💰 Salary Negotiation")
        salary = interview_data['salary_negotiation']
        if salary.get('market_range'):
            add_body_text(f"Market Range: {salary['market_range']}")
        if salary.get('negotiation_tips'):
            add_subsection("Tips:")
            for tip in salary['negotiation_tips']:
                add_bullet(tip)
        if salary.get('value_propositions'):
            add_subsection("Your Value Propositions:")
            for prop in salary['value_propositions']:
                add_bullet(prop)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_ats_resume_docx(resume_data, job_title=""):
    """Create a beautifully formatted, ATS-compliant Word document."""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Define colors
    primary_color = RGBColor(30, 58, 95)  # Dark blue
    secondary_color = RGBColor(102, 126, 234)  # Purple-blue
    text_color = RGBColor(51, 51, 51)  # Dark gray
    
    # Configure Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = text_color
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    contact = resume_data.get('contact', {})
    
    # Name - Large and prominent
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(contact.get('name', 'Your Name').upper())
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = primary_color
    name_run.font.name = 'Calibri'
    name_para.space_after = Pt(4)
    
    # Contact info line
    contact_parts = []
    if contact.get('location'):
        contact_parts.append(contact['location'])
    if contact.get('phone'):
        contact_parts.append(contact['phone'])
    if contact.get('email'):
        contact_parts.append(contact['email'])
    if contact.get('linkedin'):
        linkedin = contact['linkedin'].replace('https://', '').replace('http://', '')
        contact_parts.append(linkedin)
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run('  |  '.join(contact_parts))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = text_color
        contact_run.font.name = 'Calibri'
        contact_para.space_after = Pt(12)
    
    def add_section_header(title):
        """Add a styled section header with underline."""
        para = doc.add_paragraph()
        para.space_before = Pt(12)
        para.space_after = Pt(6)
        run = para.add_run(title.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = primary_color
        run.font.name = 'Calibri'
        
        # Add bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '667EEA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        return para
    
    def add_bullet_point(text, indent_level=0):
        """Add a bullet point with proper formatting."""
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.space_before = Pt(0)
        
        # Clear default content and add formatted text
        para.clear()
        run = para.add_run('• ' + text)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = text_color
        
        return para
    
    # Professional Summary
    if resume_data.get('summary'):
        add_section_header('Professional Summary')
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(resume_data['summary'])
        summary_run.font.size = Pt(10.5)
        summary_run.font.name = 'Calibri'
        summary_run.font.color.rgb = text_color
        summary_para.space_after = Pt(6)
    
    # Experience
    if resume_data.get('experience'):
        add_section_header('Professional Experience')
        
        for i, job in enumerate(resume_data['experience']):
            # Job title and company
            job_para = doc.add_paragraph()
            job_para.space_before = Pt(8) if i > 0 else Pt(2)
            job_para.space_after = Pt(2)
            
            title_run = job_para.add_run(job.get('title', ''))
            title_run.font.bold = True
            title_run.font.size = Pt(11)
            title_run.font.name = 'Calibri'
            title_run.font.color.rgb = text_color
            
            # Company and dates on same line
            company_para = doc.add_paragraph()
            company_para.space_after = Pt(4)
            
            company_text = job.get('company', '')
            if job.get('location'):
                company_text += f", {job['location']}"
            
            company_run = company_para.add_run(company_text)
            company_run.font.size = Pt(10.5)
            company_run.font.name = 'Calibri'
            company_run.font.color.rgb = text_color
            
            # Add tab and dates
            if job.get('dates'):
                tab_stops = company_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.3), WD_ALIGN_PARAGRAPH.RIGHT)
                company_para.add_run('\t')
                dates_run = company_para.add_run(job['dates'])
                dates_run.font.size = Pt(10.5)
                dates_run.font.italic = True
                dates_run.font.name = 'Calibri'
                dates_run.font.color.rgb = text_color
            
            # Bullet points
            for bullet in job.get('bullets', []):
                add_bullet_point(bullet)
    
    # Education
    if resume_data.get('education'):
        add_section_header('Education')
        
        for edu in resume_data['education']:
            edu_para = doc.add_paragraph()
            edu_para.space_after = Pt(2)
            
            degree_run = edu_para.add_run(edu.get('degree', ''))
            degree_run.font.bold = True
            degree_run.font.size = Pt(11)
            degree_run.font.name = 'Calibri'
            
            school_para = doc.add_paragraph()
            school_para.space_after = Pt(4)
            
            school_text = edu.get('school', '')
            if edu.get('location'):
                school_text += f", {edu['location']}"
            
            school_run = school_para.add_run(school_text)
            school_run.font.size = Pt(10.5)
            school_run.font.name = 'Calibri'
            
            if edu.get('date'):
                tab_stops = school_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.3), WD_ALIGN_PARAGRAPH.RIGHT)
                school_para.add_run('\t')
                date_run = school_para.add_run(edu['date'])
                date_run.font.size = Pt(10.5)
                date_run.font.italic = True
                date_run.font.name = 'Calibri'
            
            if edu.get('details'):
                details_para = doc.add_paragraph()
                details_para.space_after = Pt(4)
                details_run = details_para.add_run(edu['details'])
                details_run.font.size = Pt(10)
                details_run.font.name = 'Calibri'
                details_run.font.color.rgb = text_color
    
    # Skills
    if resume_data.get('skills'):
        add_section_header('Skills & Certifications')
        skills = resume_data['skills']
        
        if skills.get('technical'):
            skills_para = doc.add_paragraph()
            skills_para.space_after = Pt(4)
            label_run = skills_para.add_run('Technical Skills: ')
            label_run.font.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = 'Calibri'
            skills_run = skills_para.add_run(', '.join(skills['technical']))
            skills_run.font.size = Pt(10.5)
            skills_run.font.name = 'Calibri'
        
        if skills.get('tools'):
            tools_para = doc.add_paragraph()
            tools_para.space_after = Pt(4)
            label_run = tools_para.add_run('Tools & Platforms: ')
            label_run.font.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = 'Calibri'
            tools_run = tools_para.add_run(', '.join(skills['tools']))
            tools_run.font.size = Pt(10.5)
            tools_run.font.name = 'Calibri'
        
        if skills.get('certifications'):
            certs_para = doc.add_paragraph()
            certs_para.space_after = Pt(4)
            label_run = certs_para.add_run('Certifications: ')
            label_run.font.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = 'Calibri'
            certs_run = certs_para.add_run(', '.join(skills['certifications']))
            certs_run.font.size = Pt(10.5)
            certs_run.font.name = 'Calibri'
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


# Main UI
st.markdown('<h1 class="main-header">📄 Resume Tailoring Agent</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">AI-powered resume optimization for ATS systems and keyword matching</p>', unsafe_allow_html=True)

# Sidebar configuration
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    
    api_key = st.text_input(
        "OpenAI API Key",
        type="password",
        help="Enter your OpenAI API key for AI-powered features"
    )
    
    st.markdown("---")
    st.markdown("### 📊 How It Works")
    st.markdown("""
    1. **Upload** your current resume
    2. **Enter** the job posting URL
    3. **Analyze** keyword matches
    4. **Generate** tailored resume
    5. **Create** cover letter
    6. **Prepare** for interviews
    7. **Download** all documents
    """)
    
    st.markdown("---")
    st.markdown("### 🎯 Features")
    st.markdown("""
    - ✅ ATS-Optimized Resume
    - ✅ Tailored Cover Letter
    - ✅ Behavioral Questions
    - ✅ Technical Questions
    - ✅ Coding Challenges
    - ✅ System Design Prep
    - ✅ Questions to Ask
    - ✅ Salary Negotiation Tips
    """)
    
    st.markdown("---")
    st.markdown("### 💡 Tips")
    st.info("""
    - Use a clean, text-readable resume format
    - Ensure job posting URL is accessible
    - Review AI suggestions for accuracy
    - Practice coding challenges before interviews
    """)

# Main content area
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown('<div class="section-header">📤 Upload Resume</div>', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Upload your resume (PDF or DOCX)",
        type=['pdf', 'docx'],
        help="Supported formats: PDF, DOCX"
    )
    
    if uploaded_file:
        with st.spinner("Extracting resume content..."):
            try:
                if uploaded_file.name.endswith('.pdf'):
                    st.session_state.resume_text = extract_text_from_pdf(uploaded_file)
                else:
                    st.session_state.resume_text = extract_text_from_docx(uploaded_file)
                
                st.success(f"✅ Resume loaded: {len(st.session_state.resume_text)} characters")
                
                with st.expander("Preview Resume Text"):
                    st.text(st.session_state.resume_text[:2000] + "..." if len(st.session_state.resume_text) > 2000 else st.session_state.resume_text)
                    
            except Exception as e:
                st.error(f"Error reading file: {e}")

with col2:
    st.markdown('<div class="section-header">🔗 Job Posting URL</div>', unsafe_allow_html=True)
    
    job_url = st.text_input(
        "Enter job posting URL",
        placeholder="https://careers.company.com/job/12345",
        help="Paste the full URL to the job posting"
    )
    
    # Option to paste job description directly
    with st.expander("Or paste job description directly"):
        manual_job_text = st.text_area(
            "Job Description",
            height=200,
            placeholder="Paste the job description here if URL scraping doesn't work..."
        )
    
    if job_url and st.button("🔍 Fetch Job Posting", type="primary"):
        with st.spinner("Scraping job posting..."):
            job_data = scrape_job_posting(job_url)
            
            if 'error' in job_data:
                st.error(f"Error fetching job: {job_data['error']}")
                st.info("Try pasting the job description directly using the expander above.")
            else:
                st.session_state.job_data = job_data
                st.success(f"✅ Job posting loaded: {job_data.get('title', 'Unknown Position')}")
                
                with st.expander("Preview Job Posting"):
                    st.markdown(f"**Title:** {job_data.get('title', 'N/A')}")
                    st.text(job_data.get('content', '')[:2000] + "...")
    
    elif manual_job_text:
        st.session_state.job_data = {
            'title': 'Job Position',
            'content': manual_job_text,
            'url': 'Manual Entry'
        }
        st.info("Using manually entered job description")

# Analysis section
st.markdown("---")

if st.session_state.resume_text and st.session_state.job_data:
    st.markdown('<div class="section-header">📊 Keyword Analysis</div>', unsafe_allow_html=True)
    
    if st.button("🔬 Analyze Keywords", type="primary"):
        with st.spinner("Analyzing keywords and calculating match scores..."):
            nlp = load_nlp_model()
            
            # Extract keywords
            resume_keywords = extract_keywords(st.session_state.resume_text, nlp)
            job_keywords = extract_keywords(st.session_state.job_data['content'], nlp)
            
            # Analyze matches
            keyword_analysis = analyze_keyword_match(resume_keywords, job_keywords)
            st.session_state.keyword_analysis = keyword_analysis
            
            # Calculate similarity
            similarity_score = calculate_similarity_score(
                st.session_state.resume_text,
                st.session_state.job_data['content']
            )
            
            # Display metrics
            metric_cols = st.columns(4)
            
            with metric_cols[0]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{keyword_analysis['match_score']:.0f}%</div>
                    <div class="metric-label">Keyword Match</div>
                </div>
                """, unsafe_allow_html=True)
            
            with metric_cols[1]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{similarity_score}%</div>
                    <div class="metric-label">Content Similarity</div>
                </div>
                """, unsafe_allow_html=True)
            
            with metric_cols[2]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{len(keyword_analysis['matched'])}</div>
                    <div class="metric-label">Matched Keywords</div>
                </div>
                """, unsafe_allow_html=True)
            
            with metric_cols[3]:
                st.markdown(f"""
                <div class="metric-card">
                    <div class="metric-value">{len(keyword_analysis['missing'])}</div>
                    <div class="metric-label">Missing Keywords</div>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Display keyword details
            keyword_cols = st.columns(2)
            
            with keyword_cols[0]:
                st.markdown("**✅ Matched Keywords**")
                matched_html = ""
                for kw in keyword_analysis['matched'][:25]:
                    matched_html += f'<span class="keyword-badge keyword-matched">{kw}</span>'
                st.markdown(matched_html, unsafe_allow_html=True)
            
            with keyword_cols[1]:
                st.markdown("**❌ Missing Keywords (to incorporate)**")
                missing_html = ""
                for kw in keyword_analysis['missing'][:25]:
                    missing_html += f'<span class="keyword-badge keyword-missing">{kw}</span>'
                st.markdown(missing_html, unsafe_allow_html=True)

# Resume generation section
st.markdown("---")

if st.session_state.keyword_analysis and api_key:
    st.markdown('<div class="section-header">✨ Generate Tailored Resume</div>', unsafe_allow_html=True)
    
    if st.button("🚀 Generate ATS-Optimized Resume", type="primary"):
        with st.spinner("AI is tailoring your resume... This may take 30-60 seconds."):
            progress_bar = st.progress(0)
            
            progress_bar.progress(20)
            tailored_resume = tailor_resume_with_ai(
                st.session_state.resume_text,
                st.session_state.job_data,
                st.session_state.keyword_analysis,
                api_key
            )
            
            progress_bar.progress(60)
            
            if tailored_resume:
                st.session_state.tailored_resume = tailored_resume
                progress_bar.progress(100)
                
                st.markdown('<div class="success-box">✅ Resume successfully tailored!</div>', unsafe_allow_html=True)
                
                # Show keywords added
                if tailored_resume.get('keywords_added'):
                    st.markdown("**🎯 Keywords Successfully Incorporated:**")
                    added_html = ""
                    for kw in tailored_resume['keywords_added']:
                        added_html += f'<span class="keyword-badge keyword-added">{kw}</span>'
                    st.markdown(added_html, unsafe_allow_html=True)
                
                # Preview sections
                with st.expander("📋 Preview Tailored Resume", expanded=True):
                    st.markdown(f"### {tailored_resume.get('contact', {}).get('name', 'Name')}")
                    st.markdown(f"*{tailored_resume.get('summary', '')}*")
                    
                    st.markdown("#### Experience")
                    for job in tailored_resume.get('experience', []):
                        st.markdown(f"**{job.get('title')}** at {job.get('company')}")
                        for bullet in job.get('bullets', []):
                            st.markdown(f"• {bullet}")
                    
                    st.markdown("#### Skills")
                    skills = tailored_resume.get('skills', {})
                    if skills.get('technical'):
                        st.markdown(f"**Technical:** {', '.join(skills['technical'])}")
                    if skills.get('tools'):
                        st.markdown(f"**Tools:** {', '.join(skills['tools'])}")
                
                # Generate document
                st.markdown("---")
                st.markdown("### 📥 Download Your Resume")
                
                job_title = st.session_state.job_data.get('title', 'Position').replace(' ', '_')[:30]
                filename = f"Resume_Tailored_{job_title}.docx"
                
                docx_buffer = create_ats_resume_docx(
                    tailored_resume,
                    st.session_state.job_data.get('title', '')
                )
                
                st.download_button(
                    label="📄 Download Word Document (.docx)",
                    data=docx_buffer,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
                st.info("💡 Tip: Open in Word to make final adjustments and save as PDF for submission.")
            else:
                st.error("Failed to generate tailored resume. Please check your API key and try again.")

elif st.session_state.keyword_analysis and not api_key:
    st.markdown('<div class="warning-box">⚠️ Please enter your OpenAI API key in the sidebar to generate a tailored resume.</div>', unsafe_allow_html=True)

# Cover Letter and Interview Prep Tabs
st.markdown("---")

if st.session_state.tailored_resume and api_key:
    st.markdown('<div class="section-header">📝 Additional Documents</div>', unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["✉️ Cover Letter", "🎯 Interview Preparation"])
    
    with tab1:
        st.markdown("### Generate Tailored Cover Letter")
        
        col_tone1, col_tone2 = st.columns([1, 2])
        with col_tone1:
            tone = st.selectbox(
                "Select Tone",
                ["professional", "conversational", "confident", "enthusiastic"],
                help="Choose the tone for your cover letter"
            )
        
        if st.button("✉️ Generate Cover Letter", type="primary", key="gen_cover"):
            with st.spinner("Crafting your cover letter..."):
                cover_letter = generate_cover_letter(
                    st.session_state.tailored_resume,
                    st.session_state.job_data,
                    api_key,
                    tone
                )
                
                if cover_letter:
                    st.session_state.cover_letter = cover_letter
                    st.success("✅ Cover letter generated!")
        
        if st.session_state.cover_letter:
            with st.expander("📄 Preview Cover Letter", expanded=True):
                st.markdown(st.session_state.cover_letter.get('full_letter', ''))
            
            # Download button
            contact_info = st.session_state.tailored_resume.get('contact', {})
            job_title = st.session_state.job_data.get('title', 'Position').replace(' ', '_')[:30]
            
            cover_letter_docx = create_cover_letter_docx(
                st.session_state.cover_letter,
                contact_info,
                job_title
            )
            
            st.download_button(
                label="📥 Download Cover Letter (.docx)",
                data=cover_letter_docx,
                file_name=f"Cover_Letter_{job_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cover"
            )
    
    with tab2:
        st.markdown("### Interview Preparation Materials")
        st.markdown("Generate comprehensive interview prep including behavioral questions, technical questions, and coding challenges.")
        
        if st.button("🎯 Generate Interview Prep", type="primary", key="gen_interview"):
            with st.spinner("Creating comprehensive interview preparation guide... This may take 60-90 seconds."):
                progress = st.progress(0)
                progress.progress(30)
                
                interview_prep = generate_interview_prep(
                    st.session_state.tailored_resume,
                    st.session_state.job_data,
                    api_key
                )
                
                progress.progress(100)
                
                if interview_prep:
                    st.session_state.interview_prep = interview_prep
                    st.success("✅ Interview preparation guide generated!")
        
        if st.session_state.interview_prep:
            prep = st.session_state.interview_prep
            
            # Job Analysis
            if prep.get('job_analysis'):
                with st.expander("📋 Job Analysis", expanded=True):
                    analysis = prep['job_analysis']
                    
                    if analysis.get('key_responsibilities'):
                        st.markdown("**Key Responsibilities:**")
                        for resp in analysis['key_responsibilities']:
                            st.markdown(f"• {resp}")
                    
                    if analysis.get('required_skills'):
                        st.markdown("**Required Skills:**")
                        skills_html = ""
                        for skill in analysis['required_skills']:
                            skills_html += f'<span class="keyword-badge keyword-matched">{skill}</span>'
                        st.markdown(skills_html, unsafe_allow_html=True)
                    
                    if analysis.get('red_flags_to_address'):
                        st.markdown("**Areas to Prepare:**")
                        for flag in analysis['red_flags_to_address']:
                            st.warning(flag)
            
            # Behavioral Questions
            if prep.get('behavioral_questions'):
                with st.expander(f"🗣️ Behavioral Questions ({len(prep['behavioral_questions'])} questions)"):
                    for i, q in enumerate(prep['behavioral_questions'], 1):
                        st.markdown(f"**Q{i}: {q.get('question', '')}**")
                        if q.get('why_asked'):
                            st.caption(f"*Why Asked:* {q['why_asked']}")
                        if q.get('star_framework_tips'):
                            st.info(f"💡 **STAR Tips:** {q['star_framework_tips']}")
                        if q.get('sample_answer_outline'):
                            st.markdown(f"**Answer Outline:** {q['sample_answer_outline']}")
                        st.markdown("---")
            
            # Technical Questions
            if prep.get('technical_questions'):
                with st.expander(f"💻 Technical Questions ({len(prep['technical_questions'])} questions)"):
                    for i, q in enumerate(prep['technical_questions'], 1):
                        difficulty_colors = {"Easy": "🟢", "Medium": "🟡", "Hard": "🔴"}
                        diff_icon = difficulty_colors.get(q.get('difficulty', ''), "⚪")
                        
                        st.markdown(f"**Q{i}: {q.get('question', '')}** {diff_icon} {q.get('difficulty', '')}")
                        if q.get('topic'):
                            st.caption(f"Topic: {q['topic']}")
                        if q.get('key_points'):
                            st.markdown("**Key Points:**")
                            for point in q['key_points']:
                                st.markdown(f"• {point}")
                        if q.get('sample_answer'):
                            with st.container():
                                st.markdown("**Sample Answer:**")
                                st.markdown(q['sample_answer'])
                        st.markdown("---")
            
            # Coding Challenges
            if prep.get('coding_challenges'):
                with st.expander(f"🧩 Coding Challenges ({len(prep['coding_challenges'])} challenges)"):
                    for i, challenge in enumerate(prep['coding_challenges'], 1):
                        difficulty_colors = {"Easy": "🟢", "Medium": "🟡", "Hard": "🔴"}
                        diff_icon = difficulty_colors.get(challenge.get('difficulty', ''), "⚪")
                        
                        st.markdown(f"### Challenge {i}: {challenge.get('title', '')} {diff_icon}")
                        st.markdown(f"**Problem:** {challenge.get('description', '')}")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            if challenge.get('example_input'):
                                st.code(f"Input: {challenge['example_input']}", language=None)
                        with col2:
                            if challenge.get('example_output'):
                                st.code(f"Output: {challenge['example_output']}", language=None)
                        
                        if challenge.get('hints'):
                            with st.expander("💡 Hints"):
                                for hint in challenge['hints']:
                                    st.markdown(f"• {hint}")
                        
                        if challenge.get('solution_approach'):
                            st.markdown(f"**Approach:** {challenge['solution_approach']}")
                        
                        if challenge.get('python_solution'):
                            with st.expander("🐍 Python Solution"):
                                st.code(challenge['python_solution'], language='python')
                        
                        if challenge.get('time_complexity') or challenge.get('space_complexity'):
                            st.caption(f"⏱️ Time: {challenge.get('time_complexity', 'N/A')} | 💾 Space: {challenge.get('space_complexity', 'N/A')}")
                        
                        st.markdown("---")
            
            # System Design
            if prep.get('system_design_questions'):
                with st.expander(f"🏗️ System Design ({len(prep['system_design_questions'])} questions)"):
                    for i, q in enumerate(prep['system_design_questions'], 1):
                        st.markdown(f"**Q{i}: {q.get('question', '')}**")
                        if q.get('key_components'):
                            st.markdown("**Key Components:**")
                            for comp in q['key_components']:
                                st.markdown(f"• {comp}")
                        if q.get('discussion_points'):
                            st.markdown("**Discussion Points:**")
                            for point in q['discussion_points']:
                                st.markdown(f"• {point}")
                        if q.get('follow_up_questions'):
                            st.markdown("**Possible Follow-ups:**")
                            for follow in q['follow_up_questions']:
                                st.markdown(f"• {follow}")
                        st.markdown("---")
            
            # Questions to Ask
            if prep.get('questions_to_ask'):
                with st.expander("❓ Questions to Ask the Interviewer"):
                    for q in prep['questions_to_ask']:
                        st.markdown(f"**{q.get('question', '')}**")
                        if q.get('why_effective'):
                            st.caption(f"*Why Effective:* {q['why_effective']}")
                        st.markdown("")
            
            # Salary Negotiation
            if prep.get('salary_negotiation'):
                with st.expander("💰 Salary Negotiation Tips"):
                    salary = prep['salary_negotiation']
                    if salary.get('market_range'):
                        st.success(f"**Market Range:** {salary['market_range']}")
                    if salary.get('negotiation_tips'):
                        st.markdown("**Negotiation Tips:**")
                        for tip in salary['negotiation_tips']:
                            st.markdown(f"• {tip}")
                    if salary.get('value_propositions'):
                        st.markdown("**Your Value Propositions:**")
                        for prop in salary['value_propositions']:
                            st.info(prop)
            
            # Download full prep document
            st.markdown("---")
            job_title = st.session_state.job_data.get('title', 'Position').replace(' ', '_')[:30]
            
            interview_docx = create_interview_prep_docx(
                st.session_state.interview_prep,
                st.session_state.job_data.get('title', '')
            )
            
            st.download_button(
                label="📥 Download Complete Interview Prep Guide (.docx)",
                data=interview_docx,
                file_name=f"Interview_Prep_{job_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_interview"
            )

elif st.session_state.keyword_analysis and not st.session_state.tailored_resume:
    st.info("👆 Generate your tailored resume first to unlock Cover Letter and Interview Prep features.")

# Download All Section
if st.session_state.tailored_resume:
    st.markdown("---")
    st.markdown('<div class="section-header">📦 Download All Documents</div>', unsafe_allow_html=True)
    
    col_dl1, col_dl2, col_dl3 = st.columns(3)
    
    job_title = st.session_state.job_data.get('title', 'Position').replace(' ', '_')[:30]
    
    with col_dl1:
        docx_buffer = create_ats_resume_docx(
            st.session_state.tailored_resume,
            st.session_state.job_data.get('title', '')
        )
        st.download_button(
            label="📄 Resume",
            data=docx_buffer,
            file_name=f"Resume_{job_title}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_resume_final"
        )
    
    with col_dl2:
        if st.session_state.cover_letter:
            contact_info = st.session_state.tailored_resume.get('contact', {})
            cover_docx = create_cover_letter_docx(
                st.session_state.cover_letter,
                contact_info,
                job_title
            )
            st.download_button(
                label="✉️ Cover Letter",
                data=cover_docx,
                file_name=f"Cover_Letter_{job_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_cover_final"
            )
        else:
            st.button("✉️ Cover Letter", disabled=True, help="Generate cover letter first")
    
    with col_dl3:
        if st.session_state.interview_prep:
            interview_docx = create_interview_prep_docx(
                st.session_state.interview_prep,
                st.session_state.job_data.get('title', '')
            )
            st.download_button(
                label="🎯 Interview Prep",
                data=interview_docx,
                file_name=f"Interview_Prep_{job_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_interview_final"
            )
        else:
            st.button("🎯 Interview Prep", disabled=True, help="Generate interview prep first")

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; font-size: 0.9rem;">
    <p>Resume Tailoring Agent | Built with Streamlit & OpenAI</p>
    <p>Features: ATS Resume • Cover Letter • Behavioral Questions • Technical Questions • Coding Challenges</p>
    <p>Always review AI-generated content for accuracy before submitting.</p>
</div>
""", unsafe_allow_html=True)
