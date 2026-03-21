import streamlit as st
import anthropic
import requests
from bs4 import BeautifulSoup
import json
import re
import os
import subprocess
import tempfile
import time
from pathlib import Path

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ResumeIQ · Interview Suite",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── CSS ────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500&display=swap');

:root {
    --bg:      #0a0b0f;
    --surface: #13151c;
    --border:  #1e2130;
    --accent:  #6c63ff;
    --accent2: #00d4aa;
    --danger:  #ff4d6d;
    --warn:    #ffa940;
    --text:    #e8eaf0;
    --muted:   #6b7280;
}

html, body, [data-testid="stAppViewContainer"] {
    background: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif;
}
[data-testid="stHeader"] { background: transparent !important; }
h1,h2,h3,h4 { font-family: 'Syne', sans-serif !important; color: var(--text) !important; }

.hero {
    text-align: center;
    padding: 2.5rem 1rem 1.8rem;
    background: radial-gradient(ellipse 80% 60% at 50% -10%, rgba(108,99,255,0.22) 0%, transparent 70%);
    border-bottom: 1px solid var(--border);
    margin-bottom: 2rem;
}
.hero-badge {
    display: inline-block;
    background: rgba(108,99,255,0.15);
    border: 1px solid rgba(108,99,255,0.4);
    color: #a29bfe;
    font-size: 0.7rem; font-weight: 600;
    letter-spacing: 0.15em; text-transform: uppercase;
    padding: 4px 14px; border-radius: 100px; margin-bottom: 1rem;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: clamp(1.8rem, 4vw, 3rem); font-weight: 800; line-height: 1.1;
    margin: 0 0 0.8rem;
    background: linear-gradient(135deg, #fff 30%, #a29bfe 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
}
.hero-sub { color: var(--muted); font-size: 1rem; max-width: 560px; margin: 0 auto; line-height: 1.6; }

/* Module nav pills */
.mod-nav { display: flex; gap: 0.6rem; flex-wrap: wrap; justify-content: center; margin: 1.5rem 0; }
.mod-pill {
    padding: 0.45rem 1.2rem; border-radius: 100px;
    border: 1px solid var(--border);
    background: var(--surface);
    color: var(--muted);
    font-family: 'Syne', sans-serif; font-size: 0.78rem; font-weight: 600;
    cursor: pointer; transition: all 0.15s;
}
.mod-pill.active {
    background: rgba(108,99,255,0.18);
    border-color: rgba(108,99,255,0.5);
    color: #a29bfe;
}

/* Cards */
.card {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 14px; padding: 1.4rem; margin-bottom: 1rem;
}
.card:hover { border-color: rgba(108,99,255,0.3); }

.section-hdr {
    font-family: 'Syne', sans-serif; font-size: 0.65rem; font-weight: 700;
    letter-spacing: 0.2em; text-transform: uppercase; color: var(--muted);
    margin: 1.2rem 0 0.6rem; padding-bottom: 0.35rem;
    border-bottom: 1px solid var(--border);
}

/* Pills */
.pill-row { display: flex; flex-wrap: wrap; gap: 0.45rem; margin: 0.5rem 0; }
.pill { display: inline-flex; align-items: center; gap: 5px; padding: 3px 11px; border-radius: 100px; font-size: 0.76rem; font-weight: 500; }
.pill-green  { background: rgba(0,212,170,0.12);  border: 1px solid rgba(0,212,170,0.35);  color: #00d4aa; }
.pill-red    { background: rgba(255,77,109,0.12);  border: 1px solid rgba(255,77,109,0.35);  color: #ff4d6d; }
.pill-yellow { background: rgba(255,169,64,0.12);  border: 1px solid rgba(255,169,64,0.35);  color: #ffa940; }
.pill-purple { background: rgba(108,99,255,0.12);  border: 1px solid rgba(108,99,255,0.35);  color: #a29bfe; }

/* Rec cards */
.rec-card {
    background: rgba(108,99,255,0.05); border: 1px solid rgba(108,99,255,0.18);
    border-left: 3px solid var(--accent); border-radius: 0 10px 10px 0;
    padding: 0.85rem 1rem; margin-bottom: 0.6rem; font-size: 0.88rem; line-height: 1.55;
}
.rec-title { font-weight: 600; color: #a29bfe; font-size: 0.78rem; text-transform: uppercase; letter-spacing: 0.08em; margin-bottom: 0.25rem; }

/* Chat bubbles */
.chat-wrap { display: flex; flex-direction: column; gap: 0.8rem; margin: 1rem 0; }
.bubble {
    max-width: 82%; padding: 0.85rem 1.1rem; border-radius: 14px;
    font-size: 0.9rem; line-height: 1.6;
}
.bubble-ai {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 4px 14px 14px 14px; align-self: flex-start;
    color: var(--text);
}
.bubble-user {
    background: rgba(108,99,255,0.18); border: 1px solid rgba(108,99,255,0.3);
    border-radius: 14px 4px 14px 14px; align-self: flex-end;
    color: #c8caff;
}
.bubble-label { font-size: 0.68rem; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; margin-bottom: 0.3rem; }
.bubble-ai .bubble-label { color: var(--accent2); }
.bubble-user .bubble-label { color: #a29bfe; text-align: right; }

/* Feedback card */
.feedback-card {
    background: rgba(0,212,170,0.06); border: 1px solid rgba(0,212,170,0.25);
    border-left: 3px solid var(--accent2); border-radius: 0 10px 10px 0;
    padding: 1rem 1.1rem; margin: 0.8rem 0;
}
.feedback-score {
    font-family: 'Syne', sans-serif; font-size: 2rem; font-weight: 800; color: var(--accent2);
}

/* Exam */
.q-card {
    background: var(--surface); border: 1px solid var(--border);
    border-radius: 12px; padding: 1.2rem 1.4rem; margin-bottom: 1rem;
}
.q-num { font-size: 0.68rem; font-weight: 700; color: var(--accent); letter-spacing: 0.15em; text-transform: uppercase; margin-bottom: 0.4rem; }
.q-text { font-size: 0.95rem; color: var(--text); font-weight: 500; margin-bottom: 0.8rem; line-height: 1.5; }
.answer-correct { background: rgba(0,212,170,0.1) !important; border-color: rgba(0,212,170,0.5) !important; }
.answer-wrong   { background: rgba(255,77,109,0.1) !important; border-color: rgba(255,77,109,0.5) !important; }

/* Score ring */
.score-big { font-family: 'Syne', sans-serif; font-size: 4rem; font-weight: 800; line-height: 1; }

/* Buttons */
.stButton > button {
    background: linear-gradient(135deg, var(--accent), #8b5cf6) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 600 !important;
    font-size: 0.88rem !important; padding: 0.6rem 1.6rem !important;
    transition: opacity 0.2s, transform 0.1s !important;
}
.stButton > button:hover { opacity: 0.85 !important; transform: translateY(-1px) !important; }
.stDownloadButton > button {
    background: linear-gradient(135deg, var(--accent2), #00b894) !important;
    color: #0a0b0f !important; border: none !important; border-radius: 10px !important;
    font-family: 'Syne', sans-serif !important; font-weight: 700 !important; font-size: 0.88rem !important;
}

/* Inputs */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: var(--surface) !important; border: 1px solid var(--border) !important;
    border-radius: 10px !important; color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 3px rgba(108,99,255,0.18) !important;
}
[data-testid="stFileUploader"] { background: var(--surface) !important; border: 1.5px dashed var(--border) !important; border-radius: 12px !important; }
.stProgress > div > div { background: var(--accent) !important; }
hr { border-color: var(--border) !important; margin: 1.8rem 0 !important; }
.stTabs [data-baseweb="tab-list"] { background: var(--surface); border-radius: 10px; padding: 4px; border: 1px solid var(--border); gap: 4px; }
.stTabs [data-baseweb="tab"] { border-radius: 8px !important; color: var(--muted) !important; font-family: 'Syne', sans-serif !important; font-weight: 600 !important; }
.stTabs [aria-selected="true"] { background: var(--accent) !important; color: white !important; }
details { background: var(--surface) !important; border-radius: 10px !important; border: 1px solid var(--border) !important; }
summary { color: var(--text) !important; }
.stAlert { border-radius: 10px !important; }
.resume-preview {
    background: var(--surface); border: 1px solid var(--border); border-radius: 12px;
    padding: 1.8rem; font-size: 0.86rem; line-height: 1.7; white-space: pre-wrap;
    max-height: 550px; overflow-y: auto;
}
/* Radio button styling */
.stRadio > div { gap: 0.4rem !important; }
.stRadio label { color: var(--text) !important; }
</style>
""", unsafe_allow_html=True)

# ── Anthropic client ────────────────────────────────────────────────────────────
@st.cache_resource
def get_client():
    return anthropic.Anthropic()

# ══════════════════════════════════════════════════════════════════════════════
# UTILITY FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def extract_resume_text(uploaded_file) -> str:
    """Extract plain text from .docx or .txt — no pandoc required."""
    suffix = Path(uploaded_file.name).suffix.lower()
    if suffix == ".docx":
        # Try pandoc first (best quality), fall back to python-docx
        try:
            import shutil
            if shutil.which("pandoc"):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                    f.write(uploaded_file.getvalue())
                    tmp_path = f.name
                try:
                    result = subprocess.run(
                        ["pandoc", tmp_path, "-t", "plain"],
                        capture_output=True, text=True, timeout=15
                    )
                    if result.returncode == 0 and result.stdout.strip():
                        return result.stdout.strip()
                finally:
                    os.unlink(tmp_path)
        except Exception:
            pass
        # python-docx fallback (works on Windows without pandoc)
        try:
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(uploaded_file.getvalue()))
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
            return "\n".join(lines)
        except Exception as e:
            raise RuntimeError(f"Could not read .docx file: {e}")
    else:
        # Plain text
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

def scrape_job_description(url: str) -> str:
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script","style","nav","footer","header","aside"]):
            tag.decompose()
        lines = [l.strip() for l in soup.get_text(separator="\n").splitlines() if l.strip()]
        return "\n".join(lines)[:8000]
    except Exception as e:
        return f"ERROR: {e}"

def score_color(score):
    if score >= 70: return "#00d4aa"
    if score >= 45: return "#ffa940"
    return "#ff4d6d"

def clean_json(raw: str) -> dict:
    raw = re.sub(r"^```json\s*", "", raw.strip())
    raw = re.sub(r"^```\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)

def build_docx(resume_text: str) -> bytes:
    """Convert plain text resume to .docx using python-docx (no pandoc needed)."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()

        # Narrow margins for a clean resume look
        for section in doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.9)
            section.right_margin  = Inches(0.9)

        for line in resume_text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue

            # ALL CAPS lines → section header style
            if stripped == stripped.upper() and len(stripped) > 3 and not stripped[0].isdigit():
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
                # Underline as divider
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
            # Bullet lines
            elif stripped.startswith("- ") or stripped.startswith("• "):
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(stripped.lstrip("-•").strip())
                run.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception:
        # Last resort: return UTF-8 bytes (user gets a .docx that opens as text)
        return resume_text.encode("utf-8")

def _build_docx_pandoc_unused(resume_text: str) -> bytes:
    """Kept for reference — pandoc path not used on Windows."""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as f:
        f.write(resume_text)
        txt_path = f.name
    out_path = txt_path.replace(".txt", ".docx")
    try:
        subprocess.run(["pandoc", txt_path, "-o", out_path], capture_output=True)
        with open(out_path, "rb") as fh:
            return fh.read()
    except Exception:
        return resume_text.encode("utf-8")
    finally:
        for p in [txt_path, out_path]:
            try: os.unlink(p)
            except: pass

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 1 — RESUME GAP ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

def run_analysis(resume_text, jd_text) -> dict:
    client = get_client()
    prompt = f"""You are an expert resume analyst. Analyze this resume vs job description. Return ONLY raw JSON — no markdown fences.

RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

Return this exact JSON:
{{
  "match_score": <integer 0-100>,
  "role_title": "<role title>",
  "company": "<company name>",
  "summary": "<2-3 sentence executive summary>",
  "strengths": ["<strength 1>", ...],
  "gaps": ["<gap 1>", ...],
  "partial_matches": ["<partial 1>", ...],
  "ats_keywords_missing": ["<kw1>", ...],
  "ats_keywords_present": ["<kw1>", ...],
  "recommendations": [{{"area": "<area>", "suggestion": "<suggestion>"}}, ...],
  "rejection_risks": ["<risk 1>", ...]
}}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)

def generate_tailored_resume(resume_text, jd_text, analysis) -> str:
    gaps = "\n".join(f"- {g}" for g in analysis.get("gaps", []))
    kws  = ", ".join(analysis.get("ats_keywords_missing", []))
    recs = "\n".join(f"- [{r['area']}] {r['suggestion']}" for r in analysis.get("recommendations", []))
    prompt = f"""You are an expert resume writer. Rewrite this resume to align with the job description.
Keep all facts exact — no fabrication. Reframe bullets, weave in missing keywords naturally, rewrite the summary.
Output ONLY the resume text, no commentary. Use ALL CAPS for section headers, dashes for bullets.

ORIGINAL RESUME:
{resume_text}

JOB DESCRIPTION:
{jd_text}

GAPS: {gaps}
MISSING KEYWORDS: {kws}
RECOMMENDATIONS: {recs}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    return msg.content[0].text.strip()

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 2 — BEHAVIORAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════

BEHAVIORAL_SYSTEM = """You are an expert behavioral interview coach specializing in STAR method responses.
You have deep knowledge of the job description and the candidate's resume provided to you.

Your role:
1. Ask ONE focused behavioral question at a time, tailored to the specific role and company
2. After the candidate responds, give structured STAR feedback:
   - Score the response 1-10
   - Identify what was strong
   - Point out what was missing (Situation/Task/Action/Result gaps)
   - Give a concrete suggestion to strengthen the answer
3. Then ask the next question or offer a follow-up
4. Vary question themes: leadership, conflict, failure, collaboration, achievement, adaptability

Always be encouraging but honest. Keep responses concise and actionable.
Format feedback clearly with sections: SCORE, STRENGTHS, GAPS, SUGGESTION.
Then end with: "Ready for the next question? (yes/no)" or ask a follow-up."""

def behavioral_chat(messages, resume_text, jd_text) -> str:
    system = f"""{BEHAVIORAL_SYSTEM}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{jd_text[:3000]}"""

    # Build message history
    api_messages = []
    for m in messages:
        api_messages.append({"role": m["role"], "content": m["content"]})

    response = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1000,
        system=system,
        messages=api_messages
    )
    return response.content[0].text

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3 — TECHNICAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════

TECHNICAL_SYSTEM = """You are a senior technical interviewer conducting a rigorous technical interview.
You have the job description and candidate resume. Your questions should directly reflect the technical stack and requirements of the role.

Your role:
1. Ask ONE technical question at a time — ranging from conceptual to hands-on
2. Progress from foundational → intermediate → advanced as the conversation develops
3. After the candidate answers, give structured technical feedback:
   - SCORE: 1-10
   - CORRECT: what they got right
   - MISSED: important points they didn't cover
   - IDEAL ANSWER: brief version of what a strong answer looks like
   - FOLLOW-UP or next question

Question types to rotate through:
- Conceptual ("Explain how X works")
- Practical ("How would you approach building Y?")
- Debugging ("What's wrong with this approach?")
- Design ("How would you design a system for Z?")
- Code/SQL/Python snippets when relevant

Be rigorous but fair. If an answer is weak, say so clearly. If it's strong, acknowledge it."""

def technical_chat(messages, resume_text, jd_text) -> str:
    system = f"""{TECHNICAL_SYSTEM}

CANDIDATE RESUME:
{resume_text[:3000]}

JOB DESCRIPTION (focus your questions on this tech stack and requirements):
{jd_text[:3000]}"""

    api_messages = [{"role": m["role"], "content": m["content"]} for m in messages]
    response = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=1200,
        system=system,
        messages=api_messages
    )
    return response.content[0].text

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 4 — TECHNICAL EXAM
# ══════════════════════════════════════════════════════════════════════════════

def generate_exam(jd_text, resume_text) -> dict:
    prompt = f"""You are a technical hiring manager. Create a timed technical exam based on this job description.
Return ONLY raw JSON — no markdown fences.

JOB DESCRIPTION:
{jd_text}

CANDIDATE BACKGROUND (use this to calibrate difficulty):
{resume_text[:2000]}

Generate an exam with this exact JSON structure:
{{
  "title": "<Role> Technical Assessment",
  "duration_minutes": 30,
  "instructions": "<2 sentence exam instructions>",
  "sections": [
    {{
      "name": "Section 1: <topic>",
      "questions": [
        {{
          "id": "q1",
          "type": "multiple_choice",
          "question": "<question text>",
          "options": {{"A": "<option>", "B": "<option>", "C": "<option>", "D": "<option>"}},
          "correct": "A",
          "explanation": "<why this is correct and others are wrong>",
          "difficulty": "easy|medium|hard",
          "topic": "<topic tag>"
        }},
        {{
          "id": "q2",
          "type": "multiple_choice",
          "question": "<question>",
          "options": {{"A": "<option>", "B": "<option>", "C": "<option>", "D": "<option>"}},
          "correct": "B",
          "explanation": "<explanation>",
          "difficulty": "medium",
          "topic": "<topic>"
        }}
      ]
    }},
    {{
      "name": "Section 2: <topic>",
      "questions": [
        {{
          "id": "q6",
          "type": "short_answer",
          "question": "<question requiring a written answer>",
          "sample_answer": "<what a good answer looks like>",
          "key_points": ["<point 1>", "<point 2>", "<point 3>"],
          "difficulty": "medium",
          "topic": "<topic>"
        }}
      ]
    }}
  ]
}}

Requirements:
- 3 sections total covering the main technical areas from the JD
- Section 1: 5 multiple choice questions (fundamentals)
- Section 2: 5 multiple choice questions (intermediate/applied)
- Section 3: 3 short answer questions (design/scenario-based)
- Questions must be directly relevant to THIS specific job description
- Difficulty should be appropriate for the role level implied by the JD
- Make questions substantive and realistic — not trivial"""

    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)

def grade_short_answer(question: str, sample_answer: str, key_points: list, user_answer: str) -> dict:
    prompt = f"""Grade this short answer response for a technical interview exam.

QUESTION: {question}

CANDIDATE ANSWER: {user_answer}

SAMPLE ANSWER: {sample_answer}

KEY POINTS EXPECTED: {json.dumps(key_points)}

Return ONLY raw JSON:
{{
  "score": <integer 0-10>,
  "points_covered": ["<point covered>"],
  "points_missed": ["<point missed>"],
  "feedback": "<2-3 sentence constructive feedback>",
  "grade": "Excellent|Good|Partial|Needs Work"
}}"""
    msg = get_client().messages.create(
        model="claude-sonnet-4-20250514", max_tokens=600,
        messages=[{"role": "user", "content": prompt}]
    )
    return clean_json(msg.content[0].text)

# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ══════════════════════════════════════════════════════════════════════════════
defaults = {
    "analysis": None, "tailored_resume": None,
    "resume_text": None, "jd_text": None,
    "active_module": "resume",
    # Behavioral
    "behavioral_messages": [],
    "behavioral_started": False,
    # Technical
    "technical_messages": [],
    "technical_started": False,
    # Exam
    "exam_data": None,
    "exam_answers": {},
    "exam_sa_answers": {},
    "exam_submitted": False,
    "exam_sa_grades": {},
    "exam_started": False,
    "exam_start_time": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ══════════════════════════════════════════════════════════════════════════════
# HERO
# ══════════════════════════════════════════════════════════════════════════════
st.markdown("""
<div class="hero">
  <div class="hero-badge">AI-Powered · Career Intelligence Suite</div>
  <div class="hero-title">ResumeIQ</div>
  <div class="hero-sub">Resume analysis, tailored rewrites, behavioral coaching, technical interviews, and role-specific exams — all from one job URL.</div>
</div>
""", unsafe_allow_html=True)

# ── Module nav ─────────────────────────────────────────────────────────────────
modules = [
    ("resume",     "🎯 Resume Analyzer"),
    ("behavioral", "🗣️ Behavioral Interview"),
    ("technical",  "💻 Technical Interview"),
    ("exam",       "📝 Technical Exam"),
]

cols = st.columns(len(modules))
for col, (key, label) in zip(cols, modules):
    with col:
        if st.button(label, key=f"nav_{key}", use_container_width=True):
            st.session_state.active_module = key
            st.rerun()

# Highlight active module
active_labels = {k: l for k, l in modules}
st.markdown(f"""
<div style="text-align:center; margin: 0.3rem 0 1.5rem;">
  <span style="font-family:'Syne',sans-serif; font-size:0.75rem; color:#a29bfe; font-weight:600;">
    Active: {active_labels.get(st.session_state.active_module, '')}
  </span>
</div>
""", unsafe_allow_html=True)

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# SHARED INPUTS (always shown — needed for all modules)
# ══════════════════════════════════════════════════════════════════════════════
needs_inputs = not st.session_state.resume_text or not st.session_state.jd_text

if needs_inputs or st.session_state.active_module == "resume":
    with st.expander("⚙️ Setup — Resume & Job Description", expanded=needs_inputs):
        ic1, ic2 = st.columns(2, gap="large")
        with ic1:
            st.markdown('<div style="font-family:Syne,sans-serif;font-size:0.8rem;font-weight:700;color:#6c63ff;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:0.5rem;">Upload Resume</div>', unsafe_allow_html=True)
            uploaded = st.file_uploader("Resume (.docx or .txt)", type=["docx","txt"], label_visibility="collapsed")
            if uploaded:
                st.success(f"✓ {uploaded.name}")
        with ic2:
            st.markdown('<div style="font-family:Syne,sans-serif;font-size:0.8rem;font-weight:700;color:#6c63ff;text-transform:uppercase;letter-spacing:0.12em;margin-bottom:0.5rem;">Job Posting</div>', unsafe_allow_html=True)
            job_url = st.text_input("Job URL", placeholder="https://boards.greenhouse.io/...", label_visibility="collapsed")
            st.markdown('<div style="font-size:0.75rem;color:#6b7280;margin:0.3rem 0;">— or paste directly —</div>', unsafe_allow_html=True)
            jd_paste = st.text_area("Job description text", placeholder="Paste job description here...", height=100, label_visibility="collapsed")

        _, load_col, _ = st.columns([2,1,2])
        with load_col:
            load_btn = st.button("🔄 Load & Analyze", use_container_width=True)

        if load_btn:
            if not uploaded:
                st.error("Please upload your resume.")
            elif not job_url and not jd_paste:
                st.error("Please provide a job URL or paste the description.")
            else:
                with st.spinner("Reading resume…"):
                    st.session_state.resume_text = extract_resume_text(uploaded)
                jd_text = ""
                if job_url:
                    with st.spinner("Scraping job description…"):
                        jd_text = scrape_job_description(job_url)
                        if jd_text.startswith("ERROR:"):
                            st.warning("Couldn't scrape URL — using pasted text.")
                            jd_text = jd_paste
                if not jd_text:
                    jd_text = jd_paste
                if not jd_text:
                    st.error("Could not get job description.")
                else:
                    st.session_state.jd_text = jd_text
                    with st.spinner("Running gap analysis…"):
                        try:
                            st.session_state.analysis = run_analysis(
                                st.session_state.resume_text, jd_text
                            )
                            # Reset interview/exam state when new JD loaded
                            st.session_state.behavioral_messages = []
                            st.session_state.behavioral_started = False
                            st.session_state.technical_messages = []
                            st.session_state.technical_started = False
                            st.session_state.exam_data = None
                            st.session_state.exam_submitted = False
                            st.session_state.exam_answers = {}
                            st.session_state.exam_sa_answers = {}
                            st.session_state.exam_started = False
                            st.rerun()
                        except Exception as e:
                            st.error(f"Analysis error: {e}")

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 1: RESUME ANALYZER
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.active_module == "resume":
    if not st.session_state.analysis:
        st.markdown("""
        <div style="text-align:center;padding:3rem;color:#6b7280;">
          <div style="font-size:2.5rem;margin-bottom:1rem;">📄</div>
          <div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;color:#e8eaf0;margin-bottom:0.5rem;">Upload your resume and a job URL to get started</div>
          <div style="font-size:0.9rem;">Use the Setup panel above</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        a = st.session_state.analysis
        score = a.get("match_score", 0)
        color = score_color(score)

        # Score header
        st.markdown(f"""
        <div style="text-align:center;margin-bottom:2rem;">
          <div style="font-family:Syne,sans-serif;font-size:0.72rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;color:#6b7280;margin-bottom:0.5rem;">Match Analysis</div>
          <div style="font-family:Syne,sans-serif;font-size:1.3rem;font-weight:700;color:#e8eaf0;margin-bottom:0.2rem;">{a.get('role_title','Role')} · {a.get('company','Company')}</div>
          <div style="font-family:Syne,sans-serif;font-size:4.5rem;font-weight:800;color:{color};line-height:1;">{score}</div>
          <div style="color:#6b7280;font-size:0.82rem;">Match Score / 100</div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown(f"""<div class="rec-card"><div class="rec-title">Executive Summary</div>{a.get('summary','')}</div>""", unsafe_allow_html=True)

        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown('<div class="section-hdr">✅ Strengths</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">✓ {s}</span>' for s in a.get("strengths",[])) + '</div>', unsafe_allow_html=True)
        with c2:
            st.markdown('<div class="section-hdr">⚠️ Partial Matches</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-yellow">~ {s}</span>' for s in a.get("partial_matches",[])) + '</div>', unsafe_allow_html=True)
        with c3:
            st.markdown('<div class="section-hdr">❌ Gaps</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">✗ {s}</span>' for s in a.get("gaps",[])) + '</div>', unsafe_allow_html=True)

        st.markdown("---")
        k1, k2 = st.columns(2)
        with k1:
            st.markdown('<div class="section-hdr">🔍 ATS Keywords — Present</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">{k}</span>' for k in a.get("ats_keywords_present",[])) + '</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-hdr" style="margin-top:1rem;">Missing Keywords</div>', unsafe_allow_html=True)
            st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">{k}</span>' for k in a.get("ats_keywords_missing",[])) + '</div>', unsafe_allow_html=True)
        with k2:
            st.markdown('<div class="section-hdr">🚨 Rejection Risks</div>', unsafe_allow_html=True)
            for risk in a.get("rejection_risks",[]):
                st.markdown(f'<div class="rec-card" style="border-left-color:#ff4d6d;"><div class="rec-title" style="color:#ff4d6d;">Risk</div>{risk}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown('<div class="section-hdr">💡 Recommendations</div>', unsafe_allow_html=True)
        for rec in a.get("recommendations",[]):
            st.markdown(f'<div class="rec-card"><div class="rec-title">{rec.get("area","")}</div>{rec.get("suggestion","")}</div>', unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("""<div style="text-align:center;margin-bottom:1.2rem;"><div style="font-family:Syne,sans-serif;font-size:1.2rem;font-weight:700;color:#e8eaf0;">✨ Generate Tailored Resume</div><div style="color:#6b7280;font-size:0.85rem;margin-top:0.3rem;">Reframes your existing experience — no fabrication.</div></div>""", unsafe_allow_html=True)
        _, gc, _ = st.columns([2,1,2])
        with gc:
            if st.button("🚀 Generate Tailored Resume", use_container_width=True):
                with st.spinner("Crafting your tailored resume…"):
                    st.session_state.tailored_resume = generate_tailored_resume(
                        st.session_state.resume_text, st.session_state.jd_text, a
                    )

        if st.session_state.tailored_resume:
            st.markdown("---")
            t1, t2 = st.tabs(["📝 Tailored Resume", "📊 Side-by-Side"])
            with t1:
                st.markdown(f'<div class="resume-preview">{st.session_state.tailored_resume}</div>', unsafe_allow_html=True)
                st.markdown("<br>", unsafe_allow_html=True)
                d1, d2 = st.columns(2)
                with d1:
                    st.download_button("⬇️ Download .txt", data=st.session_state.tailored_resume.encode(), file_name=f"tailored_{a.get('company','role').replace(' ','_')}.txt", mime="text/plain", use_container_width=True)
                with d2:
                    st.download_button("⬇️ Download .docx", data=build_docx(st.session_state.tailored_resume), file_name=f"tailored_{a.get('company','role').replace(' ','_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            with t2:
                oc, tc = st.columns(2)
                with oc:
                    st.markdown("**Original**")
                    st.markdown(f'<div class="resume-preview">{st.session_state.resume_text}</div>', unsafe_allow_html=True)
                with tc:
                    st.markdown("**Tailored**")
                    st.markdown(f'<div class="resume-preview">{st.session_state.tailored_resume}</div>', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 2: BEHAVIORAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "behavioral":
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.4rem;">🗣️ Behavioral Interview Coach</div>
      <div style="color:#6b7280;font-size:0.9rem;">STAR-method coaching with role-specific questions drawn from the job description. Get scored feedback on every response.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # Tips
        with st.expander("📖 STAR Method Quick Reference"):
            sc1, sc2, sc3, sc4 = st.columns(4)
            for col, letter, name, desc in [
                (sc1, "S", "Situation", "Set the scene — context, timeframe, your role"),
                (sc2, "T", "Task", "What was required of you specifically?"),
                (sc3, "A", "Action", "What YOU did — be specific, use 'I' not 'we'"),
                (sc4, "R", "Result", "Quantified outcome — % improvement, time saved, etc."),
            ]:
                with col:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-family:Syne,sans-serif;font-size:1.8rem;font-weight:800;color:#6c63ff;">{letter}</div>
                      <div style="font-weight:600;color:#e8eaf0;font-size:0.85rem;">{name}</div>
                      <div style="color:#6b7280;font-size:0.78rem;margin-top:0.3rem;">{desc}</div>
                    </div>""", unsafe_allow_html=True)

        # Start / Reset
        bc1, bc2, bc3 = st.columns([1,1,3])
        with bc1:
            if st.button("▶️ Start Interview" if not st.session_state.behavioral_started else "🔄 New Session", use_container_width=True):
                st.session_state.behavioral_messages = []
                st.session_state.behavioral_started = True
                with st.spinner("Preparing your first question…"):
                    role = st.session_state.analysis.get("role_title","this role") if st.session_state.analysis else "this role"
                    company = st.session_state.analysis.get("company","the company") if st.session_state.analysis else "the company"
                    opener = f"Please start the behavioral interview. I'm interviewing for the {role} position at {company}. Ask me your first behavioral question."
                    st.session_state.behavioral_messages.append({"role":"user","content":opener})
                    reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                st.rerun()

        # Chat history
        if st.session_state.behavioral_started and st.session_state.behavioral_messages:
            st.markdown('<div class="section-hdr">Interview Session</div>', unsafe_allow_html=True)
            st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
            for msg in st.session_state.behavioral_messages:
                if msg["role"] == "assistant":
                    st.markdown(f"""<div class="bubble bubble-ai"><div class="bubble-label">Interview Coach</div>{msg["content"].replace(chr(10),"<br>")}</div>""", unsafe_allow_html=True)
                elif msg["role"] == "user" and not msg["content"].startswith("Please start"):
                    st.markdown(f"""<div class="bubble bubble-user"><div class="bubble-label">You</div>{msg["content"].replace(chr(10),"<br>")}</div>""", unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

            # Input
            st.markdown("<br>", unsafe_allow_html=True)
            user_input = st.text_area("Your answer", placeholder="Type your STAR response here...", height=140, key="behavioral_input", label_visibility="collapsed")
            ri1, ri2, _ = st.columns([1,1,3])
            with ri1:
                send_btn = st.button("📤 Send Response", use_container_width=True)
            with ri2:
                hint_btn = st.button("💡 Give me a hint", use_container_width=True)

            if send_btn and user_input.strip():
                st.session_state.behavioral_messages.append({"role":"user","content":user_input.strip()})
                with st.spinner("Analyzing your response…"):
                    reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                st.rerun()

            if hint_btn:
                hint_prompt = "I'm struggling with this question. Can you give me a brief hint about what you're looking for — without giving away the full answer?"
                st.session_state.behavioral_messages.append({"role":"user","content":hint_prompt})
                with st.spinner("Getting hint…"):
                    reply = behavioral_chat(st.session_state.behavioral_messages, st.session_state.resume_text, st.session_state.jd_text)
                st.session_state.behavioral_messages.append({"role":"assistant","content":reply})
                st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 3: TECHNICAL INTERVIEW AGENT
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "technical":
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.4rem;">💻 Technical Interview</div>
      <div style="color:#6b7280;font-size:0.9rem;">Questions drawn directly from the job's tech stack and requirements. Scored feedback with ideal answer examples after each response.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # Difficulty selector
        diff_col, _, _ = st.columns([1,2,2])
        with diff_col:
            difficulty = st.selectbox("Difficulty level", ["Progressive (recommended)", "Beginner", "Intermediate", "Senior"], key="tech_difficulty")

        tc1, tc2, _ = st.columns([1,1,3])
        with tc1:
            if st.button("▶️ Start Interview" if not st.session_state.technical_started else "🔄 New Session", use_container_width=True):
                st.session_state.technical_messages = []
                st.session_state.technical_started = True
                with st.spinner("Preparing first technical question…"):
                    role = st.session_state.analysis.get("role_title","this role") if st.session_state.analysis else "this role"
                    opener = f"Please start the technical interview for the {role} position. Difficulty preference: {difficulty}. Ask me your first technical question."
                    st.session_state.technical_messages.append({"role":"user","content":opener})
                    reply = technical_chat(st.session_state.technical_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.technical_messages.append({"role":"assistant","content":reply})
                st.rerun()

        if st.session_state.technical_started and st.session_state.technical_messages:
            st.markdown('<div class="section-hdr">Technical Interview Session</div>', unsafe_allow_html=True)
            st.markdown('<div class="chat-wrap">', unsafe_allow_html=True)
            for msg in st.session_state.technical_messages:
                if msg["role"] == "assistant":
                    st.markdown(f"""<div class="bubble bubble-ai"><div class="bubble-label">Technical Interviewer</div>{msg["content"].replace(chr(10),"<br>")}</div>""", unsafe_allow_html=True)
                elif msg["role"] == "user" and not msg["content"].startswith("Please start"):
                    st.markdown(f"""<div class="bubble bubble-user"><div class="bubble-label">You</div>{msg["content"].replace(chr(10),"<br>")}</div>""", unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)
            tech_input = st.text_area("Your answer", placeholder="Type your technical answer here. Include code if relevant...", height=160, key="technical_input", label_visibility="collapsed")

            ti1, ti2, ti3, _ = st.columns([1,1,1,2])
            with ti1:
                if st.button("📤 Submit Answer", use_container_width=True) and tech_input.strip():
                    st.session_state.technical_messages.append({"role":"user","content":tech_input.strip()})
                    with st.spinner("Evaluating response…"):
                        reply = technical_chat(st.session_state.technical_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.technical_messages.append({"role":"assistant","content":reply})
                    st.rerun()
            with ti2:
                if st.button("⏭️ Skip Question", use_container_width=True):
                    skip_msg = "I'd like to skip this question and move to the next one."
                    st.session_state.technical_messages.append({"role":"user","content":skip_msg})
                    with st.spinner("Moving on…"):
                        reply = technical_chat(st.session_state.technical_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.technical_messages.append({"role":"assistant","content":reply})
                    st.rerun()
            with ti3:
                if st.button("🔍 Show Ideal Answer", use_container_width=True):
                    show_msg = "Please show me the ideal answer for the last question you asked, with a full explanation."
                    st.session_state.technical_messages.append({"role":"user","content":show_msg})
                    with st.spinner("Preparing ideal answer…"):
                        reply = technical_chat(st.session_state.technical_messages, st.session_state.resume_text, st.session_state.jd_text)
                    st.session_state.technical_messages.append({"role":"assistant","content":reply})
                    st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# MODULE 4: TECHNICAL EXAM
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.active_module == "exam":
    st.markdown("""
    <div style="margin-bottom:1.5rem;">
      <div style="font-family:Syne,sans-serif;font-size:1.6rem;font-weight:800;color:#e8eaf0;margin-bottom:0.4rem;">📝 Technical Exam</div>
      <div style="color:#6b7280;font-size:0.9rem;">A timed, role-specific exam generated from the job description. Multiple choice + short answer. Scored with full explanations.</div>
    </div>
    """, unsafe_allow_html=True)

    if not st.session_state.resume_text or not st.session_state.jd_text:
        st.warning("Please load your resume and job description first using the Setup panel above.")
    else:
        # Generate exam
        if not st.session_state.exam_data:
            _, gen_col, _ = st.columns([2,1,2])
            with gen_col:
                if st.button("⚡ Generate Exam", use_container_width=True):
                    with st.spinner("Building your custom exam from the job description…"):
                        try:
                            st.session_state.exam_data = generate_exam(st.session_state.jd_text, st.session_state.resume_text)
                            st.session_state.exam_answers = {}
                            st.session_state.exam_sa_answers = {}
                            st.session_state.exam_submitted = False
                            st.session_state.exam_started = False
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to generate exam: {e}")
        else:
            exam = st.session_state.exam_data

            # Exam header
            if not st.session_state.exam_submitted:
                eh1, eh2, eh3 = st.columns([3,1,1])
                with eh1:
                    st.markdown(f"""<div style="font-family:Syne,sans-serif;font-size:1.1rem;font-weight:700;color:#e8eaf0;">{exam.get('title','Technical Assessment')}</div>
                    <div style="color:#6b7280;font-size:0.82rem;margin-top:0.3rem;">{exam.get('instructions','')}</div>""", unsafe_allow_html=True)
                with eh2:
                    total_q = sum(len(s.get("questions",[])) for s in exam.get("sections",[]))
                    answered = len(st.session_state.exam_answers) + len({k for k,v in st.session_state.exam_sa_answers.items() if v.strip()})
                    st.metric("Progress", f"{answered}/{total_q}")
                with eh3:
                    st.metric("Duration", f"{exam.get('duration_minutes',30)} min")

                if not st.session_state.exam_started:
                    _, sc, _ = st.columns([2,1,2])
                    with sc:
                        if st.button("▶️ Begin Exam", use_container_width=True):
                            st.session_state.exam_started = True
                            st.session_state.exam_start_time = time.time()
                            st.rerun()
                else:
                    # Timer display
                    elapsed = int(time.time() - (st.session_state.exam_start_time or time.time()))
                    limit   = exam.get("duration_minutes", 30) * 60
                    remain  = max(0, limit - elapsed)
                    mins, secs = divmod(remain, 60)
                    timer_color = "#ff4d6d" if remain < 300 else "#ffa940" if remain < 600 else "#00d4aa"
                    st.markdown(f"""<div style="text-align:right;font-family:Syne,sans-serif;font-size:1rem;font-weight:700;color:{timer_color};margin-bottom:1rem;">⏱ {mins:02d}:{secs:02d} remaining</div>""", unsafe_allow_html=True)

                    # Questions
                    for section in exam.get("sections", []):
                        st.markdown(f'<div class="section-hdr">{section.get("name","Section")}</div>', unsafe_allow_html=True)

                        for q in section.get("questions", []):
                            qid   = q.get("id","")
                            qtype = q.get("type","multiple_choice")
                            diff_badge = {"easy":"🟢","medium":"🟡","hard":"🔴"}.get(q.get("difficulty","medium"),"🟡")

                            st.markdown(f"""<div class="q-card">
                              <div class="q-num">{qid.upper()}  ·  {q.get('topic','')}  {diff_badge} {q.get('difficulty','').capitalize()}</div>
                              <div class="q-text">{q.get('question','')}</div>""", unsafe_allow_html=True)

                            if qtype == "multiple_choice":
                                opts = q.get("options", {})
                                choices = [f"{k}: {v}" for k, v in opts.items()]
                                current = st.session_state.exam_answers.get(qid)
                                idx = None
                                if current:
                                    try: idx = choices.index(next(c for c in choices if c.startswith(current)))
                                    except: idx = None
                                answer = st.radio("Select answer", choices, index=idx, key=f"mc_{qid}", label_visibility="collapsed")
                                if answer:
                                    st.session_state.exam_answers[qid] = answer[0]  # store just the letter
                            else:
                                sa_val = st.session_state.exam_sa_answers.get(qid, "")
                                new_val = st.text_area("Your answer", value=sa_val, height=100, key=f"sa_{qid}", label_visibility="collapsed", placeholder="Type your answer here...")
                                st.session_state.exam_sa_answers[qid] = new_val

                            st.markdown("</div>", unsafe_allow_html=True)

                    # Submit
                    st.markdown("<br>", unsafe_allow_html=True)
                    _, sub_col, _ = st.columns([2,1,2])
                    with sub_col:
                        if st.button("✅ Submit Exam", use_container_width=True):
                            st.session_state.exam_submitted = True
                            # Grade short answers
                            for section in exam.get("sections",[]):
                                for q in section.get("questions",[]):
                                    if q.get("type") == "short_answer":
                                        qid = q.get("id","")
                                        user_ans = st.session_state.exam_sa_answers.get(qid,"")
                                        if user_ans.strip():
                                            with st.spinner(f"Grading {qid}…"):
                                                grade = grade_short_answer(
                                                    q.get("question",""),
                                                    q.get("sample_answer",""),
                                                    q.get("key_points",[]),
                                                    user_ans
                                                )
                                                st.session_state.exam_sa_grades[qid] = grade
                            st.rerun()

            # ── RESULTS ───────────────────────────────────────────────────────
            else:
                exam = st.session_state.exam_data
                mc_correct = 0
                mc_total   = 0
                sa_score   = 0
                sa_total   = 0

                for section in exam.get("sections",[]):
                    for q in section.get("questions",[]):
                        if q.get("type") == "multiple_choice":
                            mc_total += 1
                            if st.session_state.exam_answers.get(q.get("id","")) == q.get("correct",""):
                                mc_correct += 1
                        elif q.get("type") == "short_answer":
                            grade = st.session_state.exam_sa_grades.get(q.get("id",""),{})
                            sa_score += grade.get("score",0)
                            sa_total += 10

                total_possible = mc_total * 10 + sa_total
                total_earned   = mc_correct * 10 + sa_score
                overall_pct    = int(total_earned / total_possible * 100) if total_possible > 0 else 0
                grade_letter   = "A" if overall_pct >= 90 else "B" if overall_pct >= 80 else "C" if overall_pct >= 70 else "D" if overall_pct >= 60 else "F"
                res_color      = score_color(overall_pct)

                st.markdown(f"""
                <div style="text-align:center;padding:2rem 0 1.5rem;">
                  <div style="font-family:Syne,sans-serif;font-size:0.72rem;font-weight:700;letter-spacing:0.18em;text-transform:uppercase;color:#6b7280;margin-bottom:0.5rem;">Exam Results</div>
                  <div style="font-family:Syne,sans-serif;font-size:5rem;font-weight:800;color:{res_color};line-height:1;">{overall_pct}%</div>
                  <div style="font-family:Syne,sans-serif;font-size:1.5rem;font-weight:700;color:{res_color};">Grade: {grade_letter}</div>
                  <div style="color:#6b7280;font-size:0.85rem;margin-top:0.5rem;">{total_earned}/{total_possible} points</div>
                </div>
                """, unsafe_allow_html=True)

                rm1, rm2, rm3 = st.columns(3)
                with rm1:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Multiple Choice</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{score_color(int(mc_correct/mc_total*100) if mc_total else 0)};">{mc_correct}/{mc_total}</div>
                    </div>""", unsafe_allow_html=True)
                with rm2:
                    sa_pct = int(sa_score/sa_total*100) if sa_total else 0
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Short Answer</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{score_color(sa_pct)};">{sa_score}/{sa_total}</div>
                    </div>""", unsafe_allow_html=True)
                with rm3:
                    st.markdown(f"""<div class="card" style="text-align:center;">
                      <div style="font-size:0.75rem;color:#6b7280;text-transform:uppercase;letter-spacing:0.1em;margin-bottom:0.3rem;">Overall Score</div>
                      <div style="font-family:Syne,sans-serif;font-size:2rem;font-weight:800;color:{res_color};">{overall_pct}%</div>
                    </div>""", unsafe_allow_html=True)

                st.markdown("---")
                st.markdown('<div class="section-hdr">📋 Detailed Review</div>', unsafe_allow_html=True)

                for section in exam.get("sections",[]):
                    st.markdown(f"**{section.get('name','')}**")
                    for q in section.get("questions",[]):
                        qid   = q.get("id","")
                        qtype = q.get("type","multiple_choice")

                        if qtype == "multiple_choice":
                            user_ans    = st.session_state.exam_answers.get(qid,"—")
                            correct_ans = q.get("correct","")
                            is_correct  = user_ans == correct_ans
                            border_col  = "#00d4aa" if is_correct else "#ff4d6d"
                            icon        = "✅" if is_correct else "❌"
                            opts        = q.get("options",{})
                            opts_html   = " · ".join(f"<strong>{k}</strong>: {v}" for k,v in opts.items())

                            st.markdown(f"""<div class="q-card" style="border-left:3px solid {border_col};">
                              <div class="q-num">{qid.upper()} — {icon} {'Correct' if is_correct else 'Incorrect'}</div>
                              <div class="q-text">{q.get('question','')}</div>
                              <div style="font-size:0.82rem;color:#6b7280;margin-bottom:0.5rem;">{opts_html}</div>
                              <div style="font-size:0.85rem;margin-bottom:0.3rem;">Your answer: <strong style="color:{'#00d4aa' if is_correct else '#ff4d6d'};">{user_ans}</strong> · Correct: <strong style="color:#00d4aa;">{correct_ans}</strong></div>
                              <div style="font-size:0.83rem;color:#a0a8b8;background:rgba(255,255,255,0.04);padding:0.6rem 0.8rem;border-radius:8px;">💡 {q.get('explanation','')}</div>
                            </div>""", unsafe_allow_html=True)

                        elif qtype == "short_answer":
                            grade    = st.session_state.exam_sa_grades.get(qid,{})
                            sa_score_q = grade.get("score",0)
                            grade_label = grade.get("grade","Not graded")
                            grade_color = score_color(sa_score_q * 10)

                            st.markdown(f"""<div class="q-card" style="border-left:3px solid {grade_color};">
                              <div class="q-num">{qid.upper()} — Short Answer · {grade_label} ({sa_score_q}/10)</div>
                              <div class="q-text">{q.get('question','')}</div>""", unsafe_allow_html=True)

                            st.markdown("**Your answer:**")
                            st.markdown(f'<div style="background:rgba(255,255,255,0.04);padding:0.7rem 0.9rem;border-radius:8px;font-size:0.85rem;color:#c8cad4;margin-bottom:0.6rem;">{st.session_state.exam_sa_answers.get(qid,"No answer provided")}</div>', unsafe_allow_html=True)

                            if grade:
                                covered = grade.get("points_covered",[])
                                missed  = grade.get("points_missed",[])
                                if covered:
                                    st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-green">✓ {p}</span>' for p in covered) + '</div>', unsafe_allow_html=True)
                                if missed:
                                    st.markdown('<div class="pill-row">' + "".join(f'<span class="pill pill-red">✗ {p}</span>' for p in missed) + '</div>', unsafe_allow_html=True)
                                st.markdown(f'<div style="font-size:0.83rem;color:#a0a8b8;background:rgba(255,255,255,0.04);padding:0.6rem 0.8rem;border-radius:8px;margin-top:0.4rem;">💡 {grade.get("feedback","")}</div>', unsafe_allow_html=True)

                            with st.expander("📖 Sample answer"):
                                st.markdown(q.get("sample_answer",""))
                            st.markdown("</div>", unsafe_allow_html=True)

                # Retake
                st.markdown("<br>", unsafe_allow_html=True)
                rc1, rc2, _ = st.columns([1,1,3])
                with rc1:
                    if st.button("🔄 Retake Exam", use_container_width=True):
                        st.session_state.exam_answers = {}
                        st.session_state.exam_sa_answers = {}
                        st.session_state.exam_submitted = False
                        st.session_state.exam_sa_grades = {}
                        st.session_state.exam_started = False
                        st.rerun()
                with rc2:
                    if st.button("⚡ New Exam", use_container_width=True):
                        st.session_state.exam_data = None
                        st.session_state.exam_answers = {}
                        st.session_state.exam_sa_answers = {}
                        st.session_state.exam_submitted = False
                        st.session_state.exam_sa_grades = {}
                        st.session_state.exam_started = False
                        st.rerun()
