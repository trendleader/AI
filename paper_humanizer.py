import streamlit as st
import openai
import pdfplumber
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import time
import difflib
import json
import subprocess
import tempfile
import os

# ─── Page Config ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PaperHuman · AI Detection Remover",
    page_icon="✍️",
    layout="wide",
)

# ─── Custom CSS ─────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Mono:wght@400;500&family=DM+Sans:wght@300;400;500;600&display=swap');

html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    background-color: #0d0f14;
    color: #e8e4dc;
}
#MainMenu, footer, header {visibility: hidden;}
.block-container {padding-top: 2rem; padding-bottom: 2rem; max-width: 1200px;}

.app-header {
    text-align: center;
    padding: 2.5rem 0 1.5rem 0;
    border-bottom: 1px solid #2a2d36;
    margin-bottom: 2rem;
}
.app-title {
    font-family: 'DM Serif Display', serif;
    font-size: 3rem;
    color: #f5f0e8;
    letter-spacing: -0.5px;
    margin: 0;
    line-height: 1.1;
}
.app-title em { font-style: italic; color: #c8a96e; }
.app-subtitle {
    font-size: 1rem;
    color: #6b7080;
    margin-top: 0.5rem;
    font-weight: 300;
    letter-spacing: 0.5px;
}

/* Sidebar */
section[data-testid="stSidebar"] { background-color: #12151c !important; border-right: 1px solid #1e2029; }
section[data-testid="stSidebar"] .stTextInput > div > div > input {
    background-color: #1a1d26 !important; border: 1px solid #2a2d3a !important;
    color: #e8e4dc !important; font-family: 'DM Mono', monospace !important;
    font-size: 0.78rem !important; border-radius: 6px !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 { color: #c8a96e !important; }

/* Cards */
.info-card {
    background: #14171f; border: 1px solid #1e2230;
    border-radius: 10px; padding: 1.25rem 1.5rem; margin-bottom: 1rem;
}
.info-card h4 { font-family: 'DM Serif Display', serif; color: #c8a96e; margin: 0 0 0.5rem 0; font-size: 1rem; }
.info-card p { font-size: 0.88rem; color: #8a8f9e; margin: 0; line-height: 1.6; }

/* Score badges */
.score-badge {
    display: inline-block; padding: 0.4rem 1.1rem;
    border-radius: 100px; font-family: 'DM Mono', monospace;
    font-size: 0.85rem; font-weight: 500; margin-bottom: 1rem;
}
.score-high   { background: #1a2e1a; color: #5dba6e; border: 1px solid #2e5030; }
.score-medium { background: #2a1f0d; color: #e0a93c; border: 1px solid #5a3d10; }
.score-low    { background: #2a1018; color: #e05c6e; border: 1px solid #5a1828; }

/* Model badge */
.model-badge {
    display: inline-block; padding: 0.2rem 0.7rem;
    border-radius: 6px; font-family: 'DM Mono', monospace;
    font-size: 0.72rem; font-weight: 500;
    background: #1a1d2a; color: #7a8aaa;
    border: 1px solid #2a2d3a; margin-left: 0.5rem;
}
.model-cheap { color: #5dba6e; border-color: #2e5030; background: #1a2e1a; }
.model-smart { color: #c8a96e; border-color: #5a3d10; background: #2a1f0d; }

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: #12151c; border-radius: 8px; padding: 4px; gap: 2px; border: 1px solid #1e2029;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 6px; color: #6b7080 !important;
    font-family: 'DM Sans', sans-serif; font-weight: 500; font-size: 0.85rem; padding: 0.4rem 1.2rem;
}
.stTabs [aria-selected="true"] { background-color: #1e2230 !important; color: #c8a96e !important; }

/* Text areas */
.stTextArea textarea {
    background-color: #12151c !important; border: 1px solid #2a2d3a !important;
    color: #e8e4dc !important; font-family: 'DM Mono', monospace !important;
    font-size: 0.82rem !important; border-radius: 8px !important; line-height: 1.7 !important;
}

/* Buttons */
.stButton > button {
    background: #c8a96e !important; color: #0d0f14 !important;
    font-family: 'DM Sans', sans-serif !important; font-weight: 600 !important;
    font-size: 0.9rem !important; border: none !important;
    border-radius: 8px !important; padding: 0.55rem 2rem !important;
    letter-spacing: 0.3px !important; transition: all 0.2s ease !important;
}
.stButton > button:hover { background: #dbbf82 !important; transform: translateY(-1px); box-shadow: 0 4px 16px rgba(200,169,110,0.25) !important; }

/* File uploader */
[data-testid="stFileUploader"] { border: 2px dashed #2a2d3a !important; border-radius: 10px !important; padding: 1rem !important; background: #12151c !important; }

/* Progress */
.stProgress > div > div > div > div { background-color: #c8a96e !important; }

/* Selectbox */
.stSelectbox > div > div { background-color: #14171f !important; border: 1px solid #2a2d3a !important; color: #e8e4dc !important; border-radius: 8px !important; }

/* Metrics */
[data-testid="stMetric"] { background: #14171f; border: 1px solid #1e2230; border-radius: 8px; padding: 1rem; }
[data-testid="stMetricLabel"] { color: #6b7080 !important; font-size: 0.8rem !important; }
[data-testid="stMetricValue"] { color: #c8a96e !important; font-size: 1.4rem !important; font-family: 'DM Mono', monospace !important; }

/* Diff styling */
.diff-container {
    display: grid; grid-template-columns: 1fr 1fr; gap: 1rem;
    background: #0d0f14;
}
.diff-panel {
    background: #12151c; border: 1px solid #1e2230;
    border-radius: 10px; overflow: hidden;
}
.diff-header {
    padding: 0.6rem 1rem; border-bottom: 1px solid #1e2230;
    font-family: 'DM Mono', monospace; font-size: 0.78rem; color: #6b7080;
    display: flex; align-items: center; gap: 0.5rem;
}
.diff-content {
    padding: 1rem; font-family: 'DM Mono', monospace;
    font-size: 0.79rem; line-height: 1.75; white-space: pre-wrap;
    overflow-y: auto; max-height: 500px;
}
.diff-add    { background: #0d2e13; color: #7be08a; border-left: 3px solid #2d7a3a; padding-left: 0.4rem; }
.diff-remove { background: #2e0d13; color: #e07a82; border-left: 3px solid #7a2d3a; padding-left: 0.4rem; }
.diff-equal  { color: #8a8f9e; }

hr { border-color: #1e2029; }
.stSpinner > div { border-top-color: #c8a96e !important; }
.stAlert { border-radius: 8px !important; }
</style>
""", unsafe_allow_html=True)


# ─── Model Config ────────────────────────────────────────────────────────────────
MODEL_OPTIONS = {
    "GPT-4o  (Best quality)": "gpt-4o",
    "GPT-4o Mini  (Faster & cheaper)": "gpt-4o-mini",
    "GPT-3.5 Turbo  (Most economical)": "gpt-3.5-turbo",
}

APPROX_COST = {
    "gpt-4o":       ("~$0.005", "~$0.015"),   # (input/1k, output/1k)
    "gpt-4o-mini":  ("~$0.0002", "~$0.001"),
    "gpt-3.5-turbo":("~$0.0005", "~$0.0015"),
}


# ─── Helpers ────────────────────────────────────────────────────────────────────
def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def split_into_chunks(text: str, chunk_size: int = 3000) -> list[str]:
    paragraphs = text.split("\n\n")
    chunks, current = [], ""
    for para in paragraphs:
        if len(current) + len(para) > chunk_size and current:
            chunks.append(current.strip())
            current = para
        else:
            current += "\n\n" + para if current else para
    if current.strip():
        chunks.append(current.strip())
    return chunks


def count_words(text: str) -> int:
    return len(re.findall(r'\b\w+\b', text))


def score_text_for_ai(client: openai.OpenAI, text_chunk: str, model: str) -> dict:
    system = """You are an expert AI-writing detector and academic writing analyst.
Analyze the given text and return ONLY a JSON object (no markdown, no backticks):
{
  "ai_probability": <int 0-100>,
  "ai_patterns": [<list of short strings describing detected AI patterns>],
  "human_patterns": [<list of short strings describing human-sounding elements>],
  "verdict": "<brief 1-sentence overall verdict>"
}
Common AI patterns: repetitive transitional phrases ("it is important to note", "in conclusion",
"delve", "it is worth noting"), overly symmetric structure, passive voice overuse, unnaturally
uniform sentence length, formulaic hedging, lack of specificity/personal voice, generic examples."""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": f"Analyze this text:\n\n{text_chunk}"}
        ],
        temperature=0.2,
        max_tokens=600,
    )
    raw = response.choices[0].message.content.strip()
    raw = re.sub(r"```json|```", "", raw).strip()
    try:
        return json.loads(raw)
    except Exception:
        return {
            "ai_probability": 50,
            "ai_patterns": ["Could not parse response"],
            "human_patterns": [],
            "verdict": "Analysis incomplete.",
        }


def humanize_chunk(client: openai.OpenAI, text_chunk: str, style: str, intensity: str, model: str) -> str:
    style_instructions = {
        "Academic":       "Maintain scholarly tone, formal vocabulary, and citation-ready style. Use disciplinary language authentically.",
        "Professional":   "Clear, confident business prose. First-person voice where appropriate. Concrete and direct.",
        "Casual / Blog":  "Conversational tone. Contractions fine. Occasional rhetorical questions, personal asides, humor.",
        "Student Essay":  "Sound like a thoughtful undergraduate or graduate student. Natural transitions, some uncertainty, personal opinions.",
    }
    intensity_instructions = {
        "Light — Fix Obvious AI Tells":   "Minimal changes. Vary sentence length, remove clichés, add one or two personal-sounding observations.",
        "Moderate — Rewrite for Flow":    "Rewrite sentences and paragraphs for natural rhythm. Add texture and specificity. Eliminate AI patterns.",
        "Aggressive — Full Humanization": "Fully rewrite in natural human voice. Break formulaic structure. Add personality, imperfection, and authenticity.",
    }
    system = f"""You are an expert editor specializing in transforming AI-generated text into authentic human writing.
Style target: {style_instructions.get(style, '')}
Intensity: {intensity_instructions.get(intensity, '')}

Rules:
- Preserve ALL factual content, arguments, and citations. Do NOT invent new facts.
- Vary sentence structure significantly — mix short punchy sentences with longer complex ones.
- Replace AI clichés ("it is worth noting", "in conclusion", "it is important to", "delve", "moreover") with natural alternatives.
- Use active voice predominantly. Add realistic hedging NOT AI hedging.
- The result should pass AI-detection tools as human-written.
- Return ONLY the rewritten text. No preamble, no commentary, no markdown."""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": f"Rewrite the following:\n\n{text_chunk}"}
        ],
        temperature=0.85,
        max_tokens=4000,
    )
    return response.choices[0].message.content.strip()


def build_diff_html(original: str, humanized: str) -> str:
    """Build a side-by-side HTML diff of two texts."""
    orig_lines = original.splitlines()
    new_lines  = humanized.splitlines()
    matcher = difflib.SequenceMatcher(None, orig_lines, new_lines, autojunk=False)

    orig_html, new_html = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            for line in orig_lines[i1:i2]:
                orig_html.append(f'<div class="diff-equal">{_esc(line) or "&nbsp;"}</div>')
                new_html.append(f'<div class="diff-equal">{_esc(line) or "&nbsp;"}</div>')
        elif tag == "replace":
            for line in orig_lines[i1:i2]:
                orig_html.append(f'<div class="diff-remove">{_esc(line) or "&nbsp;"}</div>')
            for line in new_lines[j1:j2]:
                new_html.append(f'<div class="diff-add">{_esc(line) or "&nbsp;"}</div>')
        elif tag == "delete":
            for line in orig_lines[i1:i2]:
                orig_html.append(f'<div class="diff-remove">{_esc(line) or "&nbsp;"}</div>')
        elif tag == "insert":
            for line in new_lines[j1:j2]:
                new_html.append(f'<div class="diff-add">{_esc(line) or "&nbsp;"}</div>')

    return "\n".join(orig_html), "\n".join(new_html)


def _esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_to_docx(original_text: str, humanized_text: str, filename: str, style: str, model: str) -> bytes:
    """Build a formatted DOCX with both versions using python-docx."""
    doc = Document()

    # Page margins
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin  = section.bottom_margin = Inches(1)

    # Title
    title = doc.add_heading("PaperHuman — Humanized Document", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2a)

    # Meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Style: {style}  |  Model: {model}  |  Source: {filename}")
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x88, 0x88, 0x99)

    doc.add_paragraph()

    # ── Humanized Version ──
    h1 = doc.add_heading("✍️  Humanized Version", level=2)
    h1.runs[0].font.color.rgb = RGBColor(0x2a, 0x1f, 0x0d)

    for para in humanized_text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(11)

    # Page break
    doc.add_page_break()

    # ── Original Version ──
    h2 = doc.add_heading("📄  Original Text", level=2)
    h2.runs[0].font.color.rgb = RGBColor(0x33, 0x22, 0x11)

    for para in original_text.split("\n\n"):
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── Session State Init ──────────────────────────────────────────────────────────
for key in ["extracted_text", "analysis_results", "humanized_text", "upload_name",
            "detect_model_used", "humanize_model_used"]:
    if key not in st.session_state:
        st.session_state[key] = None


# ─── Header ─────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
    <h1 class="app-title">Paper<em>Human</em></h1>
    <p class="app-subtitle">Detect & remove AI fingerprints · Side-by-side diff · Export to DOCX</p>
</div>
""", unsafe_allow_html=True)


# ─── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ Configuration")
    st.markdown("---")

    api_key = st.text_input(
        "OpenAI API Key",
        type="password",
        placeholder="sk-...",
        help="Used only in this session — never stored or logged."
    )
    if api_key:
        st.success("✓ API key entered", icon="🔑")
    else:
        st.warning("Enter your OpenAI key to begin.", icon="⚠️")

    st.markdown("---")
    st.markdown("### 🤖 Model Selection")

    detect_model_label = st.selectbox(
        "Detection Model",
        list(MODEL_OPTIONS.keys()),
        index=0,
        help="GPT-4o gives the most accurate AI pattern detection. Mini/3.5 are faster and cheaper."
    )
    detect_model = MODEL_OPTIONS[detect_model_label]

    humanize_model_label = st.selectbox(
        "Humanization Model",
        list(MODEL_OPTIONS.keys()),
        index=1,
        help="GPT-4o produces the most natural rewriting. Mini is a great balance of quality vs cost."
    )
    humanize_model = MODEL_OPTIONS[humanize_model_label]

    # Cost estimate display
    d_in, d_out = APPROX_COST[detect_model]
    h_in, h_out = APPROX_COST[humanize_model]
    st.markdown(f"""
    <div class="info-card">
        <h4>💰 Approx Cost / 1K tokens</h4>
        <p>
        <b>Detection</b> ({detect_model})<br>In: {d_in} · Out: {d_out}<br><br>
        <b>Humanize</b> ({humanize_model})<br>In: {h_in} · Out: {h_out}
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ✍️ Rewrite Settings")

    writing_style = st.selectbox(
        "Target Writing Style",
        ["Academic", "Professional", "Student Essay", "Casual / Blog"],
    )
    humanize_intensity = st.selectbox(
        "Humanization Intensity",
        ["Light — Fix Obvious AI Tells", "Moderate — Rewrite for Flow", "Aggressive — Full Humanization"],
        index=1,
    )

    st.markdown("---")
    st.markdown("""
    <div class="info-card">
        <h4>How it works</h4>
        <p>1. Upload your document<br>
        2. Detect AI risk per section<br>
        3. Humanize in your chosen style<br>
        4. Review the side-by-side diff<br>
        5. Download as DOCX or TXT</p>
    </div>
    """, unsafe_allow_html=True)


# ─── Upload Row ──────────────────────────────────────────────────────────────────
col_upload, col_gap, col_meta = st.columns([3, 0.2, 1.5])

with col_upload:
    st.markdown("#### 📄 Upload Your Paper")
    uploaded_file = st.file_uploader(
        "Drop a PDF, DOCX, or TXT file",
        type=["pdf", "docx", "txt"],
        label_visibility="collapsed"
    )
    if uploaded_file:
        file_bytes = uploaded_file.read()
        ext = uploaded_file.name.lower().split(".")[-1]
        with st.spinner("Reading document…"):
            try:
                if ext == "pdf":
                    text = extract_text_from_pdf(file_bytes)
                elif ext == "docx":
                    text = extract_text_from_docx(file_bytes)
                else:
                    text = extract_text_from_txt(file_bytes)
                st.session_state.extracted_text = text
                st.session_state.upload_name = uploaded_file.name
                st.session_state.analysis_results = None
                st.session_state.humanized_text   = None
            except Exception as e:
                st.error(f"Could not read file: {e}")

with col_meta:
    if st.session_state.extracted_text:
        wc     = count_words(st.session_state.extracted_text)
        chunks = split_into_chunks(st.session_state.extracted_text)
        st.metric("Words",    f"{wc:,}")
        st.metric("Chars",    f"{len(st.session_state.extracted_text):,}")
        st.metric("Sections", str(len(chunks)))


# ─── Tabs ───────────────────────────────────────────────────────────────────────
if st.session_state.extracted_text:
    st.markdown("---")
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🔍  AI Detection",
        "✍️  Humanize",
        "🔀  Side-by-Side Diff",
        "⬇️  Export",
        "📋  Original Text",
    ])

    # ── Tab 1: Detection ──────────────────────────────────────────────────────
    with tab1:
        st.markdown("#### AI Detection Analysis")
        col_desc, col_badge = st.columns([4, 1])
        with col_desc:
            st.markdown(f"Analyzing with **{detect_model_label.split('(')[0].strip()}**. "
                        "Each section is scored 0–100 for AI likelihood.")

        if not api_key:
            st.warning("Please enter your OpenAI API key in the sidebar.")
        else:
            if st.button("🔍 Run AI Detection", key="detect_btn"):
                client = openai.OpenAI(api_key=api_key)
                chunks = split_into_chunks(st.session_state.extracted_text)
                results = []
                progress    = st.progress(0, text="Analyzing…")
                status_text = st.empty()

                for i, chunk in enumerate(chunks):
                    status_text.markdown(f"*Analyzing section {i+1} of {len(chunks)}…*")
                    try:
                        result = score_text_for_ai(client, chunk, detect_model)
                        result["chunk_preview"] = chunk[:120] + "…"
                        result["chunk_index"]   = i + 1
                        results.append(result)
                    except Exception as e:
                        results.append({
                            "ai_probability": 0, "ai_patterns": [str(e)],
                            "human_patterns": [], "verdict": "Error.",
                            "chunk_preview": chunk[:120] + "…", "chunk_index": i + 1
                        })
                    progress.progress((i + 1) / len(chunks))
                    time.sleep(0.3)

                status_text.empty(); progress.empty()
                st.session_state.analysis_results   = results
                st.session_state.detect_model_used  = detect_model

            if st.session_state.analysis_results:
                results   = st.session_state.analysis_results
                avg_score = sum(r["ai_probability"] for r in results) / len(results)

                if avg_score >= 70:
                    badge_class, label = "score-low",    f"⚠️ High AI Risk — {avg_score:.0f}%"
                elif avg_score >= 40:
                    badge_class, label = "score-medium", f"⚡ Moderate AI Risk — {avg_score:.0f}%"
                else:
                    badge_class, label = "score-high",   f"✓ Low AI Risk — {avg_score:.0f}%"

                m_used = st.session_state.detect_model_used or detect_model
                st.markdown(
                    f'<span class="score-badge {badge_class}">{label}</span>'
                    f'<span class="model-badge">{m_used}</span>',
                    unsafe_allow_html=True
                )

                for r in results:
                    prob = r["ai_probability"]
                    color = "🔴" if prob >= 70 else "🟡" if prob >= 40 else "🟢"
                    with st.expander(f"{color} Section {r['chunk_index']} — {prob}% AI  |  {r['verdict']}"):
                        c1, c2 = st.columns(2)
                        with c1:
                            st.markdown("**🤖 AI Patterns**")
                            for p in r["ai_patterns"] or ["None detected"]:
                                st.markdown(f"- {p}")
                        with c2:
                            st.markdown("**🧑 Human Elements**")
                            for p in r["human_patterns"] or ["None detected"]:
                                st.markdown(f"- {p}")
                        st.caption(f"Preview: *{r['chunk_preview']}*")

    # ── Tab 2: Humanize ───────────────────────────────────────────────────────
    with tab2:
        st.markdown("#### Humanize Your Paper")
        st.markdown(
            f"Model: **{humanize_model_label.split('(')[0].strip()}** · "
            f"Style: **{writing_style}** · "
            f"Intensity: **{humanize_intensity.split('—')[0].strip()}**"
        )

        if not api_key:
            st.warning("Please enter your OpenAI API key in the sidebar.")
        else:
            if st.button("✍️ Humanize Paper", key="humanize_btn"):
                client = openai.OpenAI(api_key=api_key)
                chunks = split_into_chunks(st.session_state.extracted_text)
                humanized_parts = []
                progress    = st.progress(0, text="Rewriting…")
                status_text = st.empty()

                for i, chunk in enumerate(chunks):
                    status_text.markdown(f"*Rewriting section {i+1} of {len(chunks)}…*")
                    try:
                        rewritten = humanize_chunk(client, chunk, writing_style, humanize_intensity, humanize_model)
                        humanized_parts.append(rewritten)
                    except Exception as e:
                        humanized_parts.append(f"[Error section {i+1}: {e}]\n\n{chunk}")
                    progress.progress((i + 1) / len(chunks))
                    time.sleep(0.2)

                status_text.empty(); progress.empty()
                st.session_state.humanized_text      = "\n\n".join(humanized_parts)
                st.session_state.humanize_model_used = humanize_model
                st.success("✓ Humanization complete! Check the Diff and Export tabs.")

            if st.session_state.humanized_text:
                orig_wc = count_words(st.session_state.extracted_text)
                new_wc  = count_words(st.session_state.humanized_text)
                m1, m2, m3 = st.columns(3)
                m1.metric("Original Words",  f"{orig_wc:,}")
                m2.metric("Rewritten Words", f"{new_wc:,}")
                m3.metric("Δ Words",         f"{new_wc - orig_wc:+,}")

                st.markdown("**Humanized Output:**")
                st.text_area("", value=st.session_state.humanized_text, height=450, label_visibility="collapsed")

    # ── Tab 3: Side-by-Side Diff ──────────────────────────────────────────────
    with tab3:
        st.markdown("#### Side-by-Side Diff")
        if not st.session_state.humanized_text:
            st.info("Run humanization first to see the diff.")
        else:
            # Chunk selector
            chunks_orig = split_into_chunks(st.session_state.extracted_text)
            chunks_new  = split_into_chunks(st.session_state.humanized_text)
            n = min(len(chunks_orig), len(chunks_new))

            view_mode = st.radio(
                "View",
                ["Full document", "Section by section"],
                horizontal=True,
                label_visibility="collapsed"
            )

            if view_mode == "Section by section" and n > 1:
                section_idx = st.slider("Section", 1, n, 1) - 1
                orig_to_diff = chunks_orig[section_idx]
                new_to_diff  = chunks_new[section_idx]
            else:
                orig_to_diff = st.session_state.extracted_text
                new_to_diff  = st.session_state.humanized_text

            orig_html, new_html = build_diff_html(orig_to_diff, new_to_diff)

            st.markdown(f"""
            <div class="diff-container">
                <div class="diff-panel">
                    <div class="diff-header">
                        <span style="color:#e05c6e">●</span> Original
                    </div>
                    <div class="diff-content">{orig_html}</div>
                </div>
                <div class="diff-panel">
                    <div class="diff-header">
                        <span style="color:#5dba6e">●</span> Humanized
                        <span class="model-badge model-cheap">{st.session_state.humanize_model_used or humanize_model}</span>
                    </div>
                    <div class="diff-content">{new_html}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Stats
            st.markdown("---")
            orig_lines = orig_to_diff.splitlines()
            new_lines  = new_to_diff.splitlines()
            matcher    = difflib.SequenceMatcher(None, orig_lines, new_lines)
            added = deleted = changed = unchanged = 0
            for tag, i1, i2, j1, j2 in matcher.get_opcodes():
                if   tag == "equal":   unchanged += i2 - i1
                elif tag == "insert":  added     += j2 - j1
                elif tag == "delete":  deleted   += i2 - i1
                elif tag == "replace": changed   += max(i2 - i1, j2 - j1)

            s1, s2, s3, s4 = st.columns(4)
            s1.metric("Lines Unchanged", unchanged)
            s2.metric("Lines Added",     added)
            s3.metric("Lines Removed",   deleted)
            s4.metric("Lines Changed",   changed)

    # ── Tab 4: Export ─────────────────────────────────────────────────────────
    with tab4:
        st.markdown("#### Export Your Paper")

        if not st.session_state.humanized_text:
            st.info("Run humanization first to enable export.")
        else:
            base_name = st.session_state.upload_name.rsplit(".", 1)[0]
            m_used    = st.session_state.humanize_model_used or humanize_model

            col_e1, col_e2, col_e3 = st.columns(3)

            with col_e1:
                st.markdown("""
                <div class="info-card">
                    <h4>📝 Word Document (.docx)</h4>
                    <p>Formatted document with both humanized and original versions. Ready for editing in Word or Google Docs.</p>
                </div>
                """, unsafe_allow_html=True)
                try:
                    docx_bytes = export_to_docx(
                        st.session_state.extracted_text,
                        st.session_state.humanized_text,
                        st.session_state.upload_name,
                        writing_style,
                        m_used
                    )
                    st.download_button(
                        label="⬇️ Download DOCX",
                        data=docx_bytes,
                        file_name=f"{base_name}_humanized.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"DOCX export error: {e}")

            with col_e2:
                st.markdown("""
                <div class="info-card">
                    <h4>📄 Plain Text — Humanized</h4>
                    <p>Just the rewritten text. Paste into any editor, CMS, or submission portal.</p>
                </div>
                """, unsafe_allow_html=True)
                st.download_button(
                    label="⬇️ Download TXT (Humanized)",
                    data=st.session_state.humanized_text.encode("utf-8"),
                    file_name=f"{base_name}_humanized.txt",
                    mime="text/plain"
                )

            with col_e3:
                st.markdown("""
                <div class="info-card">
                    <h4>📄 Plain Text — Original</h4>
                    <p>The extracted original text. Useful as a reference copy alongside the rewrite.</p>
                </div>
                """, unsafe_allow_html=True)
                st.download_button(
                    label="⬇️ Download TXT (Original)",
                    data=st.session_state.extracted_text.encode("utf-8"),
                    file_name=f"{base_name}_original.txt",
                    mime="text/plain"
                )

            # Detection report export
            if st.session_state.analysis_results:
                st.markdown("---")
                st.markdown("##### 📊 Detection Report")
                results   = st.session_state.analysis_results
                avg_score = sum(r["ai_probability"] for r in results) / len(results)
                report_lines = [
                    f"PaperHuman Detection Report",
                    f"File: {st.session_state.upload_name}",
                    f"Model: {st.session_state.detect_model_used}",
                    f"Overall AI Risk Score: {avg_score:.1f}%",
                    "=" * 60,
                ]
                for r in results:
                    report_lines += [
                        f"\nSection {r['chunk_index']} — AI Score: {r['ai_probability']}%",
                        f"Verdict: {r['verdict']}",
                        "AI Patterns: " + ", ".join(r["ai_patterns"]),
                        "Human Elements: " + ", ".join(r["human_patterns"]),
                        f"Preview: {r['chunk_preview']}",
                    ]
                st.download_button(
                    label="⬇️ Download Detection Report (.txt)",
                    data="\n".join(report_lines).encode("utf-8"),
                    file_name=f"{base_name}_detection_report.txt",
                    mime="text/plain"
                )

    # ── Tab 5: Original ───────────────────────────────────────────────────────
    with tab5:
        st.markdown("#### Original Extracted Text")
        st.caption(f"`{st.session_state.upload_name}` · {count_words(st.session_state.extracted_text):,} words")
        st.text_area("", value=st.session_state.extracted_text, height=500, label_visibility="collapsed")

else:
    st.markdown("""
    <div style="text-align:center; padding: 4rem 2rem; color: #3a3d48;">
        <div style="font-size: 3rem; margin-bottom: 1rem;">📄</div>
        <div style="font-family: 'DM Serif Display', serif; font-size: 1.3rem; color: #4a4d58;">
            Upload a paper to get started
        </div>
        <div style="font-size: 0.85rem; margin-top: 0.5rem;">Supports PDF · DOCX · TXT</div>
    </div>
    """, unsafe_allow_html=True)
